"""
Generate KAU_FPO_DPR_RCD_KefiTech_v1.0.docx from the markdown source.

Styled to mirror the existing Phase 2 RCD (title block, section numbering, tables).
Reads: context/phase2/Dpr/DPR_Clarification_Requests.md
Writes: Documents/KAU_FPO_DPR_RCD_KefiTech_v1.0.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DOC = Document()

# ── Page margins ──────────────────────────────────────────────────────────
for section in DOC.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)


def add_p(text="", size=11, bold=False, italic=False, align=None, color=None, style=None):
    p = DOC.add_paragraph(style=style) if style else DOC.add_paragraph()
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color is not None:
            run.font.color.rgb = color
    return p


def add_heading(text, level=1):
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    add_p("", size=6)  # small spacer
    p = DOC.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(sizes.get(level, 11))
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x8A) if level <= 2 else RGBColor(0x33, 0x33, 0x33)


def add_bullet(text, size=10.5, indent_level=0):
    p = DOC.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + 0.6 * indent_level)
    run = p.add_run(text)
    run.font.size = Pt(size)


def add_table_row(table, cells, bold=False, header=False):
    row = table.add_row().cells
    for i, val in enumerate(cells):
        row[i].text = ""
        p = row[i].paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(9.5)
        run.font.bold = bold or header


# ── COVER ────────────────────────────────────────────────────────────────
add_p("KERALA AGRICULTURAL UNIVERSITY", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p("Communication Centre, Mannuthy, Thrissur — 680651", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p("")
add_p("AI-Assisted Detailed Project Report (DPR) Module", size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1F, 0x3A, 0x8A))
add_p("Requirements Clarification Document", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1F, 0x3A, 0x8A))
add_p("Version 1.0  |  Prepared by Kefi Tech Solutions Pvt. Ltd.", size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p("Confidential — Prepared in response to Tender Notice KAUCC/459/2025-C3", size=9.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p("")
add_p("")

# ── Addressee block ──────────────────────────────────────────────────────
add_p("Attention:", size=10.5, bold=True)
add_p("Dr. Meer Ahamed Ibrahim I. — Technical Consultant", size=10.5)
add_p("Dr. Binoo P. Bonny — Principal Investigator, KAU-FPO Linkage Programme", size=10.5)
add_p("")
add_p("Reference document: Business Requirements and Functional Specification for the AI-Assisted Detailed Project Report (DPR) Data Collection Module — Version 1.0, dated 17 July 2026", size=10, italic=True)
add_p("")

# ── 1. Purpose ───────────────────────────────────────────────────────────
add_heading("1. Purpose of this document", level=2)
add_p(
    "Following a detailed review of the AI-Assisted DPR Data Collection Module specification, "
    "Kefi Tech Solutions Pvt. Ltd. has identified certain items where clarification from "
    "Kerala Agricultural University will improve the accuracy of implementation and reduce the "
    "risk of rework during User Acceptance Testing.",
    size=10.5,
)
add_p("")
add_p(
    "The clarifications sought are grouped by criticality:",
    size=10.5,
)
add_bullet("Critical items block specific build phases and cannot be reasonably assumed by the software agency.")
add_bullet("Important items must be answered before the corresponding data element is implemented.")
add_bullet("Deferrable items can be answered during User Acceptance Testing.")
add_bullet("Section D items were discovered during implementation and are additions to the original clarification list.")

add_p("")
add_heading("2. How to Respond", level=2)
add_p("Please review each section and respond within 7 working days from the date of this document. For each item:", size=10.5)
add_bullet("If a decision is correct — confirm it.")
add_bullet("If a decision is incorrect — provide the correct information.")
add_bullet("If data is requested — provide it in the format described.")
add_bullet("If a question is asked — provide a clear answer.")
add_p("")
add_p("For a walkthrough call, contact: Athul Gopan — athul.gopan@kefitech.com", size=10.5)

DOC.add_page_break()

# ── SECTION A ────────────────────────────────────────────────────────────
add_heading("Section A — Critical Clarifications", level=1)
add_p(
    "These items directly block the architectural design and cannot be reasonably assumed by the software agency.",
    size=10.5, italic=True,
)
add_p("")

add_heading("A.1 — Component-to-Section Visibility Matrix", level=3)
add_p("Reference: Spec §1.7.2, §2.3.2, Ch 6.3–6.4 (Component-Based Dynamic Questionnaire)", size=10, italic=True)
add_p(
    "The specification states that the questionnaire shall dynamically display only those "
    "information blocks relevant to the selected Project Components. However, the mapping "
    "between each Project Component and the applicable Data Elements is not explicitly defined.",
    size=10.5,
)
add_p("Request:", size=10.5, bold=True)
add_p(
    "Please provide an authoritative matrix indicating, for each of the ~34 Project Components "
    "listed in §2.3.2, which of the 23 Data Elements are:", size=10.5,
)
add_bullet("(a) Mandatory to display and fill")
add_bullet("(b) Optional to display")
add_bullet("(c) Hidden (not applicable)")
add_p("Illustrative example rows:", size=10, italic=True)
tbl = DOC.add_table(rows=0, cols=6)
tbl.style = "Light Grid Accent 1"
add_table_row(tbl, ["Project Component", "§2.3.9", "§2.3.10", "§2.3.15", "§2.3.16", "§2.3.20"], header=True)
add_table_row(tbl, ["Cold Storage", "Mandatory", "Optional", "Mandatory", "Mandatory", "Mandatory"])
add_table_row(tbl, ["Custom Hiring Centre", "Optional", "Hidden", "Mandatory", "Optional", "Optional"])
add_table_row(tbl, ["Marketing / Retail Outlet", "Optional", "Hidden", "Optional", "Optional", "Optional"])
add_table_row(tbl, ["Processing Unit", "Mandatory", "Mandatory", "Mandatory", "Mandatory", "Mandatory"])
add_p("")
add_p("Without this matrix, the dynamic questionnaire engine cannot be built to specification.", size=10, italic=True)

add_heading("A.2 — Knowledge Base Scope for AI Content Generation", level=3)
add_p("Reference: Spec §5.7 (AI Knowledge Utilisation)", size=10, italic=True)
add_p(
    "The specification requires the AI system to generate content using \"domain knowledge sourced from "
    "KAU-approved materials, agricultural handbooks, government scheme documentation, and best-practice references.\"",
    size=10.5,
)
add_p("Request:", size=10.5, bold=True)
add_p("Please indicate whether KAU can provide (or approve) an authoritative Kerala-agriculture knowledge base to be used as reference material for the AI system, covering:", size=10.5)
add_bullet("Commodity-specific practices — Kerala commodities with production/processing/storage/marketing practices per commodity")
add_bullet("Government schemes applicable to FPOs — MIDH-SHM, PMFME, NABARD FPO Fund, AIF, KIIFB, LEAF, PMKSY and state-specific schemes with subsidy caps and eligibility criteria")
add_bullet("Standard Operating Procedures (SOPs) per enterprise type")
add_bullet("Kerala-specific statutory requirements — Panchayat Raj Act clauses, KSPCB norms, Kerala Water Authority norms, KSEB norms")
add_bullet("Kerala market characteristics — commodity price trends, seasonal patterns, dominant buyer profiles")
add_p("If a curated knowledge base is not available, please indicate whether Kefi Tech may compile a preliminary version for KAU review.", size=10.5)
add_p("Without this reference material, AI-generated content risks being geographically generic or factually inaccurate.", size=10, italic=True)

add_heading("A.3 — Balance Sheet Assumptions for Greenfield Projects", level=3)
add_p("Reference: Spec §4.5 (Projected Balance Sheet)", size=10, italic=True)
add_p(
    "For greenfield (new) enterprises, the opening balance sheet contains no historical values. "
    "Please confirm the convention for constructing the opening balance sheet:",
    size=10.5,
)
add_bullet("(a) Assume all-zero opening balances with Day-1 receipts from initial capital")
add_bullet("(b) Use the promoter contribution and initial loan disbursement as Day-1 equity and liability")
add_bullet("(c) Some other convention")
add_p("Also please confirm the number of projection years required (5, 7, or 10) for the Balance Sheet.", size=10.5)

DOC.add_page_break()

# ── SECTION B ────────────────────────────────────────────────────────────
add_heading("Section B — Important Clarifications", level=1)
add_p("These items must be resolved before the relevant Data Elements are implemented.", size=10.5, italic=True)
add_p("")

SECTION_B = [
    ("B.1 — Nature of Business: Single or Multi-select?",
     "Spec §2.3.3 (page 11)",
     "The specification lists Nature of Business as a \"Multiple Choice Checkbox\" and states that \"Multiple selections shall be permitted.\" A project may commonly have a single dominant nature.",
     "Please confirm that Nature of Business is intended as multi-select and provide the authoritative list of options. The specification lists 14 options; please confirm whether all 14 are required or if any should be removed for the Kerala-FPO context."),
    ("B.2 — Project Components Master List",
     "Spec §2.3.2 (pages 9–11)",
     "The specification provides a seed list of Project Components across 6 groups (Primary Production, Processing & Value Addition, Storage & Post-Harvest, Marketing & Business Development, Service-Based Enterprises, Supporting Infrastructure).",
     "Please confirm the authoritative list. In particular: (a) Are there Kerala-specific commodities that require dedicated components (e.g. cardamom-specific, coir-specific, rubber-specific, pepper-specific)? (b) Are there any components in the seed list that KAU wishes to exclude?"),
    ("B.3 — Land Area Unit Conventions",
     "Spec §2.3.13 (pages 55–56)",
     "The specification allows land area to be entered \"in the unit of measurement selected by the user.\"",
     "Please confirm the list of supported units for land area. Common Kerala units include: acres, cents, ares, hectares, square metres. Should the system: (a) Store the value in the entered unit and display in the same unit? (b) Auto-convert to a standard unit (e.g. acres) for calculations while displaying in the entered unit?"),
    ("B.4 — Total Project Cost Variance Tolerance",
     "Spec §2.3.4 (page 12) — Validation bullet 3",
     "The specification requires the system to \"compare the entered estimate with the automatically computed project cost and flag significant variations.\" The threshold for a \"significant variation\" is not defined.",
     "Please indicate the acceptable variance threshold: (a) A fixed monetary amount (e.g. ₹1,00,000)? (b) A percentage (e.g. 5%, 10%, 15%)? (c) Scheme-dependent (e.g. 5% for MIDH-SHM, 10% for SHM)?"),
    ("B.5 — AI Content Regeneration When User Edits Exist",
     "Spec §5.10 (User Review and Editing)",
     "The specification allows FPO users to edit AI-generated content and to regenerate individual chapters.",
     "When a user has edited an AI-generated chapter and subsequently requests regeneration, please confirm the intended behaviour: (a) Overwrite — regenerated content replaces the user's edits without confirmation; (b) Preserve edits — regeneration only fills empty chapters and skips edited ones; (c) Diff and confirm — show a side-by-side comparison and allow the user to choose paragraphs to keep."),
    ("B.6 — Ownership of Financial Assumption Defaults",
     "Spec §2.3.18 Category I (Financial Assumptions)",
     "The specification classifies financial assumptions (inflation, escalation rates, depreciation, tax rate, discount rate) as admin-configurable.",
     "Please indicate the responsible authority for setting and updating these defaults: (a) KAU Central Administrator (single set of values used by all FPOs across Kerala); (b) Kefi Tech Administrator (updated per operational requirements); (c) Per-FPO (each FPO may set its own values with system defaults as fallback). Also please confirm whether FPOs should be permitted to override system defaults, and if so, whether an audit trail should be maintained."),
    ("B.7 — Bilingual PDF Output",
     "Spec (language not explicitly stated in any chapter)",
     "The specification does not indicate the language(s) required for the generated DPR PDF.",
     "Please confirm: (a) English only; (b) English and Malayalam (bilingual output); (c) User-selectable at generation time. If bilingual, please indicate whether AI-generated narrative content is required in both languages or only in English with FPO-supplied fields in the user's preferred language."),
    ("B.8 — Multi-Parcel Land Support",
     "Spec §2.3.13 (Kefi Tech remarks, page 62)",
     "The specification's remarks state: \"The software shall support multiple land parcels under a single project and enable mapping of individual project components to specific land parcels.\"",
     "Please confirm this is a Priority 1 requirement for the initial release, or whether single-parcel support is acceptable for Phase 1 with multi-parcel deferred to a future release."),
    ("B.9 — Overall Project Risk Rating Algorithm",
     "Spec §2.3.22 (Kefi Tech remarks, page 126)",
     "The specification states the system shall \"classify the project as Low Risk, Moderate Risk, or High Risk based on cumulative assessment across all risk categories.\"",
     "Please provide the intended weighting algorithm across the six risk categories (Production, Market, Financial, Institutional, Environmental, Regulatory), or confirm whether the rating should follow a rule-based approach (e.g. Overall = High if any single category is High) or a cumulative-score approach."),
]

for title, ref, background, request in SECTION_B:
    add_heading(title, level=3)
    add_p(f"Reference: {ref}", size=10, italic=True)
    add_p(background, size=10.5)
    add_p("Request:", size=10.5, bold=True)
    add_p(request, size=10.5)

DOC.add_page_break()

# ── SECTION C ────────────────────────────────────────────────────────────
add_heading("Section C — Deferrable Clarifications", level=1)
add_p("These items can be resolved during User Acceptance Testing without blocking the build.", size=10.5, italic=True)
add_p("")

SECTION_C = [
    ("C.1 — Justifications in Project Rationale (§2.3.7)",
     "The specification requires \"Brief Justification (max 100 words)\" for each selected Rationale reason. If an FPO selects 15 reasons, the user is required to write 1,500 words of free-text justification.",
     "Please confirm whether AI may draft the initial justifications based on the FPO's project profile, allowing the user to edit rather than write from scratch."),
    ("C.2 — Historical Turnover Auto-population (§2.3.8)",
     "The specification requires the FPO to enter its current annual turnover under Baseline Information.",
     "For FPOs registered on the KAU platform with existing tier-assessment data or connected GST records, may the system auto-populate historical turnover values rather than requiring re-entry?"),
    ("C.3 — DPR Consultant User Role",
     "The specification lists \"DPR Consultants (where engaged)\" as an intended user category (Spec §1.5).",
     "Should the system provide a separate consultant user role, and if so, should consultants be permitted to: (a) access multiple FPOs' DPRs from a single login; (b) collaborate on DPR sections concurrently with the FPO; (c) sign off DPR sections as reviewed?"),
    ("C.4 — PDF Version Retention",
     "Currently the system supports unlimited versioning of generated DPR PDFs.",
     "How many versions of a generated DPR PDF must be retained per project? Please indicate if a maximum limit is preferred (e.g. last 10 versions)."),
    ("C.5 — Sensitivity Analysis Priority",
     "Spec §4.9 marks Sensitivity Analysis as Optional.",
     "Do target lending institutions require sensitivity analysis in the DPR? If not routinely required, may this feature be deferred to a Phase 2 release? Please also indicate which banks are being targeted."),
    ("C.6 — Auto-Inferred Field Transparency",
     "Spec §1.9 distinguishes between User-entered and System-generated information.",
     "For AI-derived or admin-configured values that appear as if entered by the user (e.g. an estimated inventory holding period derived from commodity defaults), should the user interface display a visible indicator such as \"Estimated by system\" alongside each such value?"),
    ("C.7 — Auto-Inferred Field Editability",
     "Related to C.6.",
     "For fields the AI infers rather than the user entering, please confirm whether the user should be permitted to override the inferred value, and if so, whether an audit trail is required."),
    ("C.8 — Onboarding and User Training",
     "Not addressed in the specification.",
     "Will KAU coordinate FPO training for the new DPR module, or should Kefi Tech include training materials and video tutorials as part of the deliverables?"),
    ("C.9 — Approved Sample DPRs for Reference",
     "Not addressed in the specification.",
     "May KAU share 2-3 sample DPRs that have previously been approved by banks or funding institutions? These would serve as reference material for calibrating AI-generated content quality and PDF layout expectations."),
]

for title, background, request in SECTION_C:
    add_heading(title, level=3)
    add_p(background, size=10.5)
    add_p("Question:", size=10.5, bold=True)
    add_p(request, size=10.5)

DOC.add_page_break()

# ── SECTION D — new items ────────────────────────────────────────────────
add_heading("Section D — Items Discovered During Implementation", level=1)
add_p(
    "The following items were not part of the original clarification list but emerged during "
    "backend build (Phase 2) and frontend build (Phase 5). They are included for KAU's review "
    "before user-acceptance testing.",
    size=10.5, italic=True,
)
add_p("")

SECTION_D = [
    ("D.1 — Field-Name Conflict Between §2.3.5 and §2.3.11",
     "Spec §2.3.5 (Proposed Products and Services, page 13) and §2.3.11 (Market Assessment, Category A, page 36)",
     "Both sections contain a dropdown named \"Product Type\" but with different intended values. §2.3.5 lists Finished / Intermediate / By-product / Service. §2.3.11 Category A lists Primary Product / Secondary Product / By-product.",
     "Interpretation adopted: §2.3.5's \"Product Type\" was treated as-is (Finished/Intermediate/By-product/Service) and a separate \"Primary/Secondary\" field was added. §2.3.11 was left as originally specified. Please confirm the intended semantic for each section. If both fields represent the same concept, please indicate which authoritative list should apply."),
    ("D.2 — Section 2.3.9 \"Major Processing / Production Activities\"",
     "Spec §2.3.9 Category C (Production Process), page 22",
     "The specification lists this field as \"Dynamic Multi-select based on enterprise\" — implying the list of activities should change based on the selected project components.",
     "As Section A.1 (Component-to-Section Visibility Matrix) has not been received, this field is currently rendered as a free-text field in the FPO wizard. Confirmation of the component-to-activity mapping via Section A.1 will allow this field to be converted to a dynamic multi-select as originally specified."),
    ("D.3 — Auto-Save Debounce (User Experience)",
     "Not specified in the specification document; UX decision by Kefi Tech.",
     "The FPO DPR wizard implements automatic save 800 milliseconds after the last field change. No manual \"Save\" button is provided.",
     "Please confirm this behaviour is acceptable for KAU users, or indicate an alternative such as (a) longer debounce interval, (b) explicit manual save, or (c) save on section change only."),
    ("D.4 — Wizard Section Groupings and Labels",
     "Not specified in the specification document; UX decision by Kefi Tech.",
     "The FPO wizard groups the 21 built data elements into two navigational categories in the sidebar: Project Definition (§2.3.2 through §2.3.9 and §2.3.10, §2.3.11) and Project Execution (§2.3.12 through §2.3.22). Section headings display only sequential step numbers (1, 2, 3…) and the section title. Specification references (e.g. \"§2.3.10\") are hidden from the FPO-facing UI to reduce jargon; they remain visible in developer and admin surfaces.",
     "Please confirm this presentation is acceptable, or indicate the preferred alternative (e.g. show spec numbers, use different grouping labels, single flat list)."),
]

for title, ref, background, request in SECTION_D:
    add_heading(title, level=3)
    add_p(f"Reference: {ref}", size=10, italic=True)
    add_p(background, size=10.5)
    add_p("Request:", size=10.5, bold=True)
    add_p(request, size=10.5)

DOC.add_page_break()

# ── Closing ──────────────────────────────────────────────────────────────
add_heading("Response Requested", level=1)
add_p(
    "Kefi Tech Solutions Pvt. Ltd. requests written responses to the above clarifications, "
    "in particular the three Critical items (Section A), at the earliest convenience of Kerala Agricultural University.",
    size=10.5,
)
add_p("")
add_p(
    "The Critical items directly block Phase 3 (Dynamic Questionnaire Engine) and Phase 4 "
    "(AI Content Generation) of the build plan. Delayed responses will proportionally delay the corresponding deliverables.",
    size=10.5,
)
add_p("")
add_p("Kefi Tech will proceed in parallel with:", size=10.5)
add_bullet("Phase 0 — Foundation (v2 skeleton, master data seeding, admin CRUD) — completed")
add_bullet("Phase 1 — Pilot pattern (§2.3.10 Raw Material end-to-end) — completed")
add_bullet("Phase 2 — Data element expansion for elements not blocked by Section A items — completed for all 21 sections")
add_bullet("Phase 5 — Frontend wizard for all 21 sections and admin master data — completed")
add_p("")
add_p("Should Kerala Agricultural University require clarification on any of these items, or wish to schedule a review meeting, please contact:", size=10.5)
add_p("")
add_p("Athul Gopan", size=10.5, bold=True)
add_p("Kefi Tech Solutions Pvt. Ltd.", size=10.5)
add_p("Email: athul.gopan@kefitech.com", size=10.5)
add_p("")
add_p("— End of Request for Clarification Document —", size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ─────────────────────────────────────────────────────────────────
OUT = "/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/Documents/KAU_FPO_DPR_RCD_KefiTech_v1.0.docx"
DOC.save(OUT)
print(f"Saved: {OUT}")
