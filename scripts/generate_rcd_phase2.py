"""
RCD Phase 2 Document Generator — KAU-FPO Platform
Matches Phase 1 RCD house style exactly (KefiTech).
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paths ────────────────────────────────────────────────────────────────────
KAU_LOGO   = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/KAU Emblem 0.5x0.75.jpg'
LOGO_SMALL = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/Documents/kefitech_logo_small.png'
OUTPUT     = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/Documents/KAU_FPO_RCD_Phase2_KefiTech_v1.0.docx'

# ─── Colors ───────────────────────────────────────────────────────────────────
NAVY_HEX  = '1F3864'
DARK_NAVY = RGBColor(0x1F, 0x38, 0x64)
ORANGE    = RGBColor(0xE8, 0x6C, 0x1A)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREY      = RGBColor(0x60, 0x60, 0x60)
BLACK     = RGBColor(0x00, 0x00, 0x00)


# ─── XML Helpers ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, start=120, bottom=60, end=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'),    str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def remove_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'none')
        b.set(qn('w:sz'),    '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_bottom_border(paragraph, color=NAVY_HEX, sz='8'):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_page_number(run):
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)


def set_col_width(cell, width_inches):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def set_table_width(table, width_cm):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    twips = int(width_cm * 567)
    tblW.set(qn('w:w'), str(twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


# ─── Header / Footer builders ─────────────────────────────────────────────────

def build_header(header_obj):
    """Header for pages 2+ — 3-col table with logo | title | reference."""
    for p in header_obj.paragraphs:
        p._element.getparent().remove(p._element)

    tbl = header_obj.add_table(rows=1, cols=3, width=Inches(6.5))
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)

    # Logo cell
    c0 = tbl.cell(0, 0)
    set_col_width(c0, 1.15)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(0)
    r0 = p0.add_run()
    r0.add_picture(LOGO_SMALL, width=Cm(2.92), height=Cm(0.49))

    # Title cell
    c1 = tbl.cell(0, 1)
    set_col_width(c1, 3.82)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after  = Pt(0)
    r1 = p1.add_run('RCD — AI-Based Digital Platform for KAU-FPO Linkage — Phase 2')
    r1.font.bold      = True
    r1.font.size      = Pt(9)
    r1.font.color.rgb = DARK_NAVY

    # Reference cell
    c2 = tbl.cell(0, 2)
    set_col_width(c2, 1.30)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run('KAUCC/459/2025-C3')
    r2.font.size      = Pt(9)
    r2.font.color.rgb = DARK_NAVY

    # Divider line
    div = header_obj.add_paragraph()
    div.paragraph_format.space_before = Pt(2)
    div.paragraph_format.space_after  = Pt(0)
    add_bottom_border(div)


def build_first_page_header(header_obj):
    """First page header — blank (different_first_page_header_footer = True)."""
    for p in header_obj.paragraphs:
        p._element.getparent().remove(p._element)
    header_obj.add_paragraph()  # empty


def build_footer(footer_obj):
    for p in footer_obj.paragraphs:
        p._element.getparent().remove(p._element)

    tbl = footer_obj.add_table(rows=1, cols=2, width=Inches(6.5))
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)

    cl = tbl.cell(0, 0)
    set_col_width(cl, 4.5)
    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rl = pl.add_run('Kefi Tech Solutions Pvt Ltd  │  Athul.gopan@kefitech.com')
    rl.font.size      = Pt(8)
    rl.font.color.rgb = GREY

    cr = tbl.cell(0, 1)
    set_col_width(cr, 1.77)
    pr = cr.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = pr.add_run('Page ')
    rr.font.size      = Pt(8)
    rr.font.color.rgb = GREY
    rn = pr.add_run()
    rn.font.size      = Pt(8)
    rn.font.color.rgb = GREY
    add_page_number(rn)


# ─── Content Helpers ──────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold      = True
    r.font.size      = Pt(16)
    r.font.name      = 'Calibri'
    r.font.color.rgb = DARK_NAVY
    add_bottom_border(p)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold      = True
    r.font.size      = Pt(12)
    r.font.name      = 'Calibri'
    r.font.color.rgb = ORANGE
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold      = True
    r.font.size      = Pt(11)
    r.font.name      = 'Calibri'
    r.font.color.rgb = BLACK
    return p


def body(doc, text, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Cambria'
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Cambria'
    return p


def numbered_item(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    r1 = p.add_run(f'{number}.  ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Cambria'
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.name = 'Cambria'
    return p


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    return p


def make_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl    = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    tbl.style = 'Table Grid'

    # Header row
    for j, h_text in enumerate(headers):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, NAVY_HEX)
        set_cell_margins(c)
        if col_widths:
            set_col_width(c, col_widths[j])
        p = c.paragraphs[0]
        r = p.add_run(h_text)
        r.font.bold      = True
        r.font.color.rgb = WHITE
        r.font.size      = Pt(10)
        r.font.name      = 'Calibri'

    # Data rows — alternate white
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            c = tbl.rows[i + 1].cells[j]
            set_cell_margins(c)
            if col_widths:
                set_col_width(c, col_widths[j])
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = 'Cambria'

    spacer(doc, 8)
    return tbl


def action_box(doc, items):
    """
    Renders an 'Action Required' block.
    items: list of strings (one per numbered point) OR a single string.
    """
    spacer(doc, 4)
    # Rule line
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after  = Pt(4)
    add_bottom_border(rule, color='CCCCCC', sz='4')

    if isinstance(items, str):
        items = [items]

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Action Required')
    r.font.bold      = True
    r.font.size      = Pt(11)
    r.font.name      = 'Calibri'
    r.font.color.rgb = DARK_NAVY

    for idx, item in enumerate(items, start=1):
        numbered_item(doc, idx, item)

    spacer(doc, 4)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ─── Default paragraph style ──────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Cambria'
style.font.size = Pt(11)

# Page setup
sec = doc.sections[0]
sec.page_width        = Cm(21)
sec.page_height       = Cm(29.7)
sec.left_margin       = Cm(2.54)
sec.right_margin      = Cm(2.54)
sec.top_margin        = Cm(2.0)
sec.bottom_margin     = Cm(2.0)
sec.header_distance   = Cm(1.0)
sec.footer_distance   = Cm(1.0)
sec.different_first_page_header_footer = True

build_header(sec.header)
build_footer(sec.footer)
build_first_page_header(sec.first_page_header)
build_footer(sec.first_page_footer)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

# KAU Emblem — very top, centered, 9.14cm wide
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(12)
p.add_run().add_picture(KAU_LOGO, width=Cm(9.14))

# ── Title Banner (single-cell table, navy fill) ───────────────────────────────
banner_tbl = doc.add_table(rows=1, cols=1)
banner_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(banner_tbl, 16.51)
bc = banner_tbl.cell(0, 0)
set_cell_bg(bc, NAVY_HEX)
set_cell_margins(bc, top=160, start=200, bottom=160, end=200)
bp = bc.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
br = bp.add_run('REQUIREMENTS CLARIFICATION DOCUMENT — PHASE 2')
br.font.bold      = True
br.font.size      = Pt(20)
br.font.name      = 'Calibri'
br.font.color.rgb = WHITE

spacer(doc, 14)

# KAU name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('KERALA AGRICULTURAL UNIVERSITY')
r.font.bold      = True
r.font.size      = Pt(16)
r.font.name      = 'Calibri'
r.font.color.rgb = DARK_NAVY

# Address
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run('Communication Centre, Mannuthy, Thrissur — 680651')
r.font.size      = Pt(11)
r.font.name      = 'Cambria'
r.font.color.rgb = GREY

# Project name
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('AI-Based Digital Platform for KAU-FPO Linkage Programme')
r.font.bold      = True
r.font.size      = Pt(13)
r.font.name      = 'Calibri'
r.font.color.rgb = ORANGE

# Doc sub-title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run('Requirements Clarification Document — Phase 2')
r.font.bold      = True
r.font.size      = Pt(13)
r.font.name      = 'Calibri'
r.font.color.rgb = ORANGE

# Version / Prepared by
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Version 1.0  |  Prepared by Kefi Tech Solutions Pvt. Ltd.')
r.font.size      = Pt(11)
r.font.name      = 'Cambria'
r.font.color.rgb = GREY

# Confidential notice
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
r = p.add_run('Confidential — Prepared in response to Tender Notice KAUCC/459/2025-C3')
r.font.italic     = True
r.font.size       = Pt(9)
r.font.name       = 'Cambria'
r.font.color.rgb  = GREY

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '1. Introduction')

h2(doc, '1.1 Purpose')
body(doc,
    'This Requirements Clarification Document (RCD) is prepared by Kefi Tech Solutions Pvt. Ltd. '
    'for Phase 2 of the AI-Based Digital Platform for the KAU-FPO Linkage Programme.')
spacer(doc, 4)
body(doc,
    'Phase 1 of the platform has been delivered, tested, and deployed. This document covers the '
    'decisions, data requirements, and third-party credentials needed before Phase 2 development begins.')
spacer(doc, 4)
body(doc, 'This document serves two purposes:')
numbered_item(doc, 1,
    'To inform KAU of design decisions already made for Phase 2 based on the SRS and the DPR Data '
    'Collection Module specification (July 2026), so that KAU may confirm or flag corrections before '
    'development begins.')
numbered_item(doc, 2,
    'To formally request from KAU the data, content, training datasets, and third-party credentials '
    'that Kefi Tech Solutions cannot generate internally and that are required to complete Phase 2.')

h2(doc, '1.2 Scope')
body(doc, 'This document covers all 14 Phase 2 modules as specified in the SRS (§3.2):')
make_table(doc,
    ['Module ID', 'Module Name'],
    [
        ('P2-01', 'Row-Level Security'),
        ('P2-02', 'Government Portal'),
        ('P2-03', 'CBBO / NGO Portal'),
        ('P2-04', 'Auto-Translate (AI-Assisted)'),
        ('P2-05', 'GIS Integration'),
        ('P2-06', 'AI Crop Recommendations'),
        ('P2-07', 'AI-Assisted DPR Generation'),
        ('P2-08', 'Advanced Expert Booking'),
        ('P2-09', 'Analytics Dashboards'),
        ('P2-10', 'AI Chatbot'),
        ('P2-11', 'ONDC Marketplace + Farmer Connect'),
        ('P2-12', 'Public Market Linkage Hub'),
        ('P2-13', 'WhatsApp Extended Templates'),
        ('P2-14', 'AI Product Marketing Strategies'),
    ],
    col_widths=[1.3, 5.0],
)

h2(doc, '1.3 How to Respond')
body(doc,
    'Please review each section and respond within 7 working days from the date of this document. '
    'For each item:')
bullet(doc, 'If a decision is correct — confirm it.')
bullet(doc, 'If a decision is incorrect — provide the correct information.')
bullet(doc, 'If data is requested — provide it in the format described.')
bullet(doc, 'If a question is asked — provide a clear answer.')
spacer(doc, 4)
body(doc, 'For a walkthrough call, contact: Athul Gopan — Athul.gopan@kefitech.com')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — DECISIONS ALREADY MADE
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '2. Decisions Already Made — Please Confirm or Correct')
body(doc,
    'The following decisions have been made based on the SRS and the DPR Data Collection Module '
    'specification received from KAU in July 2026. Development will proceed on these assumptions. '
    'If any of the below are incorrect, please inform us immediately.')

# ── 2.1 DPR Module ────────────────────────────────────────────────────────────
h2(doc, '2.1 DPR Module — How the System Works (P2-07)')
body(doc,
    'We have read the DPR Data Collection Module specification (Version 1.0, July 2026) prepared by '
    'Meer Ahamed Ibrahim (Technical Consultant) and approved by Dr. Binoo P Bonny. Below is our '
    'complete understanding of how the DPR module will be built. Please confirm each section is correct.')

h3(doc, '2.1.1 Overall Flow')
body(doc,
    'The FPO does not prepare a DPR document manually. The platform guides the FPO through a '
    'structured data entry wizard and the system generates the complete bank-ready DPR PDF '
    'automatically. The FPO\'s role is to answer questions — the system does all calculations '
    'and writes all narrative content.')
spacer(doc, 4)
body(doc, 'The overall process:')
steps = [
    'FPO selects project components (e.g. processing unit + storage + transport)',
    'Platform shows only the sections relevant to those components (23 sections max)',
    'FPO fills in each section (data entry, not document writing)',
    'Platform calculates all financial projections automatically',
    'AI generates narrative sections based on the data entered',
    'FPO reviews and optionally edits the AI-written narratives',
    'DPR Readiness Score is shown (0–100)',
    'FPO generates the final PDF — stored in AWS S3 and downloadable',
]
for i, s in enumerate(steps, 1):
    numbered_item(doc, f'Step {i}', s)

h3(doc, '2.1.2 The 23-Section Wizard')
body(doc,
    'The wizard has 23 sections. Sections are shown or hidden based on the components the FPO '
    'selects in Section 1.')
make_table(doc,
    ['Section #', 'Section Name', 'Type'],
    [
        ('1',  'Project Overview & Component Selection',  'Data Entry'),
        ('2',  'Promoter / FPO Details',                  'Data Entry'),
        ('3',  'Market Assessment',                       'Data Entry'),
        ('4',  'Technical Details — Civil Works',         'Data Entry'),
        ('5',  'Technical Details — Plant & Machinery',   'Data Entry'),
        ('6',  'Technical Details — Utilities & Infrastructure', 'Data Entry'),
        ('7',  'Raw Material Availability',               'Data Entry'),
        ('8',  'Production Capacity & Schedule',          'Data Entry'),
        ('9',  'Manpower Requirements',                   'Data Entry'),
        ('10', 'Project Cost Summary',                    'Auto-calculated'),
        ('11', 'Means of Finance',                        'Data Entry'),
        ('12', 'Revenue Projections',                     'Data Entry + Auto-calculated'),
        ('13', 'Profit & Loss Statement',                 'Auto-calculated'),
        ('14', 'Cash Flow Statement',                     'Auto-calculated'),
        ('15', 'Balance Sheet',                           'Auto-calculated'),
        ('16', 'Financial Ratios (IRR, NPV, BCR, DSCR, Payback)', 'Auto-calculated'),
        ('17', 'Loan Repayment Schedule',                 'Auto-calculated'),
        ('18', 'Depreciation Schedule',                   'Auto-calculated'),
        ('19', 'Break-Even Analysis',                     'Auto-calculated'),
        ('20', 'Risk Assessment',                         'AI-generated + editable'),
        ('21', 'SWOT Analysis',                           'AI-generated + editable'),
        ('22', 'Environmental & Social Assessment',       'AI-generated + editable'),
        ('23', 'Implementation Timeline',                 'Data Entry + AI-generated narrative'),
    ],
    col_widths=[0.8, 4.2, 1.77],
)

h3(doc, '2.1.3 Automatic Financial Calculations')
body(doc,
    'The FPO enters basic inputs (cost of civil works, machinery cost, production capacity, selling '
    'price, raw material cost, manpower cost, loan amount). The platform calculates everything else '
    'automatically.')
make_table(doc,
    ['Calculation', 'Basis'],
    [
        ('Total project cost',        'Civil + machinery + equipment + pre-operative + contingency'),
        ('Depreciation schedule',     'Method and rates from DPRMasterConfig (admin-configurable)'),
        ('Loan repayment schedule',   'EMI / bullet / balloon — based on loan terms entered by FPO'),
        ('Profit & Loss (5 years)',   'Revenue − operating cost − depreciation − interest'),
        ('Cash Flow Statement (5 years)', 'Net profit + depreciation − loan repayment − capex'),
        ('Balance Sheet (5 years)',   'Assets and liabilities reconciled year-by-year'),
        ('IRR',                       'Internal rate of return on project cash flows'),
        ('NPV',                       'Net present value at discount rate from DPRMasterConfig'),
        ('BCR',                       'Benefit-cost ratio'),
        ('DSCR',                      'Debt service coverage ratio (year-by-year)'),
        ('Payback period',            'Simple payback from cash flows'),
        ('Break-even point',          'In units and in revenue'),
    ],
    col_widths=[2.5, 4.0],
)
body(doc,
    'All default rates (interest rates, depreciation rates, inflation, projection period) are stored '
    'in the DPRMasterConfig table. The KAU Super Admin can update these values from the admin panel '
    'at any time — no code change needed.')

h3(doc, '2.1.4 AI Narrative Generation (Claude API)')
body(doc,
    'Narrative sections of the DPR are written by the Claude AI. Sections generated by AI: '
    'Executive Summary, Market Analysis & Demand Assessment, Technical Feasibility Assessment, '
    'Risk Assessment (Section 20), SWOT Analysis (Section 21), Environmental & Social Assessment '
    '(Section 22), Implementation Narrative.')
spacer(doc, 4)
body(doc,
    'After AI generation, the FPO can review each narrative and make edits before final PDF '
    'generation. Note: AI narrative generation is dependent on KAU\'s budget approval for the '
    'Claude API (see Section 4.4). Until budget is approved, these sections will show a placeholder '
    'text — the rest of the DPR (data entry + calculations) will work fully.')

h3(doc, '2.1.5 DPR Readiness Score')
body(doc,
    'Before the FPO can generate the PDF, the platform calculates a DPR Readiness Score (0–100).')
make_table(doc,
    ['Validation Type', 'What it checks'],
    [
        ('Mandatory',   'All required fields are filled'),
        ('Logical',     'Values are logically consistent (e.g. revenue > raw material cost)'),
        ('Technical',   'Technical inputs are within feasible ranges for the selected commodity'),
        ('Financial',   'Financial ratios meet bank-acceptability thresholds (e.g. DSCR ≥ 1.25)'),
        ('Cross-element', 'Data in different sections is consistent with each other'),
    ],
    col_widths=[1.8, 4.5],
)
body(doc,
    'Validation errors block PDF generation. Validation warnings are shown to the FPO but do '
    'not block PDF generation.')

h3(doc, '2.1.6 PDF Output')
body(doc,
    'The final output is a professionally formatted PDF. It includes: Cover page (FPO name, '
    'project name, date, KAU branding), all 23 sections in standard DPR format, all financial '
    'tables (P&L, cash flow, balance sheet, loan schedule, ratios), AI-generated narrative sections '
    '(edited by FPO if chosen), and annexures as applicable. The PDF is stored in AWS S3 and '
    'previous versions are retained if the FPO regenerates the DPR.')

action_box(doc, [
    'Are the 23 sections listed in 2.1.2 correct and complete?',
    'Are the auto-calculated financial statements in 2.1.3 the complete set required?',
    'Confirm the AI narrative sections listed in 2.1.4 are correct.',
    'Confirm the DPR Readiness Score approach (5 validation types) in 2.1.5 is acceptable.',
    'If any section of the July 2026 specification has been updated, please share the revised document.',
])

# ── 2.2 AI Chatbot ────────────────────────────────────────────────────────────
h2(doc, '2.2 AI Chatbot — How It Works (P2-10)')
body(doc, 'SRS Reference: §3.2.9')
spacer(doc, 4)
body(doc,
    'Chatbot is available to approved FPO users only — Primary and Secondary roles. Government '
    'officials, CBBO officers, experts, and unauthenticated visitors have no chatbot access.')
spacer(doc, 4)
body(doc,
    'How it works: FPO user sends a message; the platform forwards it to the Claude API along with '
    'context about the FPO (district, commodities, tier, schemes enrolled) — no personally '
    'identifiable information is sent. The AI responds in the user\'s preferred language (English or '
    'Malayalam). Full conversation history is stored and can be recalled. Each message and model '
    'version used is logged for audit purposes. The Claude API key is stored encrypted in the '
    'platform\'s admin panel — KAU Admin can update the key without a deployment. If the Claude API '
    'is unavailable, the system returns a graceful error message.')

action_box(doc, [
    'Confirm chatbot access is for FPO users only (Primary and Secondary) — not public/anonymous users.',
    'Confirm the chatbot should respond in the user\'s preferred language (EN or ML).',
    'Confirm that FPO profile context (district, commodities, tier) is used to improve relevance — and confirm no other data (financials, documents) should be sent to the AI.',
])

# ── 2.3 Row-Level Security ────────────────────────────────────────────────────
h2(doc, '2.3 Row-Level Security — Sub-Admin and CBBO Assignment (P2-01)')
body(doc,
    'Phase 1 already stores assigned_to (sub-admin) on each FPO. Phase 2 adds assigned_cbbo '
    '(CBBO officer). Both fields enforce data isolation.')
make_table(doc,
    ['Role', 'Sees'],
    [
        ('Super Admin',    'All FPOs on the platform'),
        ('Sub-Admin',      'Only FPOs where FPO.assigned_to = this sub-admin'),
        ('CBBO Officer',   'Only FPOs where FPO.assigned_cbbo = this CBBO officer'),
    ],
    col_widths=[2.0, 4.3],
)
body(doc,
    'Assignment is done by the Super Admin from the admin panel. All assignments are written to the '
    'audit log. One sub-admin per FPO and one CBBO per FPO — one-to-one per role.')

action_box(doc, [
    'Confirm the one-to-one rule: one FPO can have only one assigned sub-admin and one assigned CBBO at a time.',
    'Confirm the Super Admin is the only role that can assign and reassign FPOs.',
])

# ── 2.4 Government Portal ─────────────────────────────────────────────────────
h2(doc, '2.4 Government Portal — How It Works (P2-02)')
body(doc,
    'Government officials are created by the KAU Super Admin. Each official is given one of two '
    'jurisdiction levels.')
make_table(doc,
    ['Jurisdiction', 'What they see'],
    [
        ('District-level',
         'All FPOs in their assigned district only — HTTP 403 if they try to access another district'),
        ('State-level',
         'All FPOs across all districts in Kerala'),
    ],
    col_widths=[1.8, 4.5],
)
body(doc,
    'The government portal is strictly read-only — zero write operations. Officials cannot approve, '
    'reject, comment on, or modify any FPO data.')
spacer(doc, 4)
body(doc,
    'What they can access: FPO dashboard (counts, tier distribution, scheme utilisation), FPO listing '
    'filtered by jurisdiction, scheme utilisation reports, district-level reports downloadable as '
    'PDF or Excel.')

action_box(doc, [
    'Confirm the two jurisdiction levels (district and state) are the only levels needed. If taluk-level jurisdiction is required, inform us.',
    'Confirm zero write access — government officials cannot comment, flag, or request actions on any FPO.',
    'Confirm which government departments/designations should have access — see Section 3.2.',
])

# ── 2.5 CBBO Portal ──────────────────────────────────────────────────────────
h2(doc, '2.5 CBBO Portal — How It Works (P2-03)')
body(doc,
    'Each CBBO officer sees only the FPOs assigned to their organisation. Complete data isolation '
    'is enforced.')
make_table(doc,
    ['Feature', 'Detail'],
    [
        ('View assigned FPOs',      'Full FPO profile for each assigned FPO'),
        ('Capacity Building Reports', 'Create, edit (draft), and submit reports per FPO'),
        ('Training Sessions',       'Create sessions with date, venue, topic, trainer'),
        ('Attendance Tracking',     'Record which FPO members attended each session — tracked by member name'),
        ('Dashboard',               'Overview of assigned FPOs, report completion status, training count'),
    ],
    col_widths=[2.2, 4.1],
)
body(doc,
    'Report lifecycle: draft → submitted → locked. Once submitted, the report cannot be edited. '
    'KAU Super Admin can view all submitted reports across all CBBOs.')

action_box(doc, [
    'Confirm the report lifecycle — draft → submitted → locked. If KAU admin needs to unlock a submitted report for correction, inform us.',
    'Please provide the required fields for a Capacity Building Report — see Section 3.10.',
    'Confirm CBBOs cannot see FPOs assigned to other CBBO organisations.',
])

# ── 2.6 Auto-Translate ────────────────────────────────────────────────────────
h2(doc, '2.6 Auto-Translate — Admin Tool (P2-04)')
body(doc,
    'This feature is not specified in the SRS. It is a Kefitech-proposed enhancement.')
spacer(doc, 4)
body(doc,
    'When a new language is added to the platform, the KAU Super Admin currently has to manually '
    'translate ~1,100 UI strings via Excel export/import. The Auto-Translate tool replaces this '
    'with a single API call: Super Admin triggers auto-translate with a language code; the system '
    'fetches all English strings that have no translation for the target language; it sends them to '
    'Claude API in batches of 50 keys with a system prompt specifying agricultural domain, Kerala '
    'government context, and formal register; translations are saved with is_verified = False for '
    'admin review before publishing.')
spacer(doc, 4)
body(doc,
    'Only the Super Admin can trigger this. No end-user sees any change until translations are '
    'marked verified and published.')

action_box(doc, [
    'Confirm whether KAU approves including this feature in Phase 2. It depends on the Claude API budget being approved (Section 4.4).',
])

# ── 2.7 AI Crop Recommendations ───────────────────────────────────────────────
h2(doc, '2.7 AI Crop Recommendations — How It Works (P2-06)')
body(doc, 'SRS Reference: §3.2.1')
spacer(doc, 4)
body(doc, 'How it works:')
steps = [
    "FPO's GPS location → GIS module detects agro-climatic zone automatically",
    'FPO selects season (Kharif / Rabi / Summer) on the recommendations screen',
    'Platform sends: district, agro-zone, soil type, season, existing commodities, FPO tier to the ML model (FastAPI service on port 8001, internal only)',
    'ML model returns ranked crop recommendations',
    'Each recommendation includes: crop name, confidence score, plain-language reasoning, estimated yield, and business guidance',
    'FPO rates each recommendation (1–5 stars) — feedback is stored and used to improve the model in future retraining cycles',
]
for i, s in enumerate(steps, 1):
    numbered_item(doc, f'Step {i}', s)
spacer(doc, 4)
body(doc,
    'Key rules: Recommendations are cached per FPO per financial year. One model version is active '
    'at a time — KAU Admin can activate/deactivate model versions from the admin panel without code '
    'change. If the ML service is unavailable, the platform returns the last cached recommendation '
    'with a note. The trained model file is the property of KAU as per the IPR clause of the tender.')
spacer(doc, 4)
body(doc,
    'Until the training dataset is provided, a mock response is used during development and testing.')

action_box(doc, [
    'See Section 3.1 for the training data request.',
])

# ── 2.8 Expert Booking ────────────────────────────────────────────────────────
h2(doc, '2.8 Expert Booking — How It Works (P2-08)')
body(doc,
    'Phase 1 built the expert directory. Phase 2 adds a calendar-based booking system on top of it.')
spacer(doc, 4)
body(doc, 'How it works:')
steps = [
    'Expert sets their available time slots (date + time, can set in bulk)',
    'FPO browses available slots for an expert and submits a booking request',
    'Expert confirms or rejects the booking request',
    'If confirmed: slot is locked, notifications sent to both parties',
    '24 hours before the booking: automated reminder sent (email + in-app)',
    'After the appointment: status marked as completed',
]
for i, s in enumerate(steps, 1):
    numbered_item(doc, f'Step {i}', s)
spacer(doc, 4)
body(doc,
    'Booking status flow: pending → confirmed → completed  |  '
    'pending → rejected  |  confirmed → cancelled (by FPO or expert)')
spacer(doc, 4)
body(doc,
    'Rules: Only approved FPOs can book expert slots. A confirmed slot is locked — it cannot be '
    'double-booked. Notifications triggered at: new request, confirmation, rejection, cancellation, '
    'and 24-hour reminder.')

action_box(doc, [
    'Confirm the booking flow above is correct.',
    'Should FPOs be able to book experts from any district, or only experts in their own district?',
    'Should experts be able to set recurring availability (e.g. every Monday 10am–12pm), or only individual slots?',
])

# ── 2.9 Analytics Dashboards ─────────────────────────────────────────────────
h2(doc, '2.9 Analytics Dashboards — How It Works (P2-09)')
body(doc,
    'Analytics data is pre-calculated every night at 2:00 AM IST by a Celery Beat scheduled task. '
    'Data is at most 24 hours stale — this is displayed to the user via a last_updated timestamp.')
make_table(doc,
    ['Role', 'Analytics scope'],
    [
        ('Super Admin',                  'All Kerala — all districts'),
        ('Sub-Admin',                    'Only their assigned FPOs'),
        ('Government Official (district)', 'Their assigned district only'),
        ('Government Official (state)',   'All Kerala'),
    ],
    col_widths=[2.5, 3.8],
)
body(doc,
    'Metrics included in the snapshot: Total FPOs registered, FPOs by tier (A/B/C/D), FPOs by '
    'district, FPOs by commodity, scheme utilisation, member statistics (total members, women '
    'members), monthly registration trend.')
spacer(doc, 4)
body(doc,
    'Export formats: PDF (WeasyPrint) or Excel (openpyxl). Filters: date range and commodity.')

action_box(doc, [
    'Confirm the metrics listed above are the complete set KAU requires. If additional metrics are needed (e.g. DPR completion rate, expert booking count, marketplace activity), list them.',
    'Confirm the nightly 2:00 AM refresh schedule is acceptable.',
])

# ── 2.10 ONDC Marketplace ────────────────────────────────────────────────────
h2(doc, '2.10 ONDC Marketplace — How It Works (P2-11)')
body(doc,
    'Approved FPOs can list their agricultural products for sale. Products are visible to verified '
    'buyers in the buyer directory and — once ONDC integration is active — on the ONDC network.')
spacer(doc, 4)
body(doc,
    'Product lifecycle: draft → active → sold / expired. FPO creates a product listing; product '
    'becomes active once published; products past their availability date are automatically expired '
    'by a nightly Celery task.')
spacer(doc, 4)
body(doc,
    'Buyer-seller matching: The platform maintains a Buyer Directory managed by KAU Admin. When a '
    'new product is listed, the matching engine runs daily and identifies buyers whose requirements '
    'match the product.')
spacer(doc, 4)
body(doc,
    'ONDC integration: Each product has an is_ondc_listed flag. When ONDC credentials are active, '
    'listing a product also registers it on the ONDC network. ONDC requires quality certification '
    'for products.')
spacer(doc, 4)
body(doc,
    'Commodity prices: Daily commodity prices (from AGMARKNET via data.gov.in) are displayed '
    'alongside product listings. Prices are refreshed daily.')

action_box(doc, [
    'Confirm only approved FPOs can list products.',
    'Confirm product listings are visible to registered buyers in the buyer directory. Are they also visible to the general public, or registered users only?',
    'Confirm the buyer directory is managed by KAU Admin — not self-registration by buyers.',
    'For ONDC: confirm KAU will register as the ONDC entity — see Section 4.1.',
])

# ── 2.11 DPR Financial Calculation Defaults ───────────────────────────────────
h2(doc, '2.11 DPR Financial Calculation Defaults')
body(doc,
    'The DPR calculation engine uses configurable default assumptions stored in the database. '
    'The KAU Super Admin can update these values from the admin panel at any time.')
make_table(doc,
    ['Parameter', 'Default Value'],
    [
        ('Interest rate (term loan)',            '12% per annum'),
        ('Interest rate (working capital)',      '14% per annum'),
        ('Depreciation — Buildings',             '5% per annum (straight-line)'),
        ('Depreciation — Plant & Machinery',     '15% per annum'),
        ('Depreciation — Equipment',             '20% per annum'),
        ('Depreciation — Vehicles',              '20% per annum'),
        ('Depreciation — Furniture & Fixtures',  '10% per annum'),
        ('Contingency',                          '5% of civil + machinery cost'),
        ('Pre-operative expenses',               '2% of project cost'),
        ('Projection period',                    '5 years'),
        ('Inflation assumption',                 '6% per annum'),
    ],
    col_widths=[3.0, 3.3],
)

action_box(doc, [
    'Confirm these defaults are correct for Kerala agricultural projects. If KAU\'s agricultural economists have preferred standard values, please provide them.',
])


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — DATA REQUIRED FROM KAU
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
h1(doc, '3. Data Required from KAU')
body(doc,
    'The following data must be provided by KAU before the listed modules can be completed.')

h2(doc, '3.1 Crop Recommendation — Training Data')
body(doc, 'Needed for: P2-06 (AI Crop Recommendations)')
spacer(doc, 4)
body(doc,
    'Kefitech will train the crop recommendation model from scratch. We require a training dataset '
    'with labelled crop suitability records for Kerala. Minimum required fields: District, Soil type, '
    'Season (Kharif / Rabi / Summer), Commodity / Crop, Suitability rating or yield data.')
spacer(doc, 4)
body(doc,
    'If KAU has access to ICAR datasets, KAU research farm records, or state agriculture department '
    'datasets for Kerala, please share them. The quality and size of this dataset directly determines '
    'the accuracy of the recommendations.')

action_box(doc, [
    'Please share the training dataset or confirm the source KAU will provide. This is the longest lead-time item in the AI features — please respond on this point first.',
])

h2(doc, '3.2 Government Official Accounts')
body(doc, 'Needed for: P2-02 (Government Portal)')
spacer(doc, 4)
body(doc,
    'Please provide the initial list of Government Officials who should have access to the platform. '
    'For each official:')
make_table(doc,
    ['Field', 'Example'],
    [
        ('Full name',            'Dr. Rajan K'),
        ('Designation',          'District Agricultural Officer'),
        ('Department',           'Department of Agriculture'),
        ('Jurisdiction type',    'District / State'),
        ('Jurisdiction value',   'Thrissur (if district-level)'),
        ('Email address',        'rajan@kerala.gov.in'),
        ('Phone number',         '9XXXXXXXXX'),
    ],
    col_widths=[2.5, 3.8],
)

h2(doc, '3.3 CBBO / NGO Organisation List')
body(doc, 'Needed for: P2-03 (CBBO Portal)')
spacer(doc, 4)
body(doc,
    'Please provide a list of CBBO and NGO organisations operating under the KAU-FPO Linkage '
    'Programme. For each organisation:')
make_table(doc,
    ['Field', 'Example'],
    [
        ('Organisation name',  'Kudumbashree'),
        ('Type',               'CBBO / NGO'),
        ('Contact person',     'Meera S'),
        ('Designation',        'District Coordinator'),
        ('Email',              'meera@kudumbashree.org'),
        ('Phone',              '9XXXXXXXXX'),
        ('Districts covered',  'Thrissur, Palakkad'),
    ],
    col_widths=[2.5, 3.8],
)

h2(doc, '3.4 Agro-Climatic Zone Boundaries for Kerala')
body(doc, 'Needed for: P2-05 (GIS Integration)')
spacer(doc, 4)
body(doc,
    'The GIS module needs to display Kerala\'s agro-climatic zones on a map and assign each FPO\'s '
    'location to its zone. Please confirm: Does KAU hold GIS boundary data (Shapefile or GeoJSON) '
    'for Kerala\'s agro-climatic zones? If yes, please share the file. If not, we will source it '
    'from ISRO Bhuvan\'s WMS service.')

h2(doc, '3.5 Commodity-Specific Technical Norms (DPR Engine)')
body(doc, 'Needed for: P2-07 (DPR Generation)')
spacer(doc, 4)
body(doc,
    'The DPR calculation engine uses commodity-specific technical norms when the FPO does not '
    'provide a value. For example: Paddy processing: recovery percentage, drying loss; Banana: '
    'ripening chamber capacity norms per tonne; Coconut oil: extraction recovery per nut; Cashew: '
    'shelling recovery percentage.')
spacer(doc, 4)
body(doc,
    'Please provide KAU\'s standard technical norms for the top 20 commodities in the platform. '
    'These values go into the DPR calculation engine and affect every DPR generated.')

h2(doc, '3.6 Sub-Admin Assignment List')
body(doc, 'Needed for: P2-01 (Row-Level Security)')
spacer(doc, 4)
body(doc,
    'At Phase 2 launch, please provide: list of sub-admin email IDs currently active on the platform, '
    'and which FPOs (by Application ID) should be assigned to each sub-admin.')

h2(doc, '3.7 CBBO Officer Accounts')
body(doc, 'Needed for: P2-03 (CBBO Portal)')
spacer(doc, 4)
body(doc,
    'Once the CBBO organisations are confirmed (Section 3.3), please provide individual officer '
    'accounts to be created. For each officer: full name, designation, email address, phone number, '
    'and which CBBO organisation they belong to.')

h2(doc, '3.8 WhatsApp Message Templates')
body(doc, 'Needed for: P2-13 (WhatsApp Extended Templates)')
spacer(doc, 4)
body(doc,
    'Each WhatsApp message template must be individually approved by Meta before it can be sent. '
    'Approval typically takes 1–5 business days per template. Kefitech will draft the template '
    'content, but KAU must confirm the message text and submit it for Meta approval under the '
    'KAU WhatsApp Business Account.')
spacer(doc, 4)
body(doc, 'Templates planned for Phase 2: DPR completion reminder, tier upgrade notification, '
    'expert booking confirmation, marketplace product approved, scheme application alert, '
    'capacity building session reminder.')
spacer(doc, 4)
body(doc, 'Action required: Confirm KAU approves Kefitech drafting template content on KAU\'s behalf.')

h2(doc, '3.9 ONDC Product Categories')
body(doc, 'Needed for: P2-11 (ONDC Marketplace)')
spacer(doc, 4)
body(doc,
    'ONDC uses a specific product taxonomy. We need to map each commodity in the platform\'s '
    'commodity master list to its corresponding ONDC category code. Kefitech will prepare the '
    'mapping; KAU must confirm it before products are listed on the ONDC network.')

h2(doc, '3.10 Capacity Building Report Template')
body(doc, 'Needed for: P2-03 (CBBO Portal)')
spacer(doc, 4)
body(doc,
    'Please provide the template or format KAU uses for Capacity Building Reports submitted by '
    'CBBOs. Specifically, the list of fields required in each report — for example: date of '
    'activity, type of activity, number of participants, topics covered, outcomes, '
    'recommendations, and any supporting documents.')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — THIRD-PARTY API CREDENTIALS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
h1(doc, '4. Third-Party API Credentials Required')
body(doc,
    'None of these block development or internal testing. All of them block production go-live '
    'for the respective feature.')

h2(doc, '4.1 ONDC — Network Participant Registration')
body(doc, 'Used for: P2-11 (ONDC Marketplace)')
body(doc, 'Lead time: 2–4 months — start immediately')
spacer(doc, 4)
body(doc,
    'Kefitech recommends the DEA (Digital Enablement Agency) route: KAU registers as a DEA and '
    'partners with an existing SNP who handles the underlying ONDC network transactions. KAU '
    'remains the FPO-facing platform; the SNP handles the actual network integration.')
spacer(doc, 4)
body(doc, 'Technical findings:')
bullet(doc, 'ONDC provides a sandbox/mock environment for testing.')
bullet(doc, 'Authentication is certificate and digital-signature based.')
bullet(doc,
    'Full NP integration requires implementing Beckn protocol — no mature Python SDK exists for this.')
bullet(doc,
    'DEA route uses the SNP\'s own REST/JSON interface — significantly simpler.')
spacer(doc, 4)
body(doc, 'KAU registration requirements: Registered legal entity (KAU qualifies), GST registration, '
    'valid bank account, technical certification (Kefitech handles this).')

action_box(doc, [
    'KAU should initiate ONDC Network Participant registration under KAU\'s name.',
    'Please confirm KAU\'s GST number and legal entity name to be used for ONDC registration.',
    'Kefitech will identify a suitable SNP partner and present options to KAU before integration begins.',
])

h2(doc, '4.2 e-NAM API — Commodity Price Data')
body(doc, 'Used for: P2-11 (Marketplace), P2-12 (Market Hub)')
spacer(doc, 4)
body(doc,
    'No public developer API was found for e-NAM. e-NAM user registration exists and the platform '
    'has web dashboards with price data, but no documented public REST, SOAP, or FTP API '
    'specification was identified.')

action_box(doc, [
    'Does KAU have any existing institutional arrangement with e-NAM that provides API access? If yes, please share the API details. If not, Kefitech recommends using AGMARKNET via data.gov.in as the primary commodity price source (see Section 4.3).',
])

h2(doc, '4.3 AGMARKNET — Commodity Price Data (Recommended Primary Source)')
body(doc, 'Used for: P2-11 (Marketplace), P2-12 (Market Hub)')
body(doc, 'SRS Reference: §5.1.1')
spacer(doc, 4)
body(doc,
    'AGMARKNET data is available through the data.gov.in Data API — a confirmed public REST API. '
    'Daily commodity price data available: minimum, maximum, and modal prices by commodity, variety, '
    'market/mandi, and date. Kerala APMC data is available. Historical daily price data is available '
    'for price trend charts.')
spacer(doc, 4)
body(doc,
    'Kefitech\'s recommendation: Use AGMARKNET via data.gov.in as the primary commodity price source. '
    'e-NAM will be added only if KAU has institutional API access.')

action_box(doc, [
    'Confirm KAU approves AGMARKNET (via data.gov.in) as the primary commodity price source.',
    'Kefitech will register for a data.gov.in API key. No action needed from KAU for this step.',
    'If KAU has institutional e-NAM API access, please share those details.',
])

h2(doc, '4.4 Large Language Model (LLM) API')
body(doc, 'Used for: P2-07 (DPR Generation), P2-10 (AI Chatbot), P2-14 (Marketing Strategies), P2-04 (Auto-Translate)')
spacer(doc, 4)
body(doc,
    'SRS §3.2.9 specifies the chatbot shall be powered by a large language model API with '
    'agricultural context specific to Kerala. Kefitech has selected Anthropic Claude as the LLM '
    'provider based on its performance, API reliability, and cost structure.')
spacer(doc, 4)
body(doc,
    'Kefitech will procure the API key. This requires a budget approval from KAU for the monthly '
    'API cost.')
spacer(doc, 4)
body(doc,
    'Estimated cost: Approximately ₹8,000–₹25,000/month depending on usage volume (DPR generation '
    'volume, chatbot traffic, marketing strategy generation frequency).')

action_box(doc, [
    'Confirm KAU approves Anthropic Claude as the LLM provider and approves the estimated monthly API cost. Once approved, Kefitech will set up the account and configure credentials in the platform\'s admin panel.',
])

h2(doc, '4.5 PAN / GSTIN / CIN Verification (Live API)')
body(doc, 'Used for: Phase 1 FPO Registration (already built, currently in format-only fallback mode)')
spacer(doc, 4)
body(doc,
    'Live document verification is required during FPO registration to prevent fraudulent '
    'registrations. The platform currently falls back to format-only validation because no '
    'credentials have been provided. The Phase 1 RCD requested these credentials. They were not '
    'provided by KAU.')

action_box(doc, [
    'Please confirm which vendor KAU approves for live document verification. Common options: Signzy, IDfy, Karza, Surepass, or DigiLocker API. Without this, format-only validation remains active and fraudulent FPO registrations cannot be detected at source.',
])

h2(doc, '4.6 Bhuvan / ISRO WMS Service')
body(doc, 'Used for: P2-05 (GIS Integration)')
spacer(doc, 4)
body(doc, 'Available datasets for Kerala:')
bullet(doc, 'Administrative boundaries (district and taluk level confirmed)')
bullet(doc, 'Land Use/Land Cover (LULC)')
bullet(doc, 'Water bodies and other thematic layers')
bullet(doc, 'WMS endpoint confirmed: https://bhuvan-vec2.nrsc.gov.in/bhuvan/wms')
bullet(doc, 'CRS: EPSG:4326 — compatible with our PostGIS setup')
spacer(doc, 4)
body(doc,
    'Agro-climatic zone boundaries — not confirmed publicly available: A clearly documented '
    'downloadable Kerala agro-climatic zone Shapefile or GeoJSON was not found on Bhuvan.')

action_box(doc, [
    'Does KAU hold or have access to a Kerala agro-climatic zone boundary file (Shapefile or GeoJSON)? If yes, please share it.',
    'If KAU has an institutional data-sharing agreement with ISRO/NRSC, please share those access details.',
    'If neither is available, we will proceed with the public Bhuvan WMS for district boundaries and manually define the 5 Kerala agro-climatic zones as polygons based on KAU\'s published zone descriptions.',
])


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — OPEN QUESTIONS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
h1(doc, '5. Open Questions')
body(doc,
    'The following questions require a confirmed answer before we build the affected modules. '
    'Please respond to each item within 7 working days.')

make_table(doc,
    ['#', 'Question', 'Affects', 'Needed By'],
    [
        ('Q1',
         'Please provide the crop suitability training dataset for Kerala (see Section 3.1 for format).',
         'P2-06',
         'Before AI sprint'),
        ('Q2',
         'KAU should register as the ONDC Network Participant (DEA route recommended). Please confirm KAU\'s GST number and legal entity name for registration. Kefitech will identify SNP partner options and present to KAU.',
         'P2-11',
         'Immediately — 2–4 month lead time'),
        ('Q3',
         'Which government departments/designations should have access to the Government Portal?',
         'P2-02',
         'Before government portal sprint'),
        ('Q4',
         'Can FPO product listings on the marketplace be viewed by the general public, or only by registered buyers?',
         'P2-11',
         'Before marketplace sprint'),
        ('Q5',
         'For the Analytics Dashboard (P2-09): which metrics does KAU want to see? (e.g. FPOs by district, commodity distribution, tier distribution, scheme uptake)',
         'P2-09',
         'Before analytics sprint'),
        ('Q6',
         'Should the AI Chatbot respond in both English and Malayalam, or only the user\'s preferred language?',
         'P2-10',
         'Before chatbot sprint'),
        ('Q7',
         'For Expert Booking (P2-08): should FPOs only be able to book experts in their own district, or any expert on the platform?',
         'P2-08',
         'Before expert booking sprint'),
        ('Q8',
         'Is the Auto-Translate feature (not in SRS) approved by KAU?',
         'P2-04',
         'Before Phase 2 kickoff'),
        ('Q9',
         'Should DPR documents be downloadable by Government Officials from their portal?',
         'P2-07, P2-02',
         'Before DPR sprint'),
        ('Q10',
         'For CBBO capacity building reports — what format and fields does KAU require?',
         'P2-03',
         'Before CBBO sprint'),
    ],
    col_widths=[0.4, 3.7, 1.0, 1.2],
)


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — IMMEDIATE ACTION ITEMS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '6. Immediate Action Items for KAU')
body(doc,
    'Ordered by lead time — items with the longest lead time must be actioned first.')

make_table(doc,
    ['Priority', 'Action', 'Lead Time', 'Reason'],
    [
        ('1 — Critical',
         'Begin ONDC Network Participant registration under KAU\'s name',
         '2–4 months',
         'ONDC requires legal entity verification and technical certification — cannot be fast-tracked'),
        ('2 — Critical',
         'Provide crop suitability training dataset for Kerala',
         '4–6 weeks',
         'Kefitech trains the model from scratch; this blocks the entire AI recommendations sprint'),
        ('3 — Critical',
         'Confirm AGMARKNET (data.gov.in) as primary price source. Share e-NAM institutional API access if available.',
         '1 week',
         'No public e-NAM API found — AGMARKNET is the confirmed viable option'),
        ('4 — High',
         'Confirm PAN/GSTIN/CIN verification vendor',
         '2 weeks',
         'Required to activate fraud/spam prevention on FPO registration'),
        ('5 — High',
         'Provide commodity-specific technical norms for DPR engine',
         '2 weeks',
         'Needed before any DPR can be tested end-to-end'),
        ('6 — High',
         'Provide Government Official list for P2-02',
         '1 week',
         'Needed to create accounts and test the government portal'),
        ('7 — High',
         'Provide CBBO/NGO organisation list for P2-03',
         '1 week',
         'Needed to create CBBO accounts'),
        ('8 — Medium',
         'Approve LLM API (Anthropic Claude) budget (₹8,000–₹25,000/month estimated)',
         '1 week',
         'Kefitech will procure once approved'),
        ('9 — Medium',
         'Confirm DPR financial defaults (Section 2.11)',
         '1 week',
         'Needed before DPR calculation engine is finalised'),
        ('10 — Medium',
         'Answer all open questions in Section 5',
         '7 working days',
         'Several questions block individual module development'),
    ],
    col_widths=[1.0, 3.0, 0.9, 1.4],
)

# Closing confidential notice
spacer(doc, 20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_bottom_border(p, sz='4', color='AAAAAA')
r = p.add_run(
    'Confidential — Prepared in response to Tender Notice KAUCC/459/2025-C3  |  '
    'Kefi Tech Solutions Pvt. Ltd.  |  D1, Periyar Building, Technopark Campus, '
    'Phase I, Thiruvananthapuram, Kerala  |  Contact: Athul Gopan — Athul.gopan@kefitech.com'
)
r.font.italic     = True
r.font.size       = Pt(9)
r.font.name       = 'Cambria'
r.font.color.rgb  = GREY


# ─── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
