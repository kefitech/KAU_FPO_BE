"""
Generate User Stories DOCX for KAU-FPO Platform
=================================================
Usage:
    source venv/bin/activate && python scripts/generate_user_stories.py
Output:
    docs/KAU-FPO-User-Stories.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
KAU_GREEN  = RGBColor(0x2e, 0x7d, 0x32)
DARK_GREY  = RGBColor(0x42, 0x42, 0x42)
LIGHT_GREY = RGBColor(0x75, 0x75, 0x75)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = KAU_GREEN
    return p


def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK_GREY
    return p


def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    return p


def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.3)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def story_table(story_id, title, role, action, benefit, acceptance, priority='Medium', status='To Test'):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0].cells
    set_cell_bg(hdr[0], '2e7d32')
    set_cell_bg(hdr[1], '2e7d32')

    id_run = hdr[0].paragraphs[0].add_run(f'  {story_id}')
    id_run.bold = True
    id_run.font.color.rgb = WHITE
    id_run.font.size = Pt(10)
    hdr[0].width = Inches(1.2)

    title_run = hdr[1].paragraphs[0].add_run(f'  {title}')
    title_run.bold = True
    title_run.font.color.rgb = WHITE
    title_run.font.size = Pt(10)

    # Role / Action / Benefit
    row1 = table.add_row().cells
    set_cell_bg(row1[0], 'f1f8e9')
    row1[0].paragraphs[0].add_run('  Role').bold = True
    row1[0].paragraphs[0].runs[0].font.size = Pt(9)
    row1[1].paragraphs[0].add_run(f'  {role}').font.size = Pt(10)

    row2 = table.add_row().cells
    set_cell_bg(row2[0], 'f1f8e9')
    row2[0].paragraphs[0].add_run('  Story').bold = True
    row2[0].paragraphs[0].runs[0].font.size = Pt(9)
    story_text = f'  As a {role}, I want to {action}, so that {benefit}.'
    row2[1].paragraphs[0].add_run(story_text).font.size = Pt(10)

    # Acceptance criteria
    row3 = table.add_row().cells
    set_cell_bg(row3[0], 'f1f8e9')
    row3[0].paragraphs[0].add_run('  Acceptance\n  Criteria').bold = True
    row3[0].paragraphs[0].runs[0].font.size = Pt(9)
    criteria_text = '\n'.join([f'  ✓ {c}' for c in acceptance])
    row3[1].paragraphs[0].add_run(criteria_text).font.size = Pt(10)

    # Priority / Status
    row4 = table.add_row().cells
    set_cell_bg(row4[0], 'f1f8e9')
    row4[0].paragraphs[0].add_run('  Priority / Status').bold = True
    row4[0].paragraphs[0].runs[0].font.size = Pt(9)
    row4[1].paragraphs[0].add_run(f'  {priority}  |  {status}').font.size = Pt(10)

    doc.add_paragraph()  # spacer


# ═════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('KAU–FPO Linkage Platform')
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = KAU_GREEN

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('User Stories & Test Scenarios')
run2.bold = True
run2.font.size = Pt(16)
run2.font.color.rgb = DARK_GREY

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Kerala Agricultural University  ·  Communication Centre')
run3.font.size = Pt(11)
run3.font.color.rgb = LIGHT_GREY

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run(f'Prepared by: Kefi Tech Solutions  ·  {datetime.date.today().strftime("%B %Y")}')
run4.font.size = Pt(10)
run4.font.color.rgb = LIGHT_GREY

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# ROLES REFERENCE
# ═════════════════════════════════════════════════════════════════════════════

heading1('Platform Roles')
body('The following roles exist on the platform. Test scenarios are grouped by role.')

roles_table = doc.add_table(rows=1, cols=3)
roles_table.style = 'Table Grid'
hdr = roles_table.rows[0].cells
set_cell_bg(hdr[0], '2e7d32')
set_cell_bg(hdr[1], '2e7d32')
set_cell_bg(hdr[2], '2e7d32')
for cell, text in zip(hdr, ['  Role', '  Description', '  Access']):
    r = cell.paragraphs[0].add_run(text)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10)

for role, desc, access in [
    ('Super Admin',    'KAU platform administrator',                          'Full access to all admin functions'),
    ('Sub Admin',      'KAU staff with configurable permissions',              'Assigned permissions only'),
    ('FPO Manager',    'Primary user of a Farmer Producer Organization',       'FPO registration, dashboard, tier, team'),
    ('FPO Secondary',  'Secondary team member of an approved FPO',            'Limited FPO portal access'),
    ('Public User',    'Unregistered visitor to the landing page',             'Public pages only — no login required'),
]:
    r = roles_table.add_row().cells
    for cell, text in zip(r, [f'  {role}', f'  {desc}', f'  {access}']):
        cell.paragraphs[0].add_run(text).font.size = Pt(10)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PUBLIC LANDING PAGE
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 1 — Public Landing Page')
body('These stories apply to any visitor (no login required). Test in both English and Malayalam (X-Language: ml).')

story_table(
    'US-001', 'View Landing Page Hero',
    role='public user',
    action='view the hero section of the landing page',
    benefit='I understand what the platform is about before registering',
    acceptance=[
        'Hero headline, subheading, and description text are visible',
        'Malayalam text shown when browser/app is set to Malayalam',
        'English shown as fallback if Malayalam is unavailable',
        'Platform name and KAU branding visible',
    ],
    priority='High',
)

story_table(
    'US-002', 'View Platform Statistics Counters',
    role='public user',
    action='see live platform statistics on the landing page',
    benefit='I can see how many FPOs are using the platform',
    acceptance=[
        'Total registrations count displayed',
        'Approved FPOs count displayed',
        'Total districts covered count displayed',
        'Numbers update on page refresh',
    ],
    priority='Medium',
)

story_table(
    'US-003', 'View Announcements & News',
    role='public user',
    action='read announcements and news on the landing page',
    benefit='I stay informed about platform updates',
    acceptance=[
        'All active announcements visible on page load',
        '"All / Announcements / News" tabs filter correctly (client-side)',
        'Tamil/ML language switch updates announcement text',
        'Date shown or "Recently" if no date set',
    ],
    priority='Medium',
)

story_table(
    'US-004', 'Browse FAQs',
    role='public user',
    action='browse frequently asked questions',
    benefit='I can get answers before contacting support',
    acceptance=[
        'FAQ section has three tabs: FPO General, Schemes & Support, Platform Usage',
        'Platform Usage tab hidden if no entries exist',
        'Clicking a question expands the answer (accordion)',
        'Language switch updates FAQ text',
        '25 FPO General FAQs and 26 Schemes FAQs visible',
    ],
    priority='Medium',
)

story_table(
    'US-005', 'View How to Register Guide',
    role='public user',
    action='view the step-by-step FPO registration guide',
    benefit='I understand the legal process before starting registration',
    acceptance=[
        '"How to Register" button/link visible on landing page',
        'Clicking opens a modal/dialog with the 10-step guide',
        'Both Phase I (Legal Registration) and Phase II (Post-Incorporation) shown',
        'Modal is scrollable and closable',
    ],
    priority='Low',
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 2 — FPO REGISTRATION
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 2 — FPO User Registration & Login')

story_table(
    'US-010', 'FPO User Self-Registration',
    role='FPO primary user',
    action='create an account on the platform',
    benefit='I can start registering my FPO',
    acceptance=[
        'Registration form accepts: first name, last name, email, phone, password, confirm password',
        'Password must be min 8 chars, 1 uppercase, 1 lowercase, 1 digit, 1 special character',
        'Duplicate email shows error',
        'Successful registration sends welcome email and SMS',
        'User is redirected to login after registration',
    ],
    priority='High',
)

story_table(
    'US-011', 'FPO User Login',
    role='FPO primary user',
    action='log in to the platform',
    benefit='I can access my FPO dashboard and continue registration',
    acceptance=[
        'Login with email + password works',
        '5 consecutive wrong passwords locks account for 30 minutes',
        'Locked account shows clear error message with retry time',
        'Successful login redirects to FPO dashboard',
    ],
    priority='High',
)

story_table(
    'US-012', 'FPO Eligibility Check',
    role='FPO primary user',
    action='check if my FPO is eligible to register',
    benefit='I do not waste time filling out a form for an ineligible FPO',
    acceptance=[
        'Eligibility form asks for total members and legal structure',
        'Less than 10 members shows ineligibility message',
        'Eligible FPOs are redirected to OTP verification',
    ],
    priority='High',
)

story_table(
    'US-013', 'FPO Registration — Step 1 (Basic Info)',
    role='FPO primary user',
    action='fill in my FPO basic information',
    benefit='I can begin the registration process',
    acceptance=[
        'Form fields: FPO name, legal structure dropdown, registration number / CIN',
        'CIN required for Companies Act / Producer Companies; registration number for others',
        'CIN must be 21-character alphanumeric (validated)',
        'If PAN / CIN / email / phone already exists → duplicate detected error shown',
        'Duplicate detected → "Claim Your Business" button shown instead of error',
        'Draft FPO created on successful Step 1 submission',
    ],
    priority='High',
)

story_table(
    'US-014', 'FPO Registration — Step 2 (Contact & Location)',
    role='FPO primary user',
    action='fill in my FPO contact and location details',
    benefit='The platform can identify my FPO location',
    acceptance=[
        'Fields: district (dropdown), block (filtered by district), address, GPS coordinates',
        'Block dropdown filters based on selected district',
        'GPS latitude and longitude required',
        'Office email and phone fields present',
    ],
    priority='High',
)

story_table(
    'US-015', 'Email & Phone OTP Verification',
    role='FPO primary user',
    action='verify my FPO office email and phone number via OTP',
    benefit='The platform confirms my contact details are valid',
    acceptance=[
        'OTP sent to office email (6-digit, 10 min expiry)',
        'OTP sent to office phone via SMS (6-digit, 10 min expiry)',
        'Max 3 OTP sends per 10 minutes (rate limited — 4th attempt shows 429 error)',
        'Correct OTP marks email/phone as verified',
        'Wrong OTP shows error with remaining attempts',
        'Both email and phone must be verified before submission',
    ],
    priority='High',
)

story_table(
    'US-016', 'FPO Registration — Step 3 (Members & Directors)',
    role='FPO primary user',
    action='fill in member and director information',
    benefit='The platform has complete FPO governance information',
    acceptance=[
        'Fields: total members, women members, SC/ST members, signatory name, designation',
        'New fields: promoting agency, facilitating agency, CEO available (yes/no), accountant available (yes/no)',
        'Total directors, women directors, directors under 35 fields present',
        'Minimum 10 total members enforced',
    ],
    priority='High',
)

story_table(
    'US-017', 'FPO Registration — Step 4 (Business Details)',
    role='FPO primary user',
    action='fill in business and financial details',
    benefit='The platform has a complete picture of my FPO business',
    acceptance=[
        'Fields: commodities (multi-select), annual turnover, bank name, account number, IFSC',
        'Commodity dropdown populated from master data',
        'Turnover accepts numeric values',
    ],
    priority='High',
)

story_table(
    'US-018', 'FPO Document Upload',
    role='FPO primary user',
    action='upload required documents for my FPO',
    benefit='My application is complete and ready for submission',
    acceptance=[
        'Required documents: FPO Registration Certificate, PAN Card, Bank Details',
        'Optional documents: GST Certificate, Signatory ID, Member List, Annual Report',
        'Allowed formats: PDF, JPG, PNG (XLSX also for member list)',
        'Max file size: 5 MB standard, 10 MB for member list and annual report',
        'Document list shows: uploaded, missing required, and ready-to-submit status',
        'Uploaded documents can be deleted (only in DRAFT status)',
    ],
    priority='High',
)

story_table(
    'US-019', 'FPO Application Submission',
    role='FPO primary user',
    action='submit my completed FPO application',
    benefit='My FPO is registered and approved on the platform',
    acceptance=[
        'Submit button enabled only when all conditions met',
        'Conditions: all 4 steps complete, email verified, phone verified, 3 required docs uploaded, min 10 members, GPS coordinates present',
        'On submit: application ID generated (format: KAU-FPO-TRS-2026-0001)',
        'Status changes DRAFT → SUBMITTED → APPROVED automatically',
        'FPO user receives approval email and SMS notification',
        'Application ID visible in status page',
    ],
    priority='High',
)

story_table(
    'US-020', 'FPO Registration — Resume After Login',
    role='FPO primary user',
    action='resume my incomplete registration after logging out',
    benefit='I do not lose my progress if I close the browser',
    acceptance=[
        'GET /api/fpo/me/ returns current_step (1–4)',
        'Wizard resumes from the saved step',
        'Already-filled fields are pre-populated',
    ],
    priority='High',
)

story_table(
    'US-021', 'Ownership Claim — Duplicate FPO',
    role='FPO primary user',
    action='claim ownership of an FPO that already exists with my details',
    benefit='I can prove I am the legitimate owner and gain access',
    acceptance=[
        'Duplicate detected during Step 1 or field validation',
        'Response shows duplicate_detected: true and the duplicate field name',
        '"Claim Your Business" button shown instead of registration error',
        'Claim form shows only reason textarea (no document selector — user has no documents yet)',
        'Claim submitted successfully (201 response)',
        'Duplicate claim attempt returns 409 error',
        'Claim status page shows: pending / approved / rejected with notes',
    ],
    priority='High',
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 3 — FPO DASHBOARD & FEATURES
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 3 — FPO Dashboard & Portal Features')

story_table(
    'US-030', 'FPO Dashboard',
    role='FPO primary user',
    action='view my FPO dashboard after approval',
    benefit='I have a single overview of my FPO status and next actions',
    acceptance=[
        'FPO name, status badge, application ID visible',
        'Tier badge visible (shows "Not Assessed" if no tier yet)',
        'Quick links to documents, team, tier assessment',
        'Status timeline visible',
    ],
    priority='High',
)

story_table(
    'US-031', 'FPO Tier Assessment — Fill & Submit',
    role='FPO primary user',
    action='complete the annual tier assessment for my FPO',
    benefit='My FPO gets a tier classification (A/B/C/D) for the financial year',
    acceptance=[
        '28 questions across 6 domains presented',
        'Pre-filled questions (Q1, Q2, Q8, Q12, Q15) shown as read-only/grayed out',
        'upload_label shown on upload questions (Q6, Q17, Q27) as upload button label',
        'Conditional questions: Q19, Q20 hidden unless Q18 = Yes',
        'Answers auto-saved on each PATCH call',
        'Submit requires all mandatory questions answered',
        'On submit: tier (A/B/C/D) and total score returned',
        'Domain-wise scores shown on results screen',
    ],
    priority='High',
)

story_table(
    'US-032', 'FPO Tier Assessment — Reopen',
    role='FPO primary user',
    action='reopen a submitted tier assessment to update financial data',
    benefit='I can correct or update information for the current financial year',
    acceptance=[
        '"Edit Assessment" button visible on submitted results view',
        'Reopen returns assessment to DRAFT status',
        'Previous tier for that financial year is cleared',
        'FPO can re-submit for a new tier',
        'Reopen only allowed for current financial year assessment',
    ],
    priority='Medium',
)

story_table(
    'US-033', 'FPO Team Management',
    role='FPO primary user',
    action='invite secondary users to my approved FPO',
    benefit='My team members can access the FPO portal with appropriate permissions',
    acceptance=[
        'Invite by email — creates account with temporary password',
        'Invited user receives must_change_password prompt on first login',
        'No limit on secondary users (confirmed by KAU)',
        'Primary user can deactivate secondary users',
        'Bulk invite via file upload supported',
    ],
    priority='Medium',
)

story_table(
    'US-034', 'Browse Schemes & Subsidies',
    role='FPO primary user',
    action='browse government schemes available to my FPO',
    benefit='I can discover funding and support opportunities',
    acceptance=[
        '12 schemes visible with full details',
        'Filter by category (credit / insurance / marketing / infrastructure / capacity building)',
        'Search by keyword works',
        'Official website link opens in new tab',
        'Malayalam text shown with X-Language: ml',
    ],
    priority='Medium',
)

story_table(
    'US-035', 'Browse Expert Directory & Contact',
    role='FPO primary user',
    action='find and contact an expert for guidance',
    benefit='I can get professional support for my FPO',
    acceptance=[
        '15 KAU experts listed with designation, expertise, district',
        'Filter by category (scientist / trainer) and district',
        'Search by keyword works',
        '"Contact Expert" button visible only for approved FPOs',
        'Contact form requires min 20 character message',
        'Expert receives email notification on enquiry submission',
        'Success message shown after submission',
    ],
    priority='Medium',
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 4 — ADMIN AUTHENTICATION
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 4 — Admin Authentication')

story_table(
    'US-040', 'Super Admin Login',
    role='super admin',
    action='log in to the admin panel',
    benefit='I can manage the platform',
    acceptance=[
        'Login with email + password',
        'If 2FA is enabled → TOTP code prompt shown after password',
        '5 failed attempts → account locked for 30 min',
        'Successful login → redirected to admin dashboard',
        'HttpOnly cookie set (not visible in JS)',
    ],
    priority='High',
)

story_table(
    'US-041', 'Two-Factor Authentication Setup',
    role='super admin',
    action='set up 2FA on my admin account',
    benefit='My admin account is protected against unauthorized access',
    acceptance=[
        'QR code displayed (base64 PNG image)',
        'QR code scannable with Google Authenticator',
        'First TOTP code confirms setup',
        '8 backup codes generated and shown once',
        '2FA active on next login',
    ],
    priority='High',
)

story_table(
    'US-042', 'Forgot Password — Admin',
    role='super admin',
    action='reset my password via email',
    benefit='I can regain access if I forget my password',
    acceptance=[
        'Forgot password sends email with reset link (15 min expiry)',
        'Link is one-time use',
        'New password must meet strength requirements',
        'Success → redirected to login',
    ],
    priority='High',
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 5 — ADMIN FPO APPLICATIONS
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 5 — Admin — FPO Applications')

story_table(
    'US-050', 'View FPO Application List',
    role='admin (super or sub)',
    action='view all FPO applications with filters',
    benefit='I can monitor and manage all FPO registrations',
    acceptance=[
        'List shows: FPO name, district, status, application ID, submission date',
        'Filter by status (draft / submitted / approved / rejected / suspended / info_required)',
        'Filter by district',
        'Paginated (20 per page)',
        'Tier badge shown for approved FPOs',
    ],
    priority='High',
)

story_table(
    'US-051', 'View FPO Application Detail',
    role='admin (super or sub)',
    action='view complete details of an FPO application',
    benefit='I can review all information before taking action',
    acceptance=[
        'All 4 wizard steps data visible',
        'Uploaded documents listed with verified/unverified status',
        'Status timeline shown',
        'Current tier and tier history shown',
        'Action buttons shown based on current status',
    ],
    priority='High',
)

story_table(
    'US-052', 'Reject an FPO Application',
    role='super admin',
    action='reject an approved FPO application',
    benefit='Non-compliant FPOs are removed from the platform',
    acceptance=[
        'Reject button visible on approved FPOs',
        'Rejection reason required (min 20 characters)',
        'Status changes to REJECTED',
        'FPO user receives rejection notification',
    ],
    priority='High',
)

story_table(
    'US-053', 'Request Additional Information from FPO',
    role='admin (super or sub)',
    action='request additional information from an FPO',
    benefit='Incomplete applications can be clarified before final decision',
    acceptance=[
        'Request Info button visible on approved FPOs',
        'Reason required (min 20 characters)',
        'Status changes to INFO_REQUIRED',
        'FPO user notified',
    ],
    priority='Medium',
)

story_table(
    'US-054', 'Verify FPO Documents',
    role='admin (super or sub)',
    action='mark an FPO document as verified',
    benefit='Document review status is tracked on the platform',
    acceptance=[
        'Each document has Verify button',
        'Verified documents show verified badge',
        'Verified by and timestamp recorded',
        'Unverified documents count shown on admin dashboard',
    ],
    priority='Medium',
)

story_table(
    'US-055', 'Manually Assign Tier to FPO',
    role='super admin',
    action='manually assign a tier to an FPO',
    benefit='I can override the auto-scored tier when required (e.g. after field visit)',
    acceptance=[
        'Assign Tier button visible on approved FPOs (super admin only)',
        'Modal shows: tier dropdown (A/B/C/D), financial year, notes',
        'Notes required',
        'Tier updated on FPO profile',
        'Audit log entry created with manual_override: true',
        'Tier history table updated',
    ],
    priority='Medium',
)

story_table(
    'US-056', 'View Tier Assessment Answers (Admin)',
    role='admin (super or sub)',
    action='view the tier assessment answers submitted by an FPO',
    benefit='I can review how the FPO scored in each domain',
    acceptance=[
        'Assessment visible on application detail page',
        'All 28 questions with given answers shown',
        'Domain scores and total score visible',
        'Uploaded files (Q6, Q17, Q27) accessible',
        'Shows "Not Assessed" if no assessment submitted',
    ],
    priority='Medium',
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 6 — ADMIN DASHBOARD
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 6 — Admin Dashboard')

story_table(
    'US-060', 'View Admin Dashboard Stats',
    role='admin (super or sub)',
    action='view the admin dashboard with platform statistics',
    benefit='I have a real-time overview of platform activity',
    acceptance=[
        'Stat cards: Total Registrations, Approved FPOs, Pending Applications, Suspended',
        'Status breakdown donut chart (7 statuses, all always present)',
        'Tier distribution bar chart (A/B/C/D/not_assessed)',
        'District distribution table (ordered by count)',
        'Monthly registration trend line chart (last 12 months, no gaps)',
        'Pending actions section: ownership claims, unverified docs, info_required FPOs',
        '"No pending actions" shown if all counts are 0',
    ],
    priority='High',
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 7 — ADMIN SUB-ADMIN MANAGEMENT
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 7 — Admin — Sub-Admin Management')

story_table(
    'US-070', 'Create Sub-Admin',
    role='super admin',
    action='create a sub-admin account for KAU staff',
    benefit='KAU staff can manage FPOs with controlled access',
    acceptance=[
        'Form: first name, last name, email, phone',
        'Temporary password auto-generated',
        'Welcome email sent with temporary password',
        'Sub-admin must change password on first login',
        'Sub-admin appears in list immediately',
    ],
    priority='High',
)

story_table(
    'US-071', 'Assign Permissions to Sub-Admin',
    role='super admin',
    action='assign specific permissions to a sub-admin',
    benefit='Each KAU staff member has access only to what they need',
    acceptance=[
        'Available permissions: can_approve_fpo, can_view_all_fpos, can_request_info, can_verify_documents, can_generate_reports',
        'Permissions can be added, removed, or replaced',
        'Changes take effect immediately',
    ],
    priority='High',
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 8 — ADMIN OWNERSHIP CLAIMS
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 8 — Admin — Ownership Claims')

story_table(
    'US-080', 'Review Ownership Claim List',
    role='admin (super or sub)',
    action='view all ownership claims submitted by FPO users',
    benefit='I can identify and resolve duplicate FPO ownership disputes',
    acceptance=[
        'Table shows: claimant name, email, phone, FPO name, reason (truncated), status, submitted date',
        'Filter by status (pending / approved / rejected)',
        'Filter by FPO ID',
        'Paginated',
    ],
    priority='High',
)

story_table(
    'US-081', 'Approve Ownership Claim',
    role='super admin',
    action='approve an ownership claim',
    benefit='The legitimate owner gains control of their FPO',
    acceptance=[
        'Approve button visible (super admin only)',
        'Notes required (min 10 characters)',
        'FPO ownership transferred to claimant',
        'Old primary user access deactivated',
        'All other pending claims on same FPO auto-rejected',
        'Audit log entry created',
        'Claimant notified (when notification is wired)',
    ],
    priority='High',
)

story_table(
    'US-082', 'Reject Ownership Claim',
    role='admin (super or sub)',
    action='reject an ownership claim',
    benefit='Fraudulent claims are dismissed',
    acceptance=[
        'Reject button visible to both super admin and sub admin',
        'Notes required (min 10 characters)',
        'Claim status changes to rejected',
        'FPO ownership unchanged',
        'Claimant notified with reason (when notification is wired)',
    ],
    priority='High',
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 9 — ADMIN SCHEMES & EXPERTS
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 9 — Admin — Schemes & Expert Directory')

story_table(
    'US-090', 'Manage Schemes',
    role='admin (super or sub)',
    action='create, edit, and manage government schemes',
    benefit='FPOs always have up-to-date scheme information',
    acceptance=[
        '12 seeded schemes visible on first load',
        'Create scheme: required fields are name_en, administering_body, category, eligibility, benefit_details, application_process',
        'name_ml, objective, official_link are optional',
        'Edit (PATCH) any field',
        'Activate / Deactivate scheme (deactivated schemes hidden from FPO portal)',
        'Delete scheme (soft delete)',
    ],
    priority='Medium',
)

story_table(
    'US-091', 'Manage Expert Directory',
    role='admin (super or sub)',
    action='manage the expert directory',
    benefit='FPOs have access to current expert contact information',
    acceptance=[
        '15 seeded experts visible on first load',
        'Create expert: required fields are name_en, designation, organisation, primary_expertise, category, email',
        'Edit (PATCH) any field',
        'Activate / Deactivate expert',
        'Delete expert (soft delete)',
        '"Enquiries" tab shows messages received by that expert',
    ],
    priority='Medium',
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 10 — ADMIN CMS
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 10 — Admin — Site Content CMS')

story_table(
    'US-100', 'Edit Landing Page Content Blocks',
    role='super admin',
    action='edit the hero section and about section text on the landing page',
    benefit='KAU can update platform messaging without developer involvement',
    acceptance=[
        'All 6 blocks visible: hero_headline, hero_subheading, hero_description, about_title, about_body, how_to_register',
        'PATCH sends {"content": {"en": "...", "ml": "..."}}',
        'English required, Malayalam optional',
        'Changes visible on public landing page immediately',
    ],
    priority='Medium',
)

story_table(
    'US-101', 'Manage Announcements',
    role='admin (super or sub)',
    action='create, edit, and delete announcements',
    benefit='Platform news is always current and relevant',
    acceptance=[
        '4 seeded announcements visible on first load',
        'Create: title.en, body.en, category required; title.ml, body.ml optional',
        'Category: announcement or news',
        'Published date optional',
        'Inactive announcements hidden from public page',
        'Edit and delete work correctly',
    ],
    priority='Medium',
)

story_table(
    'US-102', 'Manage FAQs',
    role='admin (super or sub)',
    action='create, edit, and manage FAQs',
    benefit='Users can find answers to common questions in their language',
    acceptance=[
        '51 seeded FAQs visible on first load (25 FPO + 26 Schemes)',
        'Create: question.en, answer.en, category required',
        'Malayalam can be added via PATCH to existing FAQ',
        'Category: fpo_general / schemes / platform_usage',
        'Deactivated FAQs hidden from public page',
        'Order field controls display sequence',
    ],
    priority='Medium',
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 11 — ADMIN AUDIT LOGS
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 11 — Admin — Audit Logs')

story_table(
    'US-110', 'View Audit Trail',
    role='admin (super or sub)',
    action='view the platform audit trail',
    benefit='All significant actions are traceable for compliance and review',
    acceptance=[
        'List shows: action type, user, FPO (if applicable), timestamp, changes summary',
        'Filter by: action type, FPO ID, user ID, date range',
        'Action types include: FPO_STATUS_CHANGE, TIER_RECALCULATION, DOCUMENT_UPLOAD, FPO_SUBMIT, etc.',
        'Paginated',
        'changes field shows before/after values for FPO profile changes',
        'Tier entries show: tier, financial_year, total_score, domain_scores (auto) or manual_override: true (admin)',
    ],
    priority='Medium',
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 12 — MULTILINGUAL
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 12 — Multilingual Support')

story_table(
    'US-120', 'Language Switch — English to Malayalam',
    role='any user',
    action='switch the platform language from English to Malayalam',
    benefit='Malayalam-speaking users can use the platform in their preferred language',
    acceptance=[
        'Language switch sets X-Language: ml header on all API calls',
        'Hero text, about text, announcements switch to Malayalam',
        'Scheme names returned in Malayalam',
        'Expert names returned in Malayalam (if available)',
        'FAQ text returned in Malayalam (if available, else English fallback)',
        'All admin UI label keys resolved in Malayalam via Translation table',
        'Switching back to English restores English content',
    ],
    priority='High',
)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 13 — ADMIN MENU MANAGEMENT
# ═════════════════════════════════════════════════════════════════════════════

heading1('Section 13 — Admin — Menu Management')
body('Menu items control what pages appear in the sidebar for each role. Super admin manages them via the Menu page in the admin portal.')

story_table(
    'US-130', 'View All Menu Items',
    role='super admin',
    action='view all sidebar menu items and their role assignments',
    benefit='I can see which pages are visible to each role without navigating to each one',
    acceptance=[
        'Table shows all menu items: label, path, icon, assigned roles, order, active status',
        'Items for admin portal and FPO portal listed separately',
        'Role column shows which groups (super_admin, sub_admin, primary, secondary) can see each item',
        'Active / Inactive status badge visible per item',
    ],
    priority='Medium',
)

story_table(
    'US-131', 'Create a New Menu Item',
    role='super admin',
    action='add a new page to the sidebar navigation',
    benefit='New features become accessible to the correct roles without a code deployment',
    acceptance=[
        'Form fields: label key (translation key), path (frontend route), icon name, roles, order',
        'label_key must follow format menu.<key> (e.g. menu.fpo_reports)',
        'Roles is a multi-select from existing Django groups',
        'New item immediately appears in sidebar for selected roles after save',
        'Duplicate path is rejected',
    ],
    priority='Medium',
)

story_table(
    'US-132', 'Edit a Menu Item',
    role='super admin',
    action='edit an existing menu item (label, icon, order, roles)',
    benefit='I can update navigation without touching code',
    acceptance=[
        'PATCH updates any combination of: label_key, icon, order, roles',
        'Changing roles takes effect immediately — affected users see updated sidebar on next page load',
        'path (route) is editable',
        'Order field controls the position in the sidebar (ascending)',
    ],
    priority='Medium',
)

story_table(
    'US-133', 'Activate / Deactivate a Menu Item',
    role='super admin',
    action='deactivate a menu item to hide it from the sidebar',
    benefit='I can temporarily hide a page from all users without deleting it',
    acceptance=[
        'Deactivate button visible on each menu item row',
        'Deactivated items immediately disappear from all users\' sidebars',
        'Item still visible in admin menu list with Inactive badge',
        'Activate button re-enables the item',
        'GET /api/auth/me/ excludes inactive menu items from the menu list',
    ],
    priority='Medium',
)

story_table(
    'US-134', 'Role-Filtered Sidebar (Me API)',
    role='any authenticated user',
    action='see only the menu items relevant to my role when I log in',
    benefit='I am not confused by pages I don\'t have access to',
    acceptance=[
        'GET /api/auth/me/ returns menu list filtered to user\'s assigned groups',
        'Super admin sees: languages, notifications, roles, sub-admins, FPO permissions, applications, external APIs, site content, announcements, FAQs',
        'Sub admin sees: applications, announcements, FAQs (based on default seed)',
        'FPO primary/secondary see: dashboard, register, status, profile, applications, recommendations, products, market, settings',
        'Inactive items never appear regardless of role',
        'Menu labels returned in user\'s preferred language (X-Language header)',
    ],
    priority='High',
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# APPENDIX — SEED COMMANDS
# ═════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
heading1('Appendix — Seed Commands (Test Environment Setup)')
body('Run these commands to populate test data before running test scenarios.')

for cmd_title, cmd in [
    ('Translations (UI strings)',
     'source venv/bin/activate && python manage.py shell -c "\nexec(open(\'scripts/seed_translations.py\').read())\nseed_translations()\n"'),
    ('Menu items',
     'source venv/bin/activate && python manage.py shell -c "\nexec(open(\'scripts/seed_menu.py\').read())\nseed_menu()\n"'),
    ('FPO Master Data (districts, blocks, commodities)',
     'source venv/bin/activate && python manage.py shell -c "\nexec(open(\'scripts/seed_fpo_master_data.py\').read())\nseed_fpo_master_data()\n"'),
    ('Tier Assessment Framework (domains, criteria, questions)',
     'source venv/bin/activate && python manage.py shell -c "\nexec(open(\'scripts/seed_tier_framework.py\').read())\nseed_tier_framework()\n"'),
    ('Schemes (12 government schemes)',
     'source venv/bin/activate && python manage.py shell -c "\nexec(open(\'scripts/seed_schemes.py\').read())\nseed_schemes()\n"'),
    ('Expert Directory (15 KAU experts)',
     'source venv/bin/activate && python manage.py shell -c "\nexec(open(\'scripts/seed_experts.py\').read())\nseed_experts()\n"'),
    ('Site Content CMS (blocks, announcements, FAQs)',
     'source venv/bin/activate && python manage.py shell -c "\nexec(open(\'scripts/seed_cms.py\').read())\nseed_cms()\n"'),
    ('FPO Permissions',
     'source venv/bin/activate && python manage.py shell -c "\nexec(open(\'scripts/seed_fpo_permissions.py\').read())\nseed_fpo_permissions()\n"'),
]:
    heading3(cmd_title)
    p = doc.add_paragraph()
    run = p.add_run(cmd)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

# ─── Save ─────────────────────────────────────────────────────────────────────
output_path = 'docs/KAU-FPO-User-Stories.docx'
doc.save(output_path)
print(f'✅ Saved: {output_path}')
