"""
User Manual Generator — KAU-FPO Admin Portal
KAU-FPO Platform — house style (matching RCD / FPO Registration / Landing Page manuals)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paths ────────────────────────────────────────────────────────────────────
KAU_LOGO    = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/KAU Emblem 0.5x0.75.jpg'
SCREENSHOTS = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/Documents/Usermanual/screenshots/admin_portal'
OUTPUT      = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/Documents/Usermanual/KAU_FPO_UserManual_AdminPortal_v1.0.docx'

# ─── Colors ───────────────────────────────────────────────────────────────────
NAVY_HEX       = '1F3864'
DARK_NAVY      = RGBColor(0x1F, 0x38, 0x64)
ORANGE         = RGBColor(0xE8, 0x6C, 0x1A)
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
GREY           = RGBColor(0x60, 0x60, 0x60)
BLACK          = RGBColor(0x00, 0x00, 0x00)
LIGHT_GREY_HEX = 'F2F2F2'


# ─── XML Helpers ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, start=120, bottom=60, end=120):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'),    str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def remove_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'none')
        b.set(qn('w:sz'),    '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_bottom_border(paragraph, color=NAVY_HEX, sz='8'):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_page_number(run):
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)


def set_col_width(cell, width_inches):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def set_table_width(table, width_cm):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW  = OxmlElement('w:tblW')
    twips = int(width_cm * 567)
    tblW.set(qn('w:w'),    str(twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


# ─── Header / Footer ──────────────────────────────────────────────────────────

def build_header(header_obj, title='User Manual — KAU-FPO Admin Portal'):
    for p in header_obj.paragraphs:
        p._element.getparent().remove(p._element)

    tbl = header_obj.add_table(rows=1, cols=3, width=Inches(6.5))
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)

    c0 = tbl.cell(0, 0)
    set_col_width(c0, 1.15)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(0)
    r0 = p0.add_run()
    r0.add_picture(KAU_LOGO, width=Cm(0.8), height=Cm(0.49))

    c1 = tbl.cell(0, 1)
    set_col_width(c1, 3.82)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after  = Pt(0)
    r1 = p1.add_run(title)
    r1.font.bold      = True
    r1.font.size      = Pt(9)
    r1.font.color.rgb = DARK_NAVY

    c2 = tbl.cell(0, 2)
    set_col_width(c2, 1.30)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run('KAU-FPO Platform')
    r2.font.size      = Pt(9)
    r2.font.color.rgb = DARK_NAVY

    div = header_obj.add_paragraph()
    div.paragraph_format.space_before = Pt(2)
    div.paragraph_format.space_after  = Pt(0)
    add_bottom_border(div)


def build_first_page_header(header_obj):
    for p in header_obj.paragraphs:
        p._element.getparent().remove(p._element)
    header_obj.add_paragraph()


def build_footer(footer_obj):
    for p in footer_obj.paragraphs:
        p._element.getparent().remove(p._element)

    tbl = footer_obj.add_table(rows=1, cols=2, width=Inches(6.5))
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)

    cl = tbl.cell(0, 0)
    set_col_width(cl, 4.5)
    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rl = pl.add_run('Kerala Agricultural University  │  fpolinkage.kau.in')
    rl.font.size      = Pt(8)
    rl.font.color.rgb = GREY

    cr = tbl.cell(0, 1)
    set_col_width(cr, 1.77)
    pr = cr.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = pr.add_run('Page ')
    rr.font.size      = Pt(8)
    rr.font.color.rgb = GREY
    rn = pr.add_run()
    rn.font.size      = Pt(8)
    rn.font.color.rgb = GREY
    add_page_number(rn)


# ─── Content Helpers ──────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold      = True
    r.font.size      = Pt(16)
    r.font.name      = 'Calibri'
    r.font.color.rgb = DARK_NAVY
    add_bottom_border(p)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold      = True
    r.font.size      = Pt(12)
    r.font.name      = 'Calibri'
    r.font.color.rgb = ORANGE
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold      = True
    r.font.size      = Pt(11)
    r.font.name      = 'Calibri'
    r.font.color.rgb = BLACK
    return p


def body(doc, text, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Cambria'
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Cambria'
    return p


def numbered(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(18)
    r1 = p.add_run(f'{number}.  ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Cambria'
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.name = 'Cambria'
    return p


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    return p


def make_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.style = 'Table Grid'

    for j, h_text in enumerate(headers):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, NAVY_HEX)
        set_cell_margins(c)
        if col_widths:
            set_col_width(c, col_widths[j])
        p = c.paragraphs[0]
        r = p.add_run(h_text)
        r.font.bold      = True
        r.font.color.rgb = WHITE
        r.font.size      = Pt(10)
        r.font.name      = 'Calibri'

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            c = tbl.rows[i + 1].cells[j]
            set_cell_margins(c)
            if col_widths:
                set_col_width(c, col_widths[j])
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = 'Cambria'

    spacer(doc, 8)
    return tbl


def screenshot(doc, filename, caption=None):
    import os
    path = os.path.join(SCREENSHOTS, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run()
    r.add_picture(path, width=Cm(14))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        cr = cp.add_run(caption)
        cr.font.italic    = True
        cr.font.size      = Pt(9)
        cr.font.name      = 'Cambria'
        cr.font.color.rgb = GREY


def note_box(doc, text):
    spacer(doc, 4)
    tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)
    c = tbl.cell(0, 0)
    set_cell_bg(c, LIGHT_GREY_HEX)
    set_cell_margins(c, top=100, start=160, bottom=100, end=160)
    p = c.paragraphs[0]
    r1 = p.add_run('Note: ')
    r1.font.bold      = True
    r1.font.size      = Pt(10)
    r1.font.name      = 'Calibri'
    r1.font.color.rgb = DARK_NAVY
    r2 = p.add_run(text)
    r2.font.size      = Pt(10)
    r2.font.name      = 'Cambria'
    spacer(doc, 6)


def module_banner(doc, title, subtitle):
    tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)
    c = tbl.cell(0, 0)
    set_cell_bg(c, NAVY_HEX)
    set_cell_margins(c, top=100, start=160, bottom=60, end=160)
    p = c.paragraphs[0]
    r1 = p.add_run(title)
    r1.font.bold      = True
    r1.font.size      = Pt(11)
    r1.font.name      = 'Calibri'
    r1.font.color.rgb = WHITE
    if subtitle:
        p2 = c.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        r2 = p2.add_run(subtitle)
        r2.font.size      = Pt(9)
        r2.font.name      = 'Cambria'
        r2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    spacer(doc, 6)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Cambria'
style.font.size = Pt(11)

sec = doc.sections[0]
sec.page_width        = Cm(21)
sec.page_height       = Cm(29.7)
sec.left_margin       = Cm(2.54)
sec.right_margin      = Cm(2.54)
sec.top_margin        = Cm(2.0)
sec.bottom_margin     = Cm(2.0)
sec.header_distance   = Cm(1.0)
sec.footer_distance   = Cm(1.0)
sec.different_first_page_header_footer = True

build_header(sec.header)
build_footer(sec.footer)
build_first_page_header(sec.first_page_header)
build_footer(sec.first_page_footer)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(12)
p.add_run().add_picture(KAU_LOGO, width=Cm(9.14))

banner_tbl = doc.add_table(rows=1, cols=1)
banner_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(banner_tbl, 16.51)
bc = banner_tbl.cell(0, 0)
set_cell_bg(bc, NAVY_HEX)
set_cell_margins(bc, top=160, start=200, bottom=160, end=200)
bp = bc.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
br = bp.add_run('USER MANUAL — KAU-FPO ADMIN PORTAL')
br.font.bold      = True
br.font.size      = Pt(20)
br.font.name      = 'Calibri'
br.font.color.rgb = WHITE

spacer(doc, 14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('KERALA AGRICULTURAL UNIVERSITY')
r.font.bold      = True
r.font.size      = Pt(16)
r.font.name      = 'Calibri'
r.font.color.rgb = DARK_NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run('Communication Centre, Mannuthy, Thrissur — 680651')
r.font.size      = Pt(11)
r.font.name      = 'Cambria'
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('AI-Based Digital Platform for KAU-FPO Linkage Programme')
r.font.bold      = True
r.font.size      = Pt(13)
r.font.name      = 'Calibri'
r.font.color.rgb = ORANGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run('User Manual — Admin Portal')
r.font.bold      = True
r.font.size      = Pt(13)
r.font.name      = 'Calibri'
r.font.color.rgb = ORANGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Version 1.0  |  Kerala Agricultural University')
r.font.size      = Pt(11)
r.font.name      = 'Cambria'
r.font.color.rgb = GREY

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — DOCUMENT INFORMATION
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '1. Document Information')

make_table(doc,
    headers=['Field', 'Details'],
    rows=[
        ('Application Name', 'AI-Based Digital Platform for KAU-FPO Linkage Programme'),
        ('Module',           'Admin Portal'),
        ('Version',          '1.0'),
        ('Prepared By',      'Kerala Agricultural University'),
        ('Date',             'August 2026'),
        ('Audience',         'KAU Super Administrators, Sub-Administrators'),
        ('Platform URL',     'https://fpolinkage.kau.in/admin'),
    ],
    col_widths=[2.2, 4.07],
)

h2(doc, '1.1  Purpose')
body(doc,
    'This manual covers all modules available in the KAU-FPO Admin Portal. It is intended '
    'for KAU super administrators and sub-administrators who manage the platform. The portal '
    'requires an authorized admin account and is accessible only after login.')

h2(doc, '1.2  System Requirements')
make_table(doc,
    headers=['Component', 'Minimum Requirement'],
    rows=[
        ('Web Browser',       'Google Chrome 110+, Mozilla Firefox 110+, Microsoft Edge 110+, Safari 16+'),
        ('Internet',          'Broadband connection (2 Mbps or above recommended)'),
        ('Screen Resolution', '1280 × 720 or higher (1920 × 1080 recommended for admin tables)'),
        ('JavaScript',        'Must be enabled in browser settings'),
        ('Authentication',    'Admin account credentials — email and password'),
    ],
    col_widths=[2.2, 4.07],
)

h2(doc, '1.3  Admin Portal Modules')
make_table(doc,
    headers=['Module', 'Purpose', 'Who Uses It'],
    rows=[
        ('Dashboard',              'Platform overview — stats, charts, district map',         'All admins'),
        ('Languages & Translations','Manage platform languages and UI text',                  'Super admin'),
        ('Notifications',          'Manage notification templates and channel settings',      'Super admin'),
        ('Sub-Admins',             'Create and manage sub-administrator accounts',            'Super admin'),
        ('FPO Actions',            'Manage FPO permission actions, member roles, matrix',     'Super admin'),
        ('Applications',           'Review and manage FPO registration applications',         'All admins'),
        ('External APIs',          'Configure PAN / GSTIN / CIN verification API keys',      'Super admin'),
        ('Site Content',           'Edit landing page content, gallery, team, documents',     'All admins'),
        ('Announcements',          'Publish news and announcements on the public portal',     'All admins'),
        ('FAQs',                   'Manage FAQ entries shown on the public FAQ page',         'All admins'),
        ('Ownership Claims',       'Review and process FPO ownership transfer requests',      'All admins'),
        ('Audit Logs',             'View complete system event trail (login, actions)',        'Super admin'),
        ('Experts',                'Manage the expert directory visible to FPOs',             'All admins'),
        ('Schemes & Subsidies',    'Manage government schemes catalog',                       'All admins'),
        ('Settings',               'Edit own profile, phone, language, security (2FA)',       'All admins'),
    ],
    col_widths=[1.8, 2.8, 1.67],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — LOGGING IN
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '2. Logging In to the Admin Portal')

h2(doc, '2.1  How to Log In')
numbered(doc, 1, 'Open a web browser and navigate to https://fpolinkage.kau.in/admin')
numbered(doc, 2, 'Enter your Email address and Password in the login form.')
numbered(doc, 3, 'Click the Sign In button.')
numbered(doc, 4,
    'If Two-Factor Authentication (2FA) is enabled on your account, enter the '
    '6-digit code from your authenticator app when prompted.')
numbered(doc, 5, 'After successful login, the Admin Dashboard is displayed.')

note_box(doc,
    'If you are a super administrator logging in for the first time after account creation, '
    'you will be prompted to change your temporary password before accessing the portal. '
    'If 2FA is mandatory for your account, complete the 2FA setup before your first session.')

h2(doc, '2.2  Forgot Password')
numbered(doc, 1, 'On the login page, click the "Forgot Password?" link.')
numbered(doc, 2, 'Enter your registered email address and submit.')
numbered(doc, 3, 'A password reset link will be sent to your email. The link expires in 15 minutes.')
numbered(doc, 4, 'Click the link in the email, enter your new password, and confirm.')

h2(doc, '2.3  Logging Out')
bullet(doc, 'Click Logout in the bottom of the left navigation sidebar at any time.')
bullet(doc, 'Your session will be terminated and you will be redirected to the login page.')

note_box(doc,
    'Always log out when leaving the workstation unattended. Admin sessions contain '
    'sensitive FPO data. Do not share your login credentials with anyone.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PORTAL LAYOUT
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '3. Portal Layout & Navigation')
body(doc,
    'All pages in the Admin Portal follow the same layout: a fixed top bar, a left '
    'navigation sidebar, and a main content area.')

screenshot(doc, 'screenshot_06.png', 'Figure 1 — Left navigation sidebar showing all modules')

h2(doc, '3.1  Top Bar Controls')
make_table(doc,
    headers=['Control', 'Description'],
    rows=[
        ('Page title',         'Shows the name of the currently open module (e.g. "Dashboard", "Audit Logs")'),
        ('A−  A+',             'Text size controls — decrease or increase font size for accessibility'),
        ('Dark mode toggle',   'Switch between light and dark theme'),
        ('Notification bell',  'View recent platform notifications'),
        ('Language selector',  'Switch portal display language (English / Malayalam). "default" follows user preference'),
    ],
    col_widths=[2.0, 4.27],
)

h2(doc, '3.2  Left Navigation Sidebar')
body(doc,
    'The sidebar lists all available modules. The currently active module is highlighted '
    'in navy blue. Scroll down inside the sidebar to see all modules. The logged-in '
    'administrator\'s name and email are shown at the bottom of the sidebar.')

make_table(doc,
    headers=['Sidebar Item', 'Module'],
    rows=[
        ('Dashboard',            '§4 — Platform overview and statistics'),
        ('Languages & Translations', '§5 — Manage languages, translation keys, menu items'),
        ('Notifications',        '§6 — Notification templates, channel settings'),
        ('Sub-Admins',           '§7 — Create and manage sub-administrator accounts'),
        ('FPO Actions',          '§8 — FPO permission actions, member roles, matrix'),
        ('Applications',         '§9 — FPO registration applications'),
        ('External APIs',        '§10 — API credentials for PAN / GSTIN / CIN verification'),
        ('Site Content',         '§11 — Landing page content, gallery, team, documents, links'),
        ('Announcements',        '§12 — News and announcements for the public portal'),
        ('FAQs',                 '§13 — FAQ entries for the public FAQ page'),
        ('Ownership Claims',     '§14 — FPO ownership transfer requests'),
        ('Audit Logs',           '§15 — System event trail'),
        ('Experts',              '§16 — Expert directory management'),
        ('Schemes & Subsidies',  '§17 — Government schemes catalog'),
        ('Settings',             '§18 — Profile and security settings'),
        ('Logout',               'Terminate the current admin session'),
    ],
    col_widths=[2.0, 4.27],
)

h2(doc, '3.3  Common Controls on Every Module Page')
body(doc,
    'Most module pages share a consistent set of controls in the toolbar area above the '
    'data table:')
make_table(doc,
    headers=['Control', 'Description'],
    rows=[
        ('Search bar',       'Type to filter table rows by name, email, or other key fields'),
        ('Filter dropdowns', 'Narrow results by status, category, district, or other attributes'),
        ('Columns',          'Show or hide specific table columns for a customised view'),
        ('Refresh (↺)',      'Reload the table with the latest data from the server'),
        ('+ Add button',     'Open the form to create a new record (e.g. + Add Expert, + Add Scheme)'),
        ('⋯ (three dots)',   'Row-level actions menu — Edit, Activate, Deactivate, and more'),
        ('← Back',           'Return to the previous page'),
    ],
    col_widths=[1.8, 4.47],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '4. Dashboard')
body(doc,
    'The Admin Dashboard is the first page displayed after login. It provides a real-time '
    'overview of all FPO registrations, application statuses, and platform activity.')

screenshot(doc, 'screenshot_11.png', 'Figure 2 — Full Admin Dashboard')

h2(doc, '4.1  Summary Cards')
body(doc,
    'Four stat cards at the top of the dashboard show the current platform numbers at a glance.')

screenshot(doc, 'screenshot_07.png', 'Figure 3 — Dashboard summary stat cards')

make_table(doc,
    headers=['Card', 'Color', 'What It Shows'],
    rows=[
        ('Total Registrations',  'Purple',  'Total number of FPOs registered across all statuses'),
        ('Approved FPOs',        'Green',   'FPOs that are active and operational (status: Approved)'),
        ('Pending Applications', 'Orange',  'FPOs awaiting administrative review'),
        ('Suspended',            'Grey',    'FPOs whose accounts have been suspended'),
    ],
    col_widths=[2.0, 1.2, 3.07],
)

h2(doc, '4.2  Registration Trend')
body(doc,
    'The Registration Trend chart shows monthly FPO registrations for the past 12 months '
    'as a bar chart. Use this to monitor platform growth and registration activity over time.')

screenshot(doc, 'screenshot_08.png', 'Figure 4 — Monthly Registration Trend chart')

h2(doc, '4.3  Status Breakdown')
body(doc,
    'The Status Breakdown donut chart shows the distribution of all registered FPOs by '
    'their current status (Approved, Draft, Claimed, Rejected, Suspended). Hover over a '
    'segment to see the count for that status.')

screenshot(doc, 'screenshot_09.png', 'Figure 5 — Status Breakdown donut chart')

h2(doc, '4.4  Action Required')
body(doc,
    'The Action Required panel highlights items that need administrator attention — '
    'pending applications and unreviewed ownership claims. Click any item to navigate '
    'directly to the relevant module.')

h2(doc, '4.5  Tier Distribution')
body(doc,
    'The Tier Distribution section shows how many FPOs fall into each tier '
    '(Tier 1 through Tier 4) based on their annual assessment scores.')

h2(doc, '4.6  District Distribution Map')
body(doc,
    'An interactive map of Kerala shows the number of registered FPOs in each district. '
    'Districts are color-coded by FPO count — darker shading indicates higher registration '
    'density. Hover over a district to see the FPO count.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — LANGUAGES & TRANSLATIONS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '5. Languages & Translations')
body(doc,
    'The Languages & Translations module manages the languages supported by the platform '
    'and all translatable text strings. It is accessible from the left navigation sidebar.')

h2(doc, '5.1  How to Open')
numbered(doc, 1, 'Click Languages & Translations in the left navigation sidebar.')
numbered(doc, 2, 'The page opens with the Languages tab selected by default.')

h2(doc, '5.2  Tabs')
make_table(doc,
    headers=['Tab', 'Purpose'],
    rows=[
        ('Languages',     'Add, activate, or deactivate supported languages (e.g. English, Malayalam)'),
        ('Categories',    'Manage translation category groupings (e.g. auth, fpo, admin, ui)'),
        ('Translations',  'View, add, verify, bulk-import, and export individual translation strings'),
        ('Menu Items',    'Manage sidebar menu labels and assign them to roles'),
    ],
    col_widths=[1.8, 4.47],
)

h2(doc, '5.3  Language Management')
bullet(doc, 'Click + Add Language to add a new supported language.')
bullet(doc, 'Click ⋯ on any language row to Activate, Deactivate, or Set as Default.')
bullet(doc, 'Only Active languages appear in the platform language selector.')
bullet(doc, 'The Default language is used when no language preference is detected.')

h2(doc, '5.4  Translation Management')
bullet(doc, 'Click the Translations tab to see all translation keys.')
bullet(doc, 'Use the Search bar and Category filter to find specific keys.')
bullet(doc, 'Click + Add Translation to add a new translation string.')
bullet(doc, 'Translations can be Verified (confirmed accurate) or Unverified (pending review).')
bullet(doc, 'Click Export to download an Excel file of all untranslated keys for a language.')
bullet(doc, 'Click Import to upload a completed Excel file with translations filled in.')

note_box(doc,
    'Translation keys follow the format category.key_name (e.g. auth.login_success, '
    'fpo.registration_complete). Never edit the key name — only edit the translated value.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — NOTIFICATIONS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '6. Notifications')
body(doc,
    'The Notifications module manages notification templates (email, SMS, in-app, WhatsApp) '
    'and channel configuration. All platform notifications are delivered through templates '
    'configured in this module.')

h2(doc, '6.1  How to Open')
numbered(doc, 1, 'Click Notifications in the left navigation sidebar.')
numbered(doc, 2, 'The page opens with Template Codes selected by default.')

h2(doc, '6.2  Notification Sections')
make_table(doc,
    headers=['Section', 'Purpose'],
    rows=[
        ('Template Codes',    'Named event codes that trigger notifications (e.g. fpo_approved, dpr_generated)'),
        ('Templates',         'The actual message content per code, channel, and language'),
        ('Channel Settings',  'Credentials and activation status for each delivery channel'),
    ],
    col_widths=[1.8, 4.47],
)

h2(doc, '6.3  Template Codes')
body(doc,
    'Template codes are the named events that trigger a notification. Each code can have '
    'templates in multiple languages and for multiple channels.')
bullet(doc, 'Click ⋯ on a code to Activate or Deactivate it.')
bullet(doc, 'The Missing column shows which languages have no template for that code.')
bullet(doc, 'Click a code to view its linked templates.')

h2(doc, '6.4  Templates')
body(doc,
    'Each template contains the subject, body, and variables for a specific channel and language.')
bullet(doc, 'Variables in the body are written as {{variable_name}} and are replaced with real values at send time.')
bullet(doc, 'Click + Add Template to create a new template for a code.')
bullet(doc, 'Use the Test Render action to preview how a template looks with sample variable values.')

h2(doc, '6.5  Channel Settings')
make_table(doc,
    headers=['Channel', 'Config Required'],
    rows=[
        ('Email (SMTP)',   'Host, port, username, password, from_email, from_name, TLS setting'),
        ('SMS (MSG91 / KAU SOAP)', 'API key or Token, sender ID, base URL, template ID'),
        ('In-App',        'No credentials required — writes directly to the notification inbox'),
        ('WhatsApp',      'Phone Number ID, Access Token, API version (Meta Business API)'),
    ],
    col_widths=[1.8, 4.47],
)

note_box(doc,
    'Channel credentials are stored encrypted. Sensitive fields (password, token, api_key) '
    'are masked in the UI as ••••••••. Only super administrators can view or edit channel settings.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — SUB-ADMINS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '7. Sub-Admins')
body(doc,
    'The Sub-Admins module allows super administrators to create KAU staff accounts with '
    'configurable subsets of admin permissions. Sub-admins can review FPO applications '
    'and perform assigned tasks but cannot access system configuration modules.')

h2(doc, '7.1  How to Open')
numbered(doc, 1, 'Click Sub-Admins in the left navigation sidebar.')

h2(doc, '7.2  Creating a Sub-Admin')
numbered(doc, 1, 'Click + Add Sub-Admin.')
numbered(doc, 2, 'Fill in First Name, Last Name, Email, and Phone.')
numbered(doc, 3, 'Click Save. A temporary password is generated and sent to the sub-admin via email and SMS.')
numbered(doc, 4,
    'The sub-admin must change the temporary password on first login '
    '(must_change_password is set automatically).')

h2(doc, '7.3  Managing Sub-Admin Permissions')
body(doc,
    'Each sub-admin can be assigned a specific subset of permissions. Permissions are '
    'independent per user — two sub-admins can have different access levels.')

make_table(doc,
    headers=['Permission', 'What It Allows'],
    rows=[
        ('can_approve_fpo',     'Approve or reject FPO registration applications'),
        ('can_view_all_fpos',   'View all FPO profiles across all districts'),
        ('can_request_info',    'Send a Request for Information to an FPO'),
        ('can_verify_documents','Mark uploaded FPO documents as verified'),
        ('can_generate_reports','Export FPO summary reports (PDF / Excel)'),
    ],
    col_widths=[2.2, 4.07],
)

bullet(doc, 'Click ⋯ → Permissions on a sub-admin row to view or edit their permissions.')
bullet(doc, 'Use Add, Remove, or Replace actions to update the permission set.')

h2(doc, '7.4  Activating / Deactivating a Sub-Admin')
bullet(doc, 'Click ⋯ → Activate to restore access for a deactivated sub-admin.')
bullet(doc, 'Click ⋯ → Deactivate to immediately block portal access without deleting the account.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — FPO ACTIONS (PERMISSIONS MATRIX)
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '8. FPO Actions & Permissions Matrix')
body(doc,
    'The FPO Actions module manages the actions available within the FPO portal, the '
    'member roles (Primary User and Secondary User), and the permissions matrix that '
    'controls which actions each role can perform.')

h2(doc, '8.1  How to Open')
numbered(doc, 1, 'Click FPO Actions in the left navigation sidebar.')
numbered(doc, 2, 'The page opens with the FPO Actions tab selected by default.')

h2(doc, '8.2  Three Tabs')
make_table(doc,
    headers=['Tab', 'Purpose'],
    rows=[
        ('FPO Actions',      'List of all permission actions for the FPO portal (e.g. can_submit, can_upload_docs)'),
        ('Member Roles',     'The two FPO user roles: Primary User and Secondary User'),
        ('Permissions Matrix', 'A grid showing which roles have which actions enabled'),
    ],
    col_widths=[2.0, 4.27],
)

h2(doc, '8.3  FPO Actions')
make_table(doc,
    headers=['Action Code', 'Description'],
    rows=[
        ('can_view_dashboard',  'View the FPO dashboard'),
        ('can_submit',          'Submit the FPO registration application'),
        ('can_upload_docs',     'Upload registration documents'),
        ('can_delete_docs',     'Delete uploaded documents (Draft status only)'),
        ('can_view_docs',       'View uploaded documents'),
        ('can_edit_profile',    'Edit FPO profile information'),
        ('can_invite_team',     'Invite secondary users to the FPO team'),
        ('can_manage_team',     'Activate, deactivate, and reset team member passwords'),
        ('can_submit_claim',    'Submit an ownership claim for an existing FPO'),
    ],
    col_widths=[2.2, 4.07],
)

h2(doc, '8.4  Permissions Matrix')
body(doc,
    'The Permissions Matrix shows a role × action grid. Check or uncheck cells to '
    'enable or disable an action for a role. Changes take effect immediately for all '
    'users belonging to that role.')

note_box(doc,
    'The Primary User (FPO owner) always has full access regardless of the matrix — '
    'the matrix applies to Secondary Users and any additional roles.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — APPLICATIONS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '9. Applications')
body(doc,
    'The Applications module is used to review, reject, request additional information, '
    'and manage all FPO registration applications submitted through the platform.')

h2(doc, '9.1  How to Open')
numbered(doc, 1, 'Click Applications in the left navigation sidebar.')

h2(doc, '9.2  Application List')
body(doc,
    'The table lists all FPO applications. Use the filters to narrow by status, district, '
    'tier, or registration date.')
make_table(doc,
    headers=['Column', 'Description'],
    rows=[
        ('FPO Name',         'Name of the Farmer Producer Organisation'),
        ('Registration No.', 'Legal registration number (or CIN for Producer Companies)'),
        ('District',         'District where the FPO is registered'),
        ('Status',           'Current application status: Draft, Submitted, Approved, Rejected, Info Required, Suspended'),
        ('Tier',             'Current tier (Tier 1–4) based on annual assessment'),
        ('Submitted Date',   'Date the application was submitted'),
    ],
    col_widths=[1.8, 4.47],
)

h2(doc, '9.3  Application Statuses')
make_table(doc,
    headers=['Status', 'Meaning'],
    rows=[
        ('Draft',          'FPO is filling the registration wizard — not yet submitted'),
        ('Submitted',      'FPO has submitted — auto-approval in progress'),
        ('Approved',       'Application accepted — FPO is active on the platform'),
        ('Info Required',  'Admin has requested additional information from the FPO'),
        ('Rejected',       'Application rejected by admin with reason provided'),
        ('Suspended',      'Approved FPO account suspended by admin'),
        ('Claimed',        'An ownership claim has been filed against this FPO'),
    ],
    col_widths=[1.8, 4.47],
)

h2(doc, '9.4  Available Actions')
bullet(doc, 'Click ⋯ → View Detail to open the full application with all sections and uploaded documents.')
bullet(doc, 'Click ⋯ → Reject to reject an Approved application (reason required — minimum 20 characters).')
bullet(doc, 'Click ⋯ → Request Info to send an information request to the FPO.')
bullet(doc, 'Click ⋯ → Verify Document to mark a specific uploaded document as verified.')
bullet(doc, 'Click ⋯ → Assign Tier to manually assign a tier after reviewing the assessment.')

note_box(doc,
    'FPO registration is auto-approved on submission — no manual approve action is required. '
    'Admin intervention is only needed to Reject or Suspend an application after approval.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — EXTERNAL APIs
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '10. External APIs')
body(doc,
    'The External APIs module stores the API credentials used for live verification of '
    'PAN numbers, GSTIN numbers, and CIN (Company Identification Numbers) during FPO '
    'registration. These credentials are managed exclusively by super administrators.')

h2(doc, '10.1  How to Open')
numbered(doc, 1, 'Click External APIs in the left navigation sidebar.')

h2(doc, '10.2  API Configuration')
make_table(doc,
    headers=['API', 'Used For', 'Status When Not Configured'],
    rows=[
        ('PAN Verification',   'Validate FPO PAN card numbers against income tax records',    'Falls back to format-only validation'),
        ('GSTIN Verification', 'Validate GST registration numbers against GST portal',        'Falls back to format-only validation'),
        ('CIN Verification',   'Validate Company Identification Numbers against MCA records', 'Falls back to format-only validation'),
    ],
    col_widths=[1.8, 2.5, 1.97],
)

numbered(doc, 1, 'Click + Add External API to configure a new API.')
numbered(doc, 2, 'Enter the API name, base URL, and authentication credentials provided by the API vendor.')
numbered(doc, 3, 'Click Activate to enable live verification for that API.')

note_box(doc,
    'Until credentials are configured and activated, the platform falls back to '
    'format-only validation (e.g. checks that the PAN matches the correct pattern). '
    'This is safe for testing but should be replaced with live verification before go-live.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — SITE CONTENT
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '11. Site Content')
body(doc,
    'The Site Content module allows administrators to edit all public-facing content on '
    'the KAU-FPO platform landing page and other public pages — without any code changes.')

h2(doc, '11.1  How to Open')
numbered(doc, 1, 'Click Site Content in the left navigation sidebar.')

h2(doc, '11.2  Content Categories')
make_table(doc,
    headers=['Category', 'What It Controls'],
    rows=[
        ('Content Blocks', 'Hero headline, hero subheading, hero description, About section, Mission, Vision, How to Register text'),
        ('Documents',      'Downloadable documents shown in the Document Library on the public portal'),
        ('Gallery',        'Photo albums and images shown in the Gallery section'),
        ('Our Team',       'Leadership team cards — name, designation, photo, order'),
        ('Quick Links',    'External portal links shown in the Quick Links section'),
        ('Partners',       'Partner organisation logos shown in the Partners strip'),
        ('News Sources',   'External news source links shown in the News section'),
        ('Feedback',       'View messages submitted via the Contact Us / Feedback form'),
    ],
    col_widths=[1.8, 4.47],
)

h2(doc, '11.3  Editing a Content Block')
numbered(doc, 1, 'Select the Content Blocks category.')
numbered(doc, 2, 'Find the block to edit (e.g. hero_headline, about_title).')
numbered(doc, 3, 'Click ⋯ → Edit on the block row.')
numbered(doc, 4, 'Update the text for English and/or Malayalam.')
numbered(doc, 5, 'Click Save. The change is reflected on the public site immediately.')

h2(doc, '11.4  Managing the Gallery')
bullet(doc, 'Click the Gallery category.')
bullet(doc, 'Click + Add Album to create a new photo album.')
bullet(doc, 'Click on an album to open it and upload photos.')
bullet(doc, 'Use ⋯ → Delete to remove a photo or album.')

h2(doc, '11.5  Managing Our Team')
bullet(doc, 'Click the Our Team category.')
bullet(doc, 'Click + Add Team Member to add a new leadership entry.')
bullet(doc, 'Fill in Name, Designation, Photo (upload), and Order (display order on page).')
bullet(doc, 'Use ⋯ → Edit to update or ⋯ → Deactivate to hide a member without deleting.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — ANNOUNCEMENTS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '12. Announcements')
body(doc,
    'The Announcements module manages all news and announcements published on the public '
    'News & Announcements page of the platform.')

h2(doc, '12.1  How to Open')
numbered(doc, 1, 'Click Announcements in the left navigation sidebar.')

h2(doc, '12.2  Adding a New Announcement')
numbered(doc, 1, 'Click + Add Announcement.')
numbered(doc, 2, 'Fill in the Title, Category (News or Announcement), Body, and Published Date.')
numbered(doc, 3, 'Set the Display Order (lower numbers appear first).')
numbered(doc, 4, 'Set Status to Active to publish immediately, or leave Inactive to save as draft.')
numbered(doc, 5, 'Click Save.')

h2(doc, '12.3  Announcement Categories')
make_table(doc,
    headers=['Category', 'Use Case'],
    rows=[
        ('Announcement', 'Official circulars, scheme notifications, KAU communications'),
        ('News',         'News articles, event coverage, agricultural updates'),
    ],
    col_widths=[1.8, 4.47],
)

h2(doc, '12.4  Managing Existing Announcements')
bullet(doc, 'Use Search to find an announcement by title.')
bullet(doc, 'Use the Category filter to see only News or only Announcements.')
bullet(doc, 'Click ⋯ → Edit to update content, category, date, or order.')
bullet(doc, 'Click ⋯ → Deactivate to unpublish without deleting.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13 — FAQs
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '13. FAQs')
body(doc,
    'The FAQs module manages the Frequently Asked Questions shown on the public FAQ page. '
    'Each entry has a question, answer, category, display order, and status.')

h2(doc, '13.1  How to Open')
numbered(doc, 1, 'Click FAQs in the left navigation sidebar.')

h2(doc, '13.2  Adding a New FAQ')
numbered(doc, 1, 'Click + Add FAQ.')
numbered(doc, 2, 'Enter the Question and Answer.')
numbered(doc, 3, 'Select the Category: FPO General, Schemes, or Platform Usage.')
numbered(doc, 4, 'Set the Display Order (lower numbers appear first in the list).')
numbered(doc, 5, 'Set Status to Active to publish, or Inactive to save without publishing.')
numbered(doc, 6, 'Click Save.')

h2(doc, '13.3  FAQ Categories')
make_table(doc,
    headers=['Category', 'Topics'],
    rows=[
        ('FPO General',    'What is an FPO, membership, share capital, Board of Directors, CEO role'),
        ('Schemes',        'Government schemes, subsidy eligibility, application process'),
        ('Platform Usage', 'How to register, wizard help, expert directory, DPR module'),
    ],
    col_widths=[1.8, 4.47],
)

h2(doc, '13.4  Best Practices')
bullet(doc, 'Write questions exactly as a user would ask them — natural language, not technical.')
bullet(doc, 'Keep answers concise — 2 to 4 sentences where possible.')
bullet(doc, 'Review Display Order periodically to ensure the most common questions appear first.')
bullet(doc, 'Avoid duplicate questions — use Search before adding a new entry.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 14 — OWNERSHIP CLAIMS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '14. Ownership Claims')
body(doc,
    'The Ownership Claims module is used to review and process requests from users who '
    'believe they are the rightful owner of an FPO that is already registered on the '
    'platform under a different account.')

screenshot(doc, 'screenshot_00.png', 'Figure 6 — Ownership Claims list')

h2(doc, '14.1  How to Open')
numbered(doc, 1, 'Click Ownership Claims in the left navigation sidebar.')

h2(doc, '14.2  Claim Statuses')
make_table(doc,
    headers=['Status', 'Meaning'],
    rows=[
        ('Pending',  'Claim has been submitted and is awaiting admin review'),
        ('Approved', 'Admin has verified the claim and transferred ownership'),
        ('Rejected', 'Admin has reviewed and rejected the claim'),
    ],
    col_widths=[1.5, 4.77],
)

h2(doc, '14.3  Processing a Claim')
numbered(doc, 1, 'Click on a claim row to open the full claim detail.')
numbered(doc, 2,
    'Review the claimant\'s information: name, email, reason for claim, and the FPO '
    'being claimed (identified by CIN, GST, or PAN).')
numbered(doc, 3, 'Review any supporting documents uploaded by the claimant.')
numbered(doc, 4,
    'If more information is needed, use Request More Info to send a message to the claimant.')
numbered(doc, 5, 'Click Approve to transfer ownership of the FPO to the claimant\'s account.')
numbered(doc, 6, 'Click Reject with a written reason to deny the claim.')

note_box(doc,
    'Approving a claim transfers the FPO Primary User role to the claimant. The previous '
    'primary user\'s account is demoted. This action cannot be undone — review all '
    'supporting documents carefully before approving.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 15 — AUDIT LOGS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '15. Audit Logs')
body(doc,
    'The Audit Logs module provides a complete, read-only trail of all system events '
    'and administrative actions. Use it to investigate security incidents, trace data '
    'changes, and monitor platform activity.')

screenshot(doc, 'screenshot_01.png', 'Figure 7 — Audit Logs list')

h2(doc, '15.1  How to Open')
numbered(doc, 1, 'Click Audit Logs in the left navigation sidebar.')

h2(doc, '15.2  Table Columns')
make_table(doc,
    headers=['Column', 'Description'],
    rows=[
        ('Time',         'Date and time of the event (displayed in local timezone)'),
        ('Action',       'Type of event — Login, Logout, Failed Login, Document Upload, FPO Submit, etc.'),
        ('Performed By', 'The administrator or user who performed the action'),
        ('Object',       'The record affected (FPO name, document type, etc.) — empty for auth events'),
        ('Method',       'HTTP method used (POST, PATCH, DELETE) — useful for technical investigation'),
        ('IP Address',   'IP address from which the action was performed'),
    ],
    col_widths=[1.5, 4.77],
)

h2(doc, '15.3  Filtering Logs')
bullet(doc, 'Use the Search bar to find events by user name or FPO name.')
bullet(doc, 'Use the Action dropdown to filter by event type (e.g. show only Failed Login events).')
bullet(doc, 'Use From Date and To Date to narrow logs to a specific time period.')
bullet(doc, 'Use Columns to hide Method or Object if not needed for the current investigation.')

h2(doc, '15.4  Common Audit Actions')
make_table(doc,
    headers=['Action Type', 'When It Appears'],
    rows=[
        ('Login',                  'Successful admin or FPO user login'),
        ('Logout',                 'Explicit logout action'),
        ('Failed Login',           'Incorrect password — indicates possible unauthorized access attempts'),
        ('Document Upload',        'FPO uploaded a registration document'),
        ('Document Delete',        'FPO deleted a document'),
        ('FPO Submit',             'FPO submitted the registration application'),
        ('FPO Status Change',      'Admin changed the FPO status (Approve, Reject, Suspend)'),
        ('FPO Profile Change',     'FPO profile fields were edited'),
        ('FPO User Invite',        'A secondary user was invited to an FPO team'),
        ('FPO User Activate',      'A secondary user account was activated'),
        ('FPO User Deactivate',    'A secondary user account was deactivated'),
        ('Tier Recalculation',     'FPO tier was recalculated based on assessment answers'),
    ],
    col_widths=[2.2, 4.07],
)

note_box(doc,
    'Audit log data is read-only — entries cannot be edited or deleted. '
    'Monitor the IP Address column for logins from unexpected locations, which could '
    'indicate unauthorized access. Report anomalies to the KAU IT team immediately.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 16 — EXPERT DIRECTORY
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '16. Expert Directory')
body(doc,
    'The Expert Directory module manages the list of agricultural experts and KAU '
    'specialists that FPOs can contact through the platform.')

screenshot(doc, 'screenshot_02.png', 'Figure 8 — Expert Directory')

h2(doc, '16.1  How to Open')
numbered(doc, 1, 'Click Experts in the left navigation sidebar.')

h2(doc, '16.2  Expert Table Columns')
make_table(doc,
    headers=['Column', 'Description'],
    rows=[
        ('Name',         'Expert\'s full name and designation shown below'),
        ('Organisation', 'Institution or department the expert belongs to'),
        ('Category',     'Expert category — Scientist / Researcher, Extension Officer, etc.'),
        ('District',     'District where the expert is primarily available (3-letter code)'),
        ('Status',       'Active — visible to FPOs; Inactive — hidden from FPO portal'),
    ],
    col_widths=[1.5, 4.77],
)

h2(doc, '16.3  Adding a New Expert')
numbered(doc, 1, 'Click + Add Expert.')
numbered(doc, 2, 'Fill in Name, Designation, Organisation, Category, District, Contact Email, and Phone.')
numbered(doc, 3, 'Set Status to Active to make the expert visible to FPOs immediately.')
numbered(doc, 4, 'Click Save.')

h2(doc, '16.4  Finding and Filtering Experts')
bullet(doc, 'Use the Search bar to find experts by name or organisation.')
bullet(doc, 'Use the Category dropdown to filter by expert type.')
bullet(doc, 'Use the District dropdown to find experts available in a specific district.')
bullet(doc, 'Click Name or Category column headers to sort the table.')

h2(doc, '16.5  Managing Expert Status')
bullet(doc, 'Click ⋯ → Deactivate to remove an expert from the FPO-facing directory without deleting.')
bullet(doc, 'Click ⋯ → Activate to restore visibility.')
bullet(doc, 'Regularly review the Status column to ensure only currently available experts are Active.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 17 — SCHEMES & SUBSIDIES
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '17. Schemes & Subsidies')
body(doc,
    'The Schemes & Subsidies module manages the catalog of government schemes and '
    'subsidies available to FPOs. Schemes added here appear in the Schemes & Subsidies '
    'section of the FPO portal.')

screenshot(doc, 'screenshot_03.png', 'Figure 9 — Schemes & Subsidies list')

h2(doc, '17.1  How to Open')
numbered(doc, 1, 'Click Schemes & Subsidies in the left navigation sidebar.')

h2(doc, '17.2  Scheme Table Columns')
make_table(doc,
    headers=['Column', 'Description'],
    rows=[
        ('Scheme Name',       'Full name of the government scheme or subsidy programme'),
        ('Administering Body','The government department or institution managing the scheme'),
        ('Category',          'Color-coded category badge — Credit & Finance, Capacity Building, Infrastructure, etc.'),
        ('Status',            'Active — visible to FPOs; Inactive — hidden from portal'),
    ],
    col_widths=[1.8, 4.47],
)

h2(doc, '17.3  Adding a New Scheme')
numbered(doc, 1, 'Click + Add Scheme.')
numbered(doc, 2, 'Fill in Scheme Name, Administering Body, Category, Description, and Eligibility Criteria.')
numbered(doc, 3, 'Add any application links or important dates relevant to the scheme.')
numbered(doc, 4, 'Set Status to Active to publish immediately.')
numbered(doc, 5, 'Click Save.')

h2(doc, '17.4  Scheme Categories')
make_table(doc,
    headers=['Category', 'Examples'],
    rows=[
        ('Credit & Finance',    'AIF, Credit Guarantee Fund, Equity Grant Fund'),
        ('Capacity Building',   'Formation and promotion of 10,000 FPOs scheme'),
        ('Infrastructure',      'Mission for Integrated Development of Horticulture (MIDH)'),
        ('Market Linkage',      'e-NAM, ONDC FPO onboarding schemes'),
        ('Technology',          'Digital agriculture schemes, precision farming support'),
    ],
    col_widths=[1.8, 4.47],
)

note_box(doc,
    'Keep the Administering Body field accurate — FPO managers use it to identify '
    'which government department to approach for applications. Review scheme status '
    'quarterly and deactivate schemes whose application windows have closed.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 18 — SETTINGS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '18. Settings')
body(doc,
    'The Settings module allows each administrator to manage their own account profile '
    'and security preferences. Settings are personal — changes here affect only the '
    'logged-in administrator\'s account.')

screenshot(doc, 'screenshot_04.png', 'Figure 10 — Settings > Profile page')

h2(doc, '18.1  How to Open')
numbered(doc, 1, 'Click Settings in the bottom section of the left navigation sidebar.')
numbered(doc, 2, 'The page opens with the Profile tab selected by default.')

h2(doc, '18.2  Profile Tab')
make_table(doc,
    headers=['Field', 'Description'],
    rows=[
        ('Avatar',             'Initials badge generated from your name — not editable'),
        ('First Name',         'Your first name as shown in the portal and notifications'),
        ('Last Name',          'Your last name'),
        ('Phone',              'Your mobile number — used for SMS notifications and account recovery. Keep this current.'),
        ('Preferred Language', 'Language for email and SMS notifications sent to your account (English / Malayalam)'),
    ],
    col_widths=[2.0, 4.27],
)

h3(doc, 'To edit profile details:')
numbered(doc, 1, 'Click the Edit button (top right of the Profile panel).')
numbered(doc, 2, 'Update the relevant fields.')
numbered(doc, 3, 'Click Save to apply changes.')

h2(doc, '18.3  Security Tab')
bullet(doc, 'Click the Security tab in the left panel of the Settings page.')
bullet(doc, 'Use Change Password to update your current password.')
bullet(doc, 'Use Two-Factor Authentication (2FA) to set up, view status, or disable TOTP-based 2FA.')
bullet(doc, '2FA is mandatory for super administrators and sub-administrators.')

note_box(doc,
    'Keep your phone number up to date in Settings > Profile. An outdated number will '
    'cause SMS OTPs and account recovery messages to be delivered to the wrong number.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 19 — TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '19. Troubleshooting')

make_table(doc,
    headers=['Problem', 'Likely Cause', 'Solution'],
    rows=[
        ('Cannot log in — "Invalid credentials"',
         'Incorrect email or password',
         'Check caps lock. Try "Forgot Password" to reset. After 5 failures the account locks for 30 minutes.'),
        ('Account locked',
         '5 consecutive failed login attempts',
         'Wait 30 minutes and try again, or contact the super admin to check the Audit Logs.'),
        ('2FA code rejected',
         'Device clock out of sync or wrong authenticator app',
         'Ensure your phone clock is set to automatic time. Try a backup code. Contact super admin if locked out.'),
        ('Page not loading',
         'Browser cache or connectivity issue',
         'Refresh the page (Ctrl+R). Clear browser cache. Check internet connection.'),
        ('Table data not updating',
         'Browser showing cached version',
         'Click the Refresh (↺) icon on the page toolbar to reload data from the server.'),
        ('Cannot save changes',
         'Required field left empty or validation error',
         'Check for red highlighted fields on the form. Fill in all required fields before saving.'),
        ('Notification not delivered to FPO',
         'Channel not configured or inactive',
         'Check Notifications → Channel Settings. Ensure the relevant channel (email/SMS) is Active with valid credentials.'),
        ('Export / report not downloading',
         'Browser blocking file downloads',
         'Allow file downloads for fpolinkage.kau.in in browser settings. Try a different browser.'),
        ('FPO application status not updating',
         'Delayed page refresh',
         'Click the Refresh icon on the Applications page or reload the browser tab.'),
        ('Sub-admin cannot access a module',
         'Insufficient permissions assigned',
         'Go to Sub-Admins → select the user → Permissions and assign the required permission codenames.'),
    ],
    col_widths=[2.0, 2.0, 2.27],
)

note_box(doc,
    'For technical issues not covered above, contact the KAU Communication Centre at '
    'fpolinkage.kau.in. Include the page name, the action you were performing, and any '
    'error message displayed on screen.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# CLOSING
# ══════════════════════════════════════════════════════════════════════════════

spacer(doc, 20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_bottom_border(p, sz='4')

spacer(doc, 8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Kerala Agricultural University  |  Communication Centre, Mannuthy, Thrissur — 680651  |  '
    'AI-Based Digital Platform for KAU-FPO Linkage Programme  |  fpolinkage.kau.in'
)
r.font.size      = Pt(9)
r.font.name      = 'Cambria'
r.font.color.rgb = GREY


# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
