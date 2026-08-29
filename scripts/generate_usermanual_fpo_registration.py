"""
User Manual Generator — FPO Registration
KAU-FPO Platform
Matches RCD house style exactly (KefiTech).
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paths ────────────────────────────────────────────────────────────────────
KAU_LOGO    = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/KAU Emblem 0.5x0.75.jpg'
LOGO_SMALL  = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/Documents/kefitech_logo_small.png'
SCREENSHOTS = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/Documents/Usermanual/screenshots'
OUTPUT      = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/Documents/Usermanual/KAU_FPO_UserManual_FPO_Registration_v1.0.docx'

# ─── Colors ───────────────────────────────────────────────────────────────────
NAVY_HEX  = '1F3864'
DARK_NAVY = RGBColor(0x1F, 0x38, 0x64)
ORANGE    = RGBColor(0xE8, 0x6C, 0x1A)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREY      = RGBColor(0x60, 0x60, 0x60)
BLACK     = RGBColor(0x00, 0x00, 0x00)
LIGHT_GREY_HEX = 'F2F2F2'


# ─── XML Helpers ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, start=120, bottom=60, end=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'),    str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def remove_table_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'none')
        b.set(qn('w:sz'),    '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_bottom_border(paragraph, color=NAVY_HEX, sz='8'):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_page_number(run):
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)


def set_col_width(cell, width_inches):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def set_table_width(table, width_cm):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    twips = int(width_cm * 567)
    tblW.set(qn('w:w'), str(twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


# ─── Header / Footer ──────────────────────────────────────────────────────────

def build_header(header_obj, title='User Manual — FPO Registration'):
    for p in header_obj.paragraphs:
        p._element.getparent().remove(p._element)

    tbl = header_obj.add_table(rows=1, cols=3, width=Inches(6.5))
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)

    c0 = tbl.cell(0, 0)
    set_col_width(c0, 1.15)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(0)
    r0 = p0.add_run()
    r0.add_picture(KAU_LOGO, width=Cm(0.8), height=Cm(0.49))

    c1 = tbl.cell(0, 1)
    set_col_width(c1, 3.82)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after  = Pt(0)
    r1 = p1.add_run(title)
    r1.font.bold      = True
    r1.font.size      = Pt(9)
    r1.font.color.rgb = DARK_NAVY

    c2 = tbl.cell(0, 2)
    set_col_width(c2, 1.30)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run('KAU-FPO Platform')
    r2.font.size      = Pt(9)
    r2.font.color.rgb = DARK_NAVY

    div = header_obj.add_paragraph()
    div.paragraph_format.space_before = Pt(2)
    div.paragraph_format.space_after  = Pt(0)
    add_bottom_border(div)


def build_first_page_header(header_obj):
    for p in header_obj.paragraphs:
        p._element.getparent().remove(p._element)
    header_obj.add_paragraph()


def build_footer(footer_obj):
    for p in footer_obj.paragraphs:
        p._element.getparent().remove(p._element)

    tbl = footer_obj.add_table(rows=1, cols=2, width=Inches(6.5))
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)

    cl = tbl.cell(0, 0)
    set_col_width(cl, 4.5)
    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rl = pl.add_run('Kerala Agricultural University  │  fpolinkage.kau.in')
    rl.font.size      = Pt(8)
    rl.font.color.rgb = GREY

    cr = tbl.cell(0, 1)
    set_col_width(cr, 1.77)
    pr = cr.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = pr.add_run('Page ')
    rr.font.size      = Pt(8)
    rr.font.color.rgb = GREY
    rn = pr.add_run()
    rn.font.size      = Pt(8)
    rn.font.color.rgb = GREY
    add_page_number(rn)


# ─── Content Helpers ──────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold      = True
    r.font.size      = Pt(16)
    r.font.name      = 'Calibri'
    r.font.color.rgb = DARK_NAVY
    add_bottom_border(p)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold      = True
    r.font.size      = Pt(12)
    r.font.name      = 'Calibri'
    r.font.color.rgb = ORANGE
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold      = True
    r.font.size      = Pt(11)
    r.font.name      = 'Calibri'
    r.font.color.rgb = BLACK
    return p


def body(doc, text, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Cambria'
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Cambria'
    return p


def numbered(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Pt(18)
    r1 = p.add_run(f'{number}.  ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Cambria'
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.name = 'Cambria'
    return p


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    return p


def make_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.style = 'Table Grid'

    for j, h_text in enumerate(headers):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, NAVY_HEX)
        set_cell_margins(c)
        if col_widths:
            set_col_width(c, col_widths[j])
        p = c.paragraphs[0]
        r = p.add_run(h_text)
        r.font.bold      = True
        r.font.color.rgb = WHITE
        r.font.size      = Pt(10)
        r.font.name      = 'Calibri'

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            c = tbl.rows[i + 1].cells[j]
            set_cell_margins(c)
            if col_widths:
                set_col_width(c, col_widths[j])
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = 'Cambria'

    spacer(doc, 8)
    return tbl


def screenshot(doc, filename, caption=None):
    """Inserts a screenshot centered with an optional caption."""
    import os
    path = os.path.join(SCREENSHOTS, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run()
    r.add_picture(path, width=Cm(14))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        cr = cp.add_run(caption)
        cr.font.italic    = True
        cr.font.size      = Pt(9)
        cr.font.name      = 'Cambria'
        cr.font.color.rgb = GREY


def note_box(doc, text):
    """Renders a light grey tip/note box."""
    spacer(doc, 4)
    tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)
    c = tbl.cell(0, 0)
    set_cell_bg(c, LIGHT_GREY_HEX)
    set_cell_margins(c, top=100, start=160, bottom=100, end=160)
    p = c.paragraphs[0]
    r1 = p.add_run('Note: ')
    r1.font.bold      = True
    r1.font.size      = Pt(10)
    r1.font.name      = 'Calibri'
    r1.font.color.rgb = DARK_NAVY
    r2 = p.add_run(text)
    r2.font.size      = Pt(10)
    r2.font.name      = 'Cambria'
    spacer(doc, 6)


def step_banner(doc, step_num, step_title, total=7):
    """Orange-accented step header."""
    tbl = doc.add_table(rows=1, cols=2)
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)

    c0 = tbl.cell(0, 0)
    set_cell_bg(c0, 'E86C1A')
    set_col_width(c0, 1.3)
    set_cell_margins(c0, top=100, start=140, bottom=100, end=140)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(f'STEP {step_num}/{total}')
    r0.font.bold      = True
    r0.font.size      = Pt(11)
    r0.font.name      = 'Calibri'
    r0.font.color.rgb = WHITE

    c1 = tbl.cell(0, 1)
    set_cell_bg(c1, NAVY_HEX)
    set_col_width(c1, 5.47)
    set_cell_margins(c1, top=100, start=160, bottom=100, end=140)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(step_title)
    r1.font.bold      = True
    r1.font.size      = Pt(11)
    r1.font.name      = 'Calibri'
    r1.font.color.rgb = WHITE

    spacer(doc, 6)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Cambria'
style.font.size = Pt(11)

sec = doc.sections[0]
sec.page_width        = Cm(21)
sec.page_height       = Cm(29.7)
sec.left_margin       = Cm(2.54)
sec.right_margin      = Cm(2.54)
sec.top_margin        = Cm(2.0)
sec.bottom_margin     = Cm(2.0)
sec.header_distance   = Cm(1.0)
sec.footer_distance   = Cm(1.0)
sec.different_first_page_header_footer = True

build_header(sec.header)
build_footer(sec.footer)
build_first_page_header(sec.first_page_header)
build_footer(sec.first_page_footer)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(12)
p.add_run().add_picture(KAU_LOGO, width=Cm(9.14))

banner_tbl = doc.add_table(rows=1, cols=1)
banner_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(banner_tbl, 16.51)
bc = banner_tbl.cell(0, 0)
set_cell_bg(bc, NAVY_HEX)
set_cell_margins(bc, top=160, start=200, bottom=160, end=200)
bp = bc.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
br = bp.add_run('USER MANUAL — FPO REGISTRATION')
br.font.bold      = True
br.font.size      = Pt(20)
br.font.name      = 'Calibri'
br.font.color.rgb = WHITE

spacer(doc, 14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('KERALA AGRICULTURAL UNIVERSITY')
r.font.bold      = True
r.font.size      = Pt(16)
r.font.name      = 'Calibri'
r.font.color.rgb = DARK_NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run('Communication Centre, Mannuthy, Thrissur — 680651')
r.font.size      = Pt(11)
r.font.name      = 'Cambria'
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('AI-Based Digital Platform for KAU-FPO Linkage Programme')
r.font.bold      = True
r.font.size      = Pt(13)
r.font.name      = 'Calibri'
r.font.color.rgb = ORANGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run('User Manual — FPO Registration & Application')
r.font.bold      = True
r.font.size      = Pt(13)
r.font.name      = 'Calibri'
r.font.color.rgb = ORANGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Version 1.0  |  Kerala Agricultural University')
r.font.size      = Pt(11)
r.font.name      = 'Cambria'
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run('Role: FPO Primary User  |  Portal: fpolinkage.kau.in')
r.font.size      = Pt(11)
r.font.name      = 'Cambria'
r.font.color.rgb = GREY

spacer(doc, 14)

# Overview table on cover
make_table(doc,
    ['Section', 'Pages Covered'],
    [
        ('Registration',   'Eligibility Check → Verify Phone → Create Account'),
        ('Application',    'Basic Info → Contact → Signatory → Business & Bank → Verify Contacts → Documents → Review & Submit'),
        ('Post-Submission', 'Application Status & Activity Timeline'),
        ('FPO Portal',     'Expert Directory'),
    ],
    col_widths=[1.8, 4.5],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT INFORMATION
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, 'Document Information')

make_table(doc,
    ['Item', 'Details'],
    [
        ('Application Name',    'KAU-FPO Platform — fpolinkage.kau.in'),
        ('Version',             '1.0'),
        ('Prepared By',         'Kerala Agricultural University'),
        ('Date',                'August 2026'),
        ('Audience',            'FPO Primary Users — new registrations'),
        ('Document Version',    'v1.0 — Initial Release'),
    ],
    col_widths=[2.2, 4.1],
)

h2(doc, 'Purpose')
body(doc,
    'This manual guides an FPO (Farmer Producer Organization) primary user through the complete '
    'registration and application process on the KAU-FPO Platform. Using this platform, FPOs can '
    'register their organization, complete a structured application, upload supporting documents, '
    'and receive instant approval — all online without visiting any office.')

h2(doc, 'System Requirements')

h3(doc, 'Supported Devices')
make_table(doc,
    ['Device', 'Supported'],
    [
        ('Desktop / Laptop', '✔ Recommended — best experience for filling the 7-step application form'),
        ('Android Phone / Tablet', '✔ Supported'),
        ('iPhone / iPad', '✔ Supported'),
    ],
    col_widths=[2.2, 4.1],
)

h3(doc, 'Supported Browsers')
make_table(doc,
    ['Browser', 'Supported'],
    [
        ('Google Chrome',    '✔ Recommended'),
        ('Microsoft Edge',   '✔ Supported'),
        ('Safari',           '✔ Supported (iPhone/iPad)'),
        ('Mozilla Firefox',  '✔ Supported'),
        ('Internet Explorer','✘ Not supported — please use Chrome or Edge'),
    ],
    col_widths=[2.2, 4.1],
)

note_box(doc,
    'A stable internet connection is required throughout the registration and application process. '
    'Keep your FPO registration documents, PAN card, and bank details ready before starting — '
    'the process is easier to complete in one session.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '1. Overview')

body(doc,
    'This manual guides a Farmer Producer Organization (FPO) primary user through the complete '
    'registration and application process on the KAU-FPO Platform (fpolinkage.kau.in). '
    'The process is divided into two stages: Registration (3 steps) and Application (7 steps).')

spacer(doc, 6)

make_table(doc,
    ['Stage', 'Steps', 'What Happens'],
    [
        ('Registration', '3 steps', 'Eligibility check → Phone verification → Account creation'),
        ('Application',  '7 steps', 'Full FPO profile → Document upload → Submit for approval'),
        ('Post-Submission', '—', 'Auto-approval → Application ID issued → Dashboard access'),
    ],
    col_widths=[1.5, 1.0, 4.0],
)

note_box(doc,
    'The platform auto-approves the application once all requirements are met. There is no manual '
    'waiting period — once submitted successfully, the FPO receives an Application ID and is '
    'granted immediate access to the dashboard.')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — REGISTRATION
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
h1(doc, '2. Registration')

body(doc,
    'Registration creates your personal login account and confirms that your FPO meets the '
    'minimum eligibility criteria. It consists of 3 steps.')

# ── Step 1 ────────────────────────────────────────────────────────────────────
step_banner(doc, 1, 'Eligibility Check', total=3)

h2(doc, '2.1 Purpose')
body(doc,
    'The Eligibility Check confirms that your FPO meets the minimum requirements before allowing '
    'it to proceed with full registration. This prevents ineligible organizations from completing '
    'an application.')

h2(doc, '2.2 How to Access')
bullet(doc, 'Go to fpolinkage.kau.in and click Register or Apply Now on the home page.')
bullet(doc, 'The Eligibility Check page loads as Step 1 of 3.')
bullet(doc, 'If you already have an account, click Sign in instead.')

h2(doc, '2.3 Fields on This Page')
make_table(doc,
    ['Field', 'Required', 'Description'],
    [
        ('District',            'Yes', 'Select the district where your FPO is registered.'),
        ('Total Farmer Members','Yes', 'Enter the total number of farmer members. Minimum: 10.'),
        ('Requirements Checklist', 'Yes',
         'Tick all three boxes to confirm: (1) Registered under an applicable Act, '
         '(2) Holds a valid registration certificate, (3) Has an active bank account in the FPO\'s name.'),
    ],
    col_widths=[2.2, 0.8, 3.3],
)

h2(doc, '2.4 Steps to Complete')
numbered(doc, 1, 'Select your District from the dropdown.')
numbered(doc, 2, 'Enter the Total Farmer Members count (minimum 10).')
numbered(doc, 3, 'Review the three Requirements listed in the checklist box.')
numbered(doc, 4, 'Tick each checkbox individually, or click Accept all to check all three at once.')
numbered(doc, 5, 'Click Check Eligibility. If eligible, you advance to Step 2: Verify Phone.')

screenshot(doc, 'screenshot_00.png', 'Figure 1 — Eligibility Check page')

note_box(doc,
    'Use Accept all only when genuinely confident all three requirements are met — '
    'these are compliance declarations. Double-check your District selection as it affects '
    'schemes and regional rules applied to the FPO later.')

# ── Step 2 ────────────────────────────────────────────────────────────────────
step_banner(doc, 2, 'Verify Phone Number', total=3)

h2(doc, '2.5 Purpose')
body(doc,
    'Step 2 confirms that you have access to the mobile number provided, using a one-time '
    'password (OTP) sent via SMS. This ties your registration to a valid, reachable phone number.')

h2(doc, '2.6 Steps to Complete')
numbered(doc, 1, 'Confirm the phone number displayed on screen is correct.')
numbered(doc, 2, 'Wait for the OTP to arrive via SMS.')
numbered(doc, 3, 'If not received within a minute, click Resend to get a new code.')
numbered(doc, 4, 'Enter the 6-digit OTP in the Enter OTP field.')
numbered(doc, 5, 'Click Verify & Continue to validate and proceed to Step 3.')
numbered(doc, 6, 'Click ← Back if you need to correct any details from Step 1.')

note_box(doc,
    'OTPs are time-limited. If verification fails after a delay, use Resend to get a fresh code. '
    'Double-check the phone number before requesting the OTP — an incorrect number will prevent the SMS from arriving.')

# ── Step 3 ────────────────────────────────────────────────────────────────────
step_banner(doc, 3, 'Create Your Account', total=3)

h2(doc, '2.7 Purpose')
body(doc,
    'Step 3 collects your personal details to create the login account that will be used to manage '
    'the FPO\'s profile going forward.')

h2(doc, '2.8 Fields on This Page')
make_table(doc,
    ['Field', 'Required', 'Description'],
    [
        ('First Name',        'Yes', 'Your first name.'),
        ('Last Name',         'Yes', 'Your last name.'),
        ('Email Address',     'Yes', 'A valid email address — used for login and notifications.'),
        ('Phone Number',      '—',   'Pre-filled with the number verified in Step 2. No action needed.'),
        ('Password',          'Yes', 'Minimum 8 characters. Must include uppercase, lowercase, digit, and special character.'),
        ('Confirm Password',  'Yes', 'Re-enter the same password to confirm.'),
    ],
    col_widths=[1.8, 0.8, 3.7],
)

h2(doc, '2.9 Steps to Complete')
numbered(doc, 1, 'Enter First Name and Last Name.')
numbered(doc, 2, 'Enter your Email Address.')
numbered(doc, 3, 'The Phone Number field shows "Verified in previous step" — no action needed.')
numbered(doc, 4, 'Create a Password that meets the strength requirements.')
numbered(doc, 5, 'Re-enter the same password in Confirm Password.')
numbered(doc, 6, 'Use the eye icon to show/hide password text and confirm accuracy.')
numbered(doc, 7, 'Click Create Account & Continue to finalize and proceed to the Application.')

screenshot(doc, 'screenshot_01.png', 'Figure 2 — Create Your Account page')

note_box(doc,
    'Use a strong password — this account manages sensitive FPO data. '
    'Double-check the Email Address for accuracy; it will be used for all platform notifications. '
    'Password and Confirm Password must match exactly.')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — APPLICATION
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
h1(doc, '3. FPO Application')

body(doc,
    'After account creation, you are directed into the 7-step FPO Application. This collects '
    'all information about the FPO — legal details, location, signatory, business activity, '
    'bank details, contact verification, and supporting documents. You can save progress at any '
    'step and return later.')

note_box(doc,
    'Use the Save button frequently if gathering details from physical documents takes multiple sessions. '
    'All steps must be completed before submitting.')

# ── Step 1 ────────────────────────────────────────────────────────────────────
step_banner(doc, 1, 'Basic Information', total=7)

h2(doc, '3.1 Purpose')
body(doc,
    'Collects the FPO\'s core legal registration details. This forms the identity foundation '
    'of the FPO\'s profile used throughout the rest of the application.')

h2(doc, '3.2 Fields on This Page')
make_table(doc,
    ['Field', 'Required', 'Description'],
    [
        ('FPO Name (English)',    'Yes', 'Official legal name exactly as on registration documents.'),
        ('FPO Name (Malayalam)',  'No',  'Local-language name for display purposes.'),
        ('Registered Under',      'Yes', 'Legal entity type — select from dropdown (Companies Act, Producer Companies, Cooperative Society, etc.).'),
        ('Registration Number',   'Yes', 'Official number issued by the registering authority.'),
        ('CIN Number',            'No',  'Company Identification Number — required for Companies Act / Producer Companies registrations only.'),
        ('Date of Registration',  'Yes', 'Use the date picker (dd-mm-yyyy).'),
        ('PAN Number',            'Yes', 'Permanent Account Number — required for all FPOs.'),
        ('GST Number',            'No',  'If the FPO is GST-registered.'),
    ],
    col_widths=[2.0, 0.8, 3.5],
)

h2(doc, '3.3 Steps to Complete')
numbered(doc, 1, 'Enter FPO Name (English) — must match official registration documents exactly.')
numbered(doc, 2, 'Enter FPO Name (Malayalam) — optional.')
numbered(doc, 3, 'Select Registered Under from the dropdown.')
numbered(doc, 4, 'Enter Registration Number.')
numbered(doc, 5, 'Enter CIN Number if applicable (companies/producer companies only).')
numbered(doc, 6, 'Select Date of Registration using the date picker.')
numbered(doc, 7, 'Enter PAN Number.')
numbered(doc, 8, 'Enter GST Number if applicable.')
numbered(doc, 9, 'Click Save to store progress, or Get Started → to proceed to Step 2.')

screenshot(doc, 'screenshot_02.png', 'Figure 3 — Basic Information (Step 1)')

# ── Step 2 ────────────────────────────────────────────────────────────────────
step_banner(doc, 2, 'Contact & Location', total=7)

h2(doc, '3.4 Purpose')
body(doc,
    'Collects the FPO\'s office address, contact information, and precise map location. '
    'This is the primary contact point administrators and platform notifications will use.')

h2(doc, '3.5 Fields on This Page')
make_table(doc,
    ['Field', 'Required', 'Description'],
    [
        ('District',       'Yes', 'Select the FPO\'s district first — enables the Block/Taluk dropdown.'),
        ('Block / Taluk',  'Yes', 'Select after choosing District.'),
        ('Village / Town', 'Yes', 'Specific village or town name.'),
        ('Address Line 1', 'Yes', 'House/building/street details.'),
        ('Address Line 2', 'No',  'Landmark or area reference.'),
        ('Pincode',        'Yes', '6-digit postal code.'),
        ('Office Phone',   'Yes', '10-digit contact number for the FPO office.'),
        ('Office Email',   'Yes', 'FPO\'s official email address — will be verified in Step 5.'),
        ('Website',        'No',  'FPO\'s website URL if available.'),
        ('Map Location',   'Yes', 'Drop a pin on the interactive map to record GPS coordinates.'),
    ],
    col_widths=[1.8, 0.8, 3.7],
)

h2(doc, '3.6 Setting the Map Location')
body(doc, 'Four ways to place the map pin:')
numbered(doc, 1, 'Search for the location using the FPO Location on Map search field.')
numbered(doc, 2, 'Click directly on the map to drop a pin.')
numbered(doc, 3, 'Drag an existing pin to adjust its position.')
numbered(doc, 4, 'Click Use My Location to auto-detect current GPS position (only if physically at the FPO office).')

screenshot(doc, 'screenshot_04.png', 'Figure 4 — Contact & Location (Step 2)')

note_box(doc,
    'Select District first — the Block/Taluk field remains disabled until this is done. '
    'Zoom in closely on the map before finalizing the pin to ensure accuracy in rural or less-mapped areas. '
    'Office Phone and Email entered here will be verified in Step 5.')

# ── Step 3 ────────────────────────────────────────────────────────────────────
step_banner(doc, 3, 'Signatory & Members', total=7)

h2(doc, '3.7 Purpose')
body(doc,
    'Collects details about the FPO\'s authorized signatory, membership statistics, and '
    'governance/agency information. This data is used to assess the FPO\'s structure and '
    'eligibility for government schemes.')

h2(doc, '3.8 Fields on This Page')
h3(doc, '3.8.1 Authorized Signatory')
make_table(doc,
    ['Field', 'Required', 'Description'],
    [
        ('Full Name',           'Yes', 'Legal name of the person authorized to act on behalf of the FPO.'),
        ('Designation',         'Yes', 'Their official role — select from dropdown (Chairman, Secretary, CEO, etc.).'),
        ('Phone',               'Yes', '10-digit mobile number.'),
        ('Email',               'Yes', 'Email address of the signatory.'),
        ('Aadhaar Last 4 Digits','Yes', 'Last 4 digits of the signatory\'s Aadhaar number for identity verification.'),
    ],
    col_widths=[2.0, 0.8, 3.5],
)

h3(doc, '3.8.2 Membership Details')
make_table(doc,
    ['Field', 'Required', 'Description'],
    [
        ('Total Members',    'Yes', 'Total farmer members — must be minimum 10.'),
        ('Male Members',     'No',  'Number of male members.'),
        ('Female Members',   'No',  'Number of female members.'),
        ('SC-ST Members',    'No',  'Number of SC/ST members.'),
    ],
    col_widths=[2.0, 0.8, 3.5],
)

h3(doc, '3.8.3 Governance & Agencies')
make_table(doc,
    ['Field', 'Required', 'Description'],
    [
        ('Promoting Agency',        'Yes', 'Select the agency that promoted / established this FPO.'),
        ('Facilitating Agency Name','No',  'Name of the facilitating organization if different from promoting agency.'),
        ('Total Directors',         'Yes', 'Total number of board directors.'),
        ('Women Directors',         'No',  'Number of women on the board.'),
        ('Directors Under 35',      'No',  'Number of directors aged below 35.'),
        ('CEO Available',           'No',  'Tick if the FPO has a dedicated CEO.'),
        ('Accountant Available',    'No',  'Tick if the FPO has a dedicated accountant.'),
    ],
    col_widths=[2.2, 0.8, 3.3],
)

screenshot(doc, 'screenshot_05.png', 'Figure 5 — Signatory & Members (Step 3)')

note_box(doc,
    'The Total Members figure should match the count entered during the Eligibility Check. '
    'Filling in the male/female/SC-ST breakdown helps determine eligibility for inclusion-focused schemes. '
    'The Aadhaar Last 4 Digits may be used for identity verification in Step 5.')

# ── Step 4 ────────────────────────────────────────────────────────────────────
step_banner(doc, 4, 'Business & Bank Details', total=7)

h2(doc, '3.9 Purpose')
body(doc,
    'Collects the FPO\'s commodity information, financial overview, and banking details. '
    'Bank details are used for scheme disbursements and subsidies.')

h2(doc, '3.10 Fields on This Page')
make_table(doc,
    ['Field', 'Required', 'Description'],
    [
        ('Primary Commodities',   'Yes', 'Main commodity/commodities the FPO deals in — affects scheme recommendations.'),
        ('Secondary Commodities', 'No',  'Additional commodities, if any.'),
        ('Annual Turnover',       'No',  'FPO\'s turnover in lakhs, if available.'),
        ('About the FPO',         'No',  'Short description of the FPO\'s activities and goals.'),
        ('Bank Name',             'Yes', 'Search and select the bank.'),
        ('Branch',                'Yes', 'Branch name.'),
        ('Account Number',        'Yes', 'FPO\'s official bank account number.'),
        ('IFSC Code',             'Yes', 'Branch IFSC code — determines routing of funds.'),
    ],
    col_widths=[2.0, 0.8, 3.5],
)

screenshot(doc, 'screenshot_06.png', 'Figure 6 — Business & Bank Details (Step 4)')

note_box(doc,
    'Enter bank details exactly as they appear on official bank documents (cancelled cheque or passbook). '
    'A typo in the IFSC Code can route funds to the wrong branch. '
    'Select Primary Commodities carefully — this influences which schemes and expert recommendations are shown later.')

# ── Step 5 ────────────────────────────────────────────────────────────────────
step_banner(doc, 5, 'Verify Contact Details', total=7)

h2(doc, '3.11 Purpose')
body(doc,
    'Confirms that the FPO\'s office email and phone number (entered in Step 2) are valid and '
    'accessible by sending OTPs to each. Both must be verified before proceeding.')

h2(doc, '3.12 Steps to Complete')
numbered(doc, 1, 'Review the Office Email and Office Phone displayed — confirm they match what was entered in Step 2.')
numbered(doc, 2, 'Click Send OTP to email. An OTP will be sent to the office email address.')
numbered(doc, 3, 'Enter the received code in the email OTP field within 10 minutes.')
numbered(doc, 4, 'Click Send OTP to phone. An OTP will be sent via SMS to the office phone.')
numbered(doc, 5, 'Enter the received code in the phone OTP field within 10 minutes.')
numbered(doc, 6, 'Once both are verified, click Continue → to proceed to Step 6: Documents.')

screenshot(doc, 'screenshot_07.png', 'Figure 7 — Verify Contact Details (Step 5)')

note_box(doc,
    'OTPs expire after 10 minutes — complete both verifications in the same session if possible. '
    'Check spam/junk folders if the email OTP is not received. '
    'If either contact detail is incorrect, click ← Back to return to Step 2 and correct it before attempting verification.')

# ── Step 6 ────────────────────────────────────────────────────────────────────
step_banner(doc, 6, 'Upload Documents', total=7)

h2(doc, '3.13 Purpose')
body(doc,
    'Collects supporting documentary evidence for the application. Three documents are mandatory '
    'before submission. Optional documents can speed up the review process.')

h2(doc, '3.14 Required Documents')
make_table(doc,
    ['Document', 'Format', 'Max Size'],
    [
        ('FPO Registration Certificate', 'PDF, JPG, PNG', '5 MB'),
        ('Bank Statement / Details',     'PDF, JPG, PNG', '5 MB'),
        ('PAN Card',                     'PDF, JPG, PNG', '5 MB'),
    ],
    col_widths=[2.5, 1.5, 0.8],
)

h2(doc, '3.15 Optional Documents')
make_table(doc,
    ['Document', 'Format', 'Max Size'],
    [
        ('GST Certificate', 'PDF, JPG, PNG', '5 MB'),
        ('Annual Report',   'PDF, JPG, PNG', '10 MB'),
        ('Member List',     'PDF or XLSX',   '10 MB'),
        ('Other Document',  'PDF, JPG, PNG', '5 MB'),
    ],
    col_widths=[2.5, 1.5, 0.8],
)

h2(doc, '3.16 Steps to Complete')
numbered(doc, 1, 'Click Upload next to FPO Registration Certificate and select the file.')
numbered(doc, 2, 'Click Upload next to Bank Statement/Details and select the file.')
numbered(doc, 3, 'Click Upload next to PAN Card and select the file.')
numbered(doc, 4, 'Watch the "X of 3 required documents uploaded" counter — it updates as each file is added.')
numbered(doc, 5, 'Upload optional documents if available — this can help speed up approval.')
numbered(doc, 6, 'Once all 3 required documents are uploaded, click Continue to Review →.')

screenshot(doc, 'screenshot_08.png', 'Figure 8 — Upload Documents (Step 6)')

note_box(doc,
    'The Continue to Review → button only becomes active once all 3 required documents are uploaded. '
    'Upload optional documents where available — a complete application speeds up review.')

# ── Step 7 ────────────────────────────────────────────────────────────────────
step_banner(doc, 7, 'Review & Submit', total=7)

h2(doc, '3.17 Purpose')
body(doc,
    'Gives you a final consolidated summary of all information entered across all steps, '
    'confirms that all requirements are met, and allows formal submission of the application.')

h2(doc, '3.18 What You Will See on This Page')
make_table(doc,
    ['Section', 'What It Shows'],
    [
        ('Submission Checklist',
         'A green "All requirements met. Ready to submit!" message when everything is in order. '
         'If any requirement is missing, it is listed here with a link back to the relevant step.'),
        ('Application Summary',
         'A read-only summary of key details: FPO Name, Registration No., District, Total Members, '
         'Primary Commodities, Bank, and IFSC — for a final sanity check.'),
        ('Confirmation Notice',
         '"By submitting, you confirm that all information provided is accurate. The application '
         'will be reviewed by the KAU team." — submitting is a formal declaration.'),
    ],
    col_widths=[2.0, 4.3],
)

h2(doc, '3.19 Steps to Complete')
numbered(doc, 1, 'Check the Submission Checklist — confirm the green "All requirements met" message is showing.')
numbered(doc, 2, 'Review the Application Summary — check all key fields for accuracy.')
numbered(doc, 3, 'Click ← Back if any detail is incorrect and fix it in the relevant step.')
numbered(doc, 4, 'Read the confirmation notice.')
numbered(doc, 5, 'Click Submit Application to send the completed application.')

screenshot(doc, 'screenshot_09.png', 'Figure 9 — Review & Submit (Step 7)')

note_box(doc,
    'Once submitted, you cannot edit the application. Ensure all details are correct before clicking '
    'Submit Application.')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — POST-SUBMISSION
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
h1(doc, '4. Application Status (After Submission)')

body(doc,
    'Immediately after submitting, the platform redirects you to the Application Status page. '
    'Because the platform uses auto-approval, you will typically see an Approved status '
    'immediately after a successful submission.')

h2(doc, '4.1 What You Will See')
make_table(doc,
    ['Element', 'Description'],
    [
        ('Application ID',
         'A unique ID in the format: KAU-FPO-[District Code]-[Year]-[Sequence Number]. '
         'Example: KAU-FPO-TRS-2026-0001. Keep this ID for future reference and support queries.'),
        ('Current Status Card',
         'Shows the approval status clearly — Approved (green), Pending, Rejected, or Info Required.'),
        ('Go to Dashboard button',
         'Appears once the application is Approved. Click to access the full FPO Dashboard.'),
        ('Activity Timeline',
         'A chronological record of every status change with timestamps — from Draft through to Approved.'),
        ('Refresh button',
         'If the status is still pending, use Refresh to check for the latest update without reloading the full page.'),
    ],
    col_widths=[1.8, 4.5],
)

h2(doc, '4.2 Understanding the Application ID Format')
make_table(doc,
    ['Part', 'Meaning', 'Example'],
    [
        ('KAU',             'Kerala Agricultural University',              'KAU'),
        ('FPO',             'Farmer Producer Organization application',    'FPO'),
        ('[District Code]', 'Short code for the FPO\'s district',         'TRS (Thrissur)'),
        ('[Year]',          'Year the application was created',            '2026'),
        ('[Sequence No.]',  'Running number within that district and year','0001'),
    ],
    col_widths=[1.5, 3.0, 1.2],
)

screenshot(doc, 'screenshot_10.png', 'Figure 10 — Application Status page (Approved)')

note_box(doc,
    'Note this Application ID immediately after submission — it is required for any support queries '
    'or future reference. Once approved, click Go to Dashboard → to begin using all platform features.')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — EXPERT DIRECTORY
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
h1(doc, '5. Expert Directory')

body(doc,
    'The Expert Directory is available to approved FPO users from the left-hand navigation menu. '
    'It allows you to find and contact agricultural experts and KAU specialists.')

h2(doc, '5.1 Purpose')
body(doc, 'This page enables you to:')
bullet(doc, 'Browse all available experts in one place.')
bullet(doc, 'Filter experts by professional category.')
bullet(doc, 'Filter by district to find a locally accessible specialist.')
bullet(doc, 'Search for a specific expert by keyword.')
bullet(doc, 'View each expert\'s role, institution, location, and area of expertise.')
bullet(doc, 'Directly initiate contact with an expert.')

h2(doc, '5.2 How to Access')
numbered(doc, 1, 'Log in to the KAU-FPO Platform (FPO Portal).')
numbered(doc, 2, 'In the left-hand navigation menu, click Expert Directory.')

h2(doc, '5.3 Filters Available')
make_table(doc,
    ['Filter', 'How It Works'],
    [
        ('Category Filter',
         'Pill buttons at the top — All Experts, Banker / Financial Advisor, Scientist / Researcher, '
         'Trainer / Extension Worker, etc. Only one category is active at a time.'),
        ('District Filter',
         'Dropdown to narrow results to a specific district — useful if in-person consultation is preferred.'),
        ('Search Bar',
         'Free-text search — find a specific expert by name or keyword related to their expertise.'),
    ],
    col_widths=[1.8, 4.5],
)

h2(doc, '5.4 Typical Workflow')
numbered(doc, 1, 'Click a category button to narrow the list by specialist type, or keep All Experts selected.')
numbered(doc, 2, 'Optionally select a district from the dropdown.')
numbered(doc, 3, 'Use the Search bar if you already know the expert\'s name or a specific keyword.')
numbered(doc, 4, 'Review each expert card — check designation, institution, location, and expertise.')
numbered(doc, 5, 'Click Contact Expert on the relevant card to initiate contact.')

screenshot(doc, 'screenshot_03.png', 'Figure 11 — Expert Directory page')

note_box(doc,
    'Experts within the same category can have very different specialisations — read the Expertise field '
    'carefully. Combine the District filter with a category filter for the most targeted results.')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
h1(doc, '6. Troubleshooting')

make_table(doc,
    ['Issue', 'Likely Cause', 'What to Do'],
    [
        ('Check Eligibility button is blocked',
         'Total Farmer Members is below 10, or not all requirement checkboxes are ticked.',
         'Ensure member count ≥ 10 and all three requirements are checked.'),
        ('OTP not received (phone)',
         'Incorrect phone number, or SMS delivery delay.',
         'Check the number displayed, wait 60 seconds, then click Resend.'),
        ('OTP not received (email)',
         'OTP went to spam, or incorrect email address.',
         'Check spam/junk folder. If wrong email was entered, go back to Step 2 and correct it.'),
        ('OTP expired',
         'More than 10 minutes passed before entering the code.',
         'Click Resend or Send OTP again to get a fresh code.'),
        ('Block/Taluk dropdown is disabled',
         'District has not been selected yet.',
         'Select District first — Block/Taluk unlocks automatically.'),
        ('Document upload not working',
         'File type or size is not accepted.',
         'Required documents: PDF, JPG, PNG up to 5 MB. Member List: PDF or XLSX up to 10 MB.'),
        ('"All requirements met" message not showing on Step 7',
         'One or more mandatory items are incomplete.',
         'Review the checklist on the Review & Submit page — it lists exactly what is missing.'),
        ('Forgot password after account creation',
         'N/A',
         'Click Forgot Password on the login page and follow the reset link sent to your email.'),
    ],
    col_widths=[1.8, 2.2, 2.3],
)


# ══════════════════════════════════════════════════════════════════════════════
# CLOSING
# ══════════════════════════════════════════════════════════════════════════════

spacer(doc, 20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_bottom_border(p, sz='4', color='AAAAAA')
r = p.add_run(
    'Kerala Agricultural University  |  Communication Centre, Mannuthy, Thrissur — 680651  |  '
    'AI-Based Digital Platform for KAU-FPO Linkage Programme  |  fpolinkage.kau.in'
)
r.font.italic     = True
r.font.size       = Pt(9)
r.font.name       = 'Cambria'
r.font.color.rgb  = GREY


# ─── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
