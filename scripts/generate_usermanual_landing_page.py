"""
User Manual Generator — Landing Page & Public Portal
KAU-FPO Platform — house style (same as RCD and FPO Registration manual)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paths ────────────────────────────────────────────────────────────────────
KAU_LOGO    = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/KAU Emblem 0.5x0.75.jpg'
SCREENSHOTS = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/Documents/Usermanual/screenshots/landing_page'
OUTPUT      = '/home/athul_dasp/Desktop/AGRI-THRISSUR/kau-fpo-backend/Documents/Usermanual/KAU_FPO_UserManual_LandingPage_v1.0.docx'

# ─── Colors ───────────────────────────────────────────────────────────────────
NAVY_HEX       = '1F3864'
DARK_NAVY      = RGBColor(0x1F, 0x38, 0x64)
ORANGE         = RGBColor(0xE8, 0x6C, 0x1A)
WHITE          = RGBColor(0xFF, 0xFF, 0xFF)
GREY           = RGBColor(0x60, 0x60, 0x60)
BLACK          = RGBColor(0x00, 0x00, 0x00)
LIGHT_GREY_HEX = 'F2F2F2'


# ─── XML Helpers ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_margins(cell, top=60, start=120, bottom=60, end=120):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'),    str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def remove_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'none')
        b.set(qn('w:sz'),    '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_bottom_border(paragraph, color=NAVY_HEX, sz='8'):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def add_page_number(run):
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    run._r.append(instrText)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)


def set_col_width(cell, width_inches):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(width_inches * 1440)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def set_table_width(table, width_cm):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW  = OxmlElement('w:tblW')
    twips = int(width_cm * 567)
    tblW.set(qn('w:w'),    str(twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


# ─── Header / Footer ──────────────────────────────────────────────────────────

def build_header(header_obj, title='User Manual — Landing Page & Public Portal'):
    for p in header_obj.paragraphs:
        p._element.getparent().remove(p._element)

    tbl = header_obj.add_table(rows=1, cols=3, width=Inches(6.5))
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)

    c0 = tbl.cell(0, 0)
    set_col_width(c0, 1.15)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after  = Pt(0)
    r0 = p0.add_run()
    r0.add_picture(KAU_LOGO, width=Cm(0.8), height=Cm(0.49))

    c1 = tbl.cell(0, 1)
    set_col_width(c1, 3.82)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after  = Pt(0)
    r1 = p1.add_run(title)
    r1.font.bold      = True
    r1.font.size      = Pt(9)
    r1.font.color.rgb = DARK_NAVY

    c2 = tbl.cell(0, 2)
    set_col_width(c2, 1.30)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run('KAU-FPO Platform')
    r2.font.size      = Pt(9)
    r2.font.color.rgb = DARK_NAVY

    div = header_obj.add_paragraph()
    div.paragraph_format.space_before = Pt(2)
    div.paragraph_format.space_after  = Pt(0)
    add_bottom_border(div)


def build_first_page_header(header_obj):
    for p in header_obj.paragraphs:
        p._element.getparent().remove(p._element)
    header_obj.add_paragraph()


def build_footer(footer_obj):
    for p in footer_obj.paragraphs:
        p._element.getparent().remove(p._element)

    tbl = footer_obj.add_table(rows=1, cols=2, width=Inches(6.5))
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)

    cl = tbl.cell(0, 0)
    set_col_width(cl, 4.5)
    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rl = pl.add_run('Kerala Agricultural University  │  fpolinkage.kau.in')
    rl.font.size      = Pt(8)
    rl.font.color.rgb = GREY

    cr = tbl.cell(0, 1)
    set_col_width(cr, 1.77)
    pr = cr.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = pr.add_run('Page ')
    rr.font.size      = Pt(8)
    rr.font.color.rgb = GREY
    rn = pr.add_run()
    rn.font.size      = Pt(8)
    rn.font.color.rgb = GREY
    add_page_number(rn)


# ─── Content Helpers ──────────────────────────────────────────────────────────

def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold      = True
    r.font.size      = Pt(16)
    r.font.name      = 'Calibri'
    r.font.color.rgb = DARK_NAVY
    add_bottom_border(p)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold      = True
    r.font.size      = Pt(12)
    r.font.name      = 'Calibri'
    r.font.color.rgb = ORANGE
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.bold      = True
    r.font.size      = Pt(11)
    r.font.name      = 'Calibri'
    r.font.color.rgb = BLACK
    return p


def body(doc, text, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Cambria'
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.font.name = 'Cambria'
    return p


def numbered(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Pt(18)
    r1 = p.add_run(f'{number}.  ')
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Cambria'
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.name = 'Cambria'
    return p


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    return p


def make_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    tbl.style = 'Table Grid'

    for j, h_text in enumerate(headers):
        c = tbl.rows[0].cells[j]
        set_cell_bg(c, NAVY_HEX)
        set_cell_margins(c)
        if col_widths:
            set_col_width(c, col_widths[j])
        p = c.paragraphs[0]
        r = p.add_run(h_text)
        r.font.bold      = True
        r.font.color.rgb = WHITE
        r.font.size      = Pt(10)
        r.font.name      = 'Calibri'

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            c = tbl.rows[i + 1].cells[j]
            set_cell_margins(c)
            if col_widths:
                set_col_width(c, col_widths[j])
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = 'Cambria'

    spacer(doc, 8)
    return tbl


def screenshot(doc, filename, caption=None):
    import os
    path = os.path.join(SCREENSHOTS, filename)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run()
    r.add_picture(path, width=Cm(14))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        cr = cp.add_run(caption)
        cr.font.italic    = True
        cr.font.size      = Pt(9)
        cr.font.name      = 'Cambria'
        cr.font.color.rgb = GREY


def note_box(doc, text):
    spacer(doc, 4)
    tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)
    c = tbl.cell(0, 0)
    set_cell_bg(c, LIGHT_GREY_HEX)
    set_cell_margins(c, top=100, start=160, bottom=100, end=160)
    p = c.paragraphs[0]
    r1 = p.add_run('Note: ')
    r1.font.bold      = True
    r1.font.size      = Pt(10)
    r1.font.name      = 'Calibri'
    r1.font.color.rgb = DARK_NAVY
    r2 = p.add_run(text)
    r2.font.size      = Pt(10)
    r2.font.name      = 'Cambria'
    spacer(doc, 6)


def section_banner(doc, section_title):
    tbl = doc.add_table(rows=1, cols=1)
    remove_table_borders(tbl)
    set_table_width(tbl, 16.51)
    c = tbl.cell(0, 0)
    set_cell_bg(c, NAVY_HEX)
    set_cell_margins(c, top=120, start=160, bottom=120, end=160)
    p = c.paragraphs[0]
    r = p.add_run(section_title)
    r.font.bold      = True
    r.font.size      = Pt(12)
    r.font.name      = 'Calibri'
    r.font.color.rgb = WHITE
    spacer(doc, 6)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Cambria'
style.font.size = Pt(11)

sec = doc.sections[0]
sec.page_width        = Cm(21)
sec.page_height       = Cm(29.7)
sec.left_margin       = Cm(2.54)
sec.right_margin      = Cm(2.54)
sec.top_margin        = Cm(2.0)
sec.bottom_margin     = Cm(2.0)
sec.header_distance   = Cm(1.0)
sec.footer_distance   = Cm(1.0)
sec.different_first_page_header_footer = True

build_header(sec.header)
build_footer(sec.footer)
build_first_page_header(sec.first_page_header)
build_footer(sec.first_page_footer)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after  = Pt(12)
p.add_run().add_picture(KAU_LOGO, width=Cm(9.14))

banner_tbl = doc.add_table(rows=1, cols=1)
banner_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_width(banner_tbl, 16.51)
bc = banner_tbl.cell(0, 0)
set_cell_bg(bc, NAVY_HEX)
set_cell_margins(bc, top=160, start=200, bottom=160, end=200)
bp = bc.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
br = bp.add_run('USER MANUAL — LANDING PAGE & PUBLIC PORTAL')
br.font.bold      = True
br.font.size      = Pt(20)
br.font.name      = 'Calibri'
br.font.color.rgb = WHITE

spacer(doc, 14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('KERALA AGRICULTURAL UNIVERSITY')
r.font.bold      = True
r.font.size      = Pt(16)
r.font.name      = 'Calibri'
r.font.color.rgb = DARK_NAVY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run('Communication Centre, Mannuthy, Thrissur — 680651')
r.font.size      = Pt(11)
r.font.name      = 'Cambria'
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('AI-Based Digital Platform for KAU-FPO Linkage Programme')
r.font.bold      = True
r.font.size      = Pt(13)
r.font.name      = 'Calibri'
r.font.color.rgb = ORANGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(16)
r = p.add_run('User Manual — Landing Page & Public Portal')
r.font.bold      = True
r.font.size      = Pt(13)
r.font.name      = 'Calibri'
r.font.color.rgb = ORANGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Version 1.0  |  Kerala Agricultural University')
r.font.size      = Pt(11)
r.font.name      = 'Cambria'
r.font.color.rgb = GREY

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — DOCUMENT INFORMATION
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '1. Document Information')

make_table(doc,
    headers=['Field', 'Details'],
    rows=[
        ('Application Name',  'AI-Based Digital Platform for KAU-FPO Linkage Programme'),
        ('Module',            'Landing Page & Public Portal'),
        ('Version',           '1.0'),
        ('Prepared By',       'Kerala Agricultural University'),
        ('Date',              'August 2026'),
        ('Audience',          'General Public, FPO Managers, Government Officials, Experts'),
        ('Platform URL',      'https://fpolinkage.kau.in'),
    ],
    col_widths=[2.2, 4.07],
)

h2(doc, '1.1  Purpose')
body(doc,
    'This manual describes the publicly accessible pages of the KAU-FPO Linkage Platform. '
    'It covers the Landing Page, the More Info portal, the About Us page, News & Announcements, '
    'the FAQ page, and the Contact Us page. No login is required to access any of the pages '
    'described in this document.')

h2(doc, '1.2  System Requirements')
make_table(doc,
    headers=['Component', 'Minimum Requirement'],
    rows=[
        ('Web Browser',      'Google Chrome 110+, Mozilla Firefox 110+, Microsoft Edge 110+, Safari 16+'),
        ('Internet',         'Broadband connection (1 Mbps or above recommended)'),
        ('Screen Resolution','1280 × 720 or higher'),
        ('JavaScript',       'Must be enabled in browser settings'),
        ('Mobile',           'Responsive — works on smartphones and tablets'),
    ],
    col_widths=[2.2, 4.07],
)

note_box(doc,
    'The platform is fully responsive and works on mobile devices. For the best experience, '
    'use a desktop or laptop browser with a minimum resolution of 1280 × 720 pixels.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '2. Overview')
body(doc,
    'The KAU-FPO Linkage Platform is an AI-assisted digital platform developed by Kerala '
    'Agricultural University (KAU) to support Farmer Producer Organisations (FPOs) across '
    'Kerala. The public-facing portal serves as the primary entry point for all visitors — '
    'FPO managers, government officials, agricultural experts, and the general public.')

body(doc,
    'The platform is accessible at https://fpolinkage.kau.in. All pages described in this '
    'manual are publicly visible and do not require user registration or login.')

h2(doc, '2.1  Pages Covered in This Manual')
make_table(doc,
    headers=['Page', 'URL Path', 'Audience'],
    rows=[
        ('Landing Page',         '/',                'All visitors'),
        ('More Info',            '/more-info',       'All visitors'),
        ('About Us',             '/about',           'All visitors'),
        ('News & Announcements', '/news',            'All visitors'),
        ('FAQ',                  '/faq',             'All visitors'),
        ('Contact Us',           '/contact',         'All visitors'),
    ],
    col_widths=[2.0, 1.8, 2.47],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — LANDING PAGE
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '3. Landing Page')
body(doc,
    'The Landing Page is the home page of the platform. It provides a comprehensive '
    'overview of the KAU-FPO Linkage Programme and offers quick access to all major '
    'sections. Visitors can navigate to any part of the site from this page.')

screenshot(doc, 'screenshot_00.png', 'Figure 1 — Landing Page (full view)')

h2(doc, '3.1  Top Navigation Bar')
body(doc,
    'The top navigation bar is fixed at the top of every page and provides links to all '
    'major sections of the platform.')

make_table(doc,
    headers=['Navigation Item', 'Destination'],
    rows=[
        ('Home',       'Returns to the Landing Page'),
        ('Sign In',    'Opens the FPO Manager / Admin login page'),
        ('Get Started','Opens the FPO Registration page'),
        ('Pages',      'Dropdown menu — About Us, News, FAQ, Contact'),
        ('More Info',  'Opens the Agriculture Information Technologies & Services page'),
        ('Language',   'Toggle between English and Malayalam'),
    ],
    col_widths=[2.2, 4.07],
)

note_box(doc,
    'The navigation bar includes the KAU logo and the Directorate of Extension logo on the '
    'left side. Click either logo at any time to return to the home page.')

h2(doc, '3.2  Hero Banner')
body(doc,
    'The hero section is the first section visible when the page loads. It contains a '
    'prominent headline — "Empowering Farmers through FPO Linkage" — along with a brief '
    'description of the programme and a "Get Started" call-to-action button.')

bullet(doc, 'Click "Get Started" to begin the FPO registration process.')
bullet(doc, 'A Quick Access Panel on the right displays key statistics and platform highlights.')

h2(doc, '3.3  Services Offered')
body(doc,
    'The "Services Offered" section highlights the four core services available through '
    'the platform. Each service card links to the relevant feature inside the platform.')

make_table(doc,
    headers=['Service', 'Description'],
    rows=[
        ('Tier Classification',    'Automated assessment of FPO performance across 5 domains with a tier badge (Tier 1–4)'),
        ('Project Reports (DPR)',  'AI-assisted Detailed Project Report generation for bank loans and government schemes'),
        ('Market Linkage',         'Product listing, buyer directory, and ONDC marketplace integration'),
        ('Schemes & Subsidies',    'Centralised hub for government schemes, subsidies, and application links'),
    ],
    col_widths=[2.2, 4.07],
)

h2(doc, '3.4  Quick Links')
body(doc,
    'The Quick Links section provides direct links to 10 external agricultural portals and '
    'government services relevant to FPOs in Kerala. Clicking any tile opens the external '
    'website in a new browser tab.')

bullet(doc, 'Links include: KAU official site, Agri Department Kerala, e-NAM, NABARD, SFAC, and more.')
bullet(doc, 'All links open in a new tab — the KAU-FPO platform stays open in the background.')

h2(doc, '3.5  Our Team')
body(doc,
    'The "Our Team" section displays the leadership team responsible for the KAU-FPO '
    'Linkage Programme, including their names, designations, and photographs.')

h2(doc, '3.6  News & Announcements')
body(doc,
    'The latest news and announcements from KAU are shown in this section. Each card '
    'displays a headline, date, and brief summary. Click any card to read the full '
    'announcement on the News & Announcements page.')

h2(doc, '3.7  Gallery')
body(doc,
    'The Gallery section displays recent photographs from KAU-FPO events, training '
    'sessions, and field visits. Images are organised in albums and can be viewed by '
    'clicking on any photo.')

h2(doc, '3.8  KAU-FPO Linkage in Numbers')
body(doc,
    'This section displays live platform statistics including the total number of '
    'registered FPOs, active members, districts covered, and expert consultants '
    'available on the platform.')

h2(doc, '3.9  Contact / Send a Message')
body(doc,
    'Visitors can send a message directly from the Landing Page without navigating away. '
    'Fill in the Name, Email, Phone (optional), Subject, and Message fields, then click '
    '"Get in Touch". The message is delivered to the KAU Communication Centre.')

h2(doc, '3.10  Partners Strip')
body(doc,
    'The Partners section displays logos of government and institutional partners '
    'supporting the KAU-FPO Linkage Programme.')

h2(doc, '3.11  Footer')
body(doc,
    'The footer appears at the bottom of every page and contains:')
bullet(doc, 'Explore links — About Us, Meet Our Team, News & Events, Contact Us')
bullet(doc, 'Contact Info — Address (Directorate of Extension, Mannuthy PO, Thrissur — 680651), Email, Phone')
bullet(doc, 'Copyright notice')
bullet(doc, 'Support link')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — MORE INFO PAGE
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '4. More Info — Agriculture Information Technologies & Services')
body(doc,
    'The More Info page is accessible from the top navigation bar (More Info) or by '
    'clicking the "More Info" button on the Landing Page. It presents 18 agricultural '
    'information technology services provided by KAU and its partner organisations.')

screenshot(doc, 'screenshot_05.png', 'Figure 2 — More Info page — Agriculture Information Technologies & Services')

body(doc,
    'Each service is shown as a card with an icon, name, and a brief description. '
    'Clicking a card opens the external service website or the relevant KAU portal in '
    'a new browser tab.')

h2(doc, '4.1  Available Services')
make_table(doc,
    headers=['Service', 'Description'],
    rows=[
        ('Crop Husbandry',        'Crop production guidelines and best practices from KAU'),
        ('Animal Husbandry',      'Livestock management resources and advisory services'),
        ('Fisheries',             'Fishery and aquaculture information portal'),
        ('Forestry',              'Forest management and agroforestry resources'),
        ('KAU Moodle',            'KAU e-learning platform for agricultural education'),
        ('e-Crop Doctor',         'AI-based plant disease diagnosis tool'),
        ('Seed Rate & Spacing',   'Seed rate calculator for major crops'),
        ('Fertulator',            'Fertilizer requirement calculator based on soil and crop type'),
        ('Agri Almanac',          'Agricultural calendar — planting dates, weather advisories'),
        ('Farm Machinery',        'Information on farm mechanisation and machinery rental'),
        ('Agri Enterprises',      'Agri-business development resources and case studies'),
        ('e-BIO',                 'Bio-agent information and ordering portal'),
        ('Knowledge Bank',        'Repository of agricultural research publications and extension materials'),
        ('Market Intelligence',   'Commodity price trends and market analysis'),
        ('Agri Videos',           'KAU training and demonstration video library'),
        ('Weather Advisory',      'District-wise weather forecasts and agro-advisories'),
        ('Kerala Directory',      'Directory of agricultural institutions in Kerala'),
        ('Library',               'KAU digital library and e-resource portal'),
    ],
    col_widths=[2.2, 4.07],
)

note_box(doc,
    'All services on the More Info page link to external websites or KAU portals. '
    'An active internet connection is required. Some services may require a separate '
    'KAU account or registration on the respective portal.')

body(doc,
    'The More Info page also includes a sub-section titled "How to Register FPO" that '
    'outlines the legal registration steps required before using the KAU-FPO Platform. '
    'This section is intended for FPO promoters who are setting up a new Farmer Producer '
    'Organisation.')

screenshot(doc, 'screenshot_02.png', 'Figure 3 — How to Register FPO sub-page')

h2(doc, '4.2  How to Register FPO — Legal Process Overview')
body(doc,
    'Before registering on the KAU-FPO Platform, an FPO must complete legal registration '
    'as a Producer Company or Cooperative Society. The platform provides a two-phase guide:')

h3(doc, 'Phase I — Legal Registration')
numbered(doc, 1, 'Mobilise the required number of farmer members (minimum 10 members required).')
numbered(doc, 2, 'Obtain a Digital Signature Certificate (DSC) for the proposed directors.')
numbered(doc, 3, 'Apply for Director Identification Number (DIN) for each director.')
numbered(doc, 4, 'Reserve the Company Name through the Ministry of Corporate Affairs (MCA) portal.')
numbered(doc, 5, 'Prepare the Memorandum and Articles of Association (MoA/AoA).')
numbered(doc, 6, 'Submit the SPICe+ Incorporation Application.')
numbered(doc, 7, 'Obtain the Certificate of Incorporation from the Registrar of Companies.')

h3(doc, 'Phase II — Post-Incorporation Setup')
body(doc,
    'After legal incorporation, complete the following to activate the FPO on the KAU-FPO '
    'Platform:')
bullet(doc, 'Open a bank account in the name of the FPO.')
bullet(doc, 'Obtain a PAN card for the FPO entity.')
bullet(doc, 'Apply for GST registration if applicable.')
bullet(doc, 'Register on the KAU-FPO Platform using the FPO Registration wizard.')

note_box(doc,
    'Legal registration must be completed before applying on the KAU-FPO Platform. '
    'The platform requires a valid Registration Number or CIN (for Producer Companies) '
    'during the Step 1 — Basic Information stage of registration.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — ABOUT US PAGE
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '5. About Us')
body(doc,
    'The About Us page provides an overview of the KAU-FPO Linkage Programme, its mission, '
    'vision, and the partner institutions involved. It is accessible from the Pages dropdown '
    'in the top navigation bar.')

screenshot(doc, 'screenshot_01.png', 'Figure 4 — About Us page')

h2(doc, '5.1  About the KAU-FPO Linkage Programme')
body(doc,
    'The KAU-FPO Linkage Programme is a Government of Kerala initiative implemented by '
    'Kerala Agricultural University (KAU) through its Directorate of Extension. The '
    'programme aims to strengthen Farmer Producer Organisations across Kerala by providing '
    'digital tools for registration, project planning, market access, and capacity building.')

h2(doc, '5.2  The Mission')
body(doc,
    'The mission of the programme is to empower FPOs and their member farmers through '
    'technology-enabled services. The three core objectives are:')
numbered(doc, 1,
    'Development of FPOs — supporting legally registered FPOs in obtaining government '
    'schemes and institutional credit through AI-assisted project reports.')
numbered(doc, 2,
    'Capacity Building — training FPO managers, providing expert access, and enabling '
    'continuous skill development through the platform.')
numbered(doc, 3,
    'Facility Assistance — connecting FPOs with agricultural input suppliers, output '
    'market buyers, and financial institutions.')

h2(doc, '5.3  The Vision')
body(doc,
    'The vision of the KAU-FPO Linkage Programme is to establish Kerala as a model state '
    'for farmer collectivisation, where every FPO has access to digital tools, institutional '
    'markets, and professional expertise to achieve sustainable agricultural growth.')

h2(doc, '5.4  Partners')
body(doc,
    'The programme is implemented in partnership with the Directorate of Extension, '
    'Kerala Agricultural University, and supported by government departments, NABARD, '
    'SFAC, and other agricultural institutions in Kerala.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — NEWS & ANNOUNCEMENTS
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '6. News & Announcements')
body(doc,
    'The News & Announcements page displays the latest updates, circulars, and news '
    'relevant to FPOs and the KAU-FPO Linkage Programme. It is accessible from the '
    'Pages dropdown in the top navigation bar or from news cards on the Landing Page.')

h2(doc, '6.1  Page Layout')
bullet(doc, 'Two tabs are available: Announcements and News.')
bullet(doc, 'Click the "Announcements" tab to view official circulars, notifications, and KAU communications.')
bullet(doc, 'Click the "News" tab to view news articles related to agricultural developments and FPO activities.')

h2(doc, '6.2  Reading an Announcement or News Article')
numbered(doc, 1, 'Navigate to the News & Announcements page from the Pages dropdown in the navigation bar.')
numbered(doc, 2, 'Select the "Announcements" or "News" tab as required.')
numbered(doc, 3, 'Scroll through the content cards to find the relevant item.')
numbered(doc, 4, 'Click on a card to open the full announcement or news article.')
numbered(doc, 5, 'Use the browser back button to return to the list.')

note_box(doc,
    'New announcements are added by the KAU Communication Centre. All content is '
    'managed by KAU administrators and is verified before publication.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — FAQ PAGE
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '7. Frequently Asked Questions (FAQ)')
body(doc,
    'The FAQ page answers common questions about the KAU-FPO Linkage Programme, FPO '
    'registration, government schemes, and platform usage. It is accessible from the '
    'Pages dropdown in the top navigation bar.')

screenshot(doc, 'screenshot_03.png', 'Figure 5 — FAQ page')

h2(doc, '7.1  Using the FAQ Page')
numbered(doc, 1, 'Navigate to the FAQ page from the Pages dropdown in the top navigation bar.')
numbered(doc, 2,
    'Use the Search bar at the top of the page to search for a specific question. '
    'Type keywords and the list will filter automatically.')
numbered(doc, 3,
    'Use the Category Filter buttons to narrow results by topic: All, FPO General, '
    'Schemes, or Platform Usage.')
numbered(doc, 4, 'Click on any question in the accordion list to expand and read the answer.')
numbered(doc, 5, 'Click the same question again to collapse the answer.')
numbered(doc, 6, 'If more questions are available, click the "Load More" button at the bottom to view additional FAQs.')

h2(doc, '7.2  FAQ Categories')
make_table(doc,
    headers=['Category', 'Topics Covered'],
    rows=[
        ('All',            'All questions across every category'),
        ('FPO General',    'What is an FPO, membership requirements, share capital, Board of Directors, CEO role, AGM'),
        ('Schemes',        'Government schemes for FPOs, subsidy eligibility, application process'),
        ('Platform Usage', 'How to register, how to use the DPR wizard, how to access expert directory'),
    ],
    col_widths=[2.0, 4.27],
)

h2(doc, '7.3  Commonly Asked Questions')
make_table(doc,
    headers=['Question', 'Short Answer'],
    rows=[
        ('What is a Farmer Producer Organisation (FPO)?',
         'An FPO is a legal entity formed by primary producers (farmers, fishers, etc.) for the purpose of economic activities related to their produce.'),
        ('What is a Producer Company?',
         'A Producer Company is a type of company registered under the Companies Act that allows farmers to collectively manage production, procurement, processing, and sale.'),
        ('Who can become a member of an FPO?',
         'Any primary producer (farmer, fisher, artisan) can become a member by purchasing shares in the FPO.'),
        ('How many members are required to form an FPO?',
         'A minimum of 10 farmer members are required to register an FPO on the KAU-FPO Platform.'),
        ('What are the benefits of joining an FPO?',
         'Members gain access to collective bargaining, shared resources, government schemes, expert advisory services, and credit facilities.'),
    ],
    col_widths=[3.0, 3.27],
)

note_box(doc,
    'If your question is not answered in the FAQ, use the Contact Us page to send a '
    'message to the KAU Communication Centre. A team member will respond within 2 working days.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — CONTACT US PAGE
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '8. Contact Us')
body(doc,
    'The Contact Us page allows visitors to send a direct message to the KAU '
    'Communication Centre. It also displays the official contact information of the '
    'Directorate of Extension, Kerala Agricultural University.')

screenshot(doc, 'screenshot_04.png', 'Figure 6 — Contact Us page')

h2(doc, '8.1  Sending a Message')
numbered(doc, 1, 'Navigate to the Contact Us page from the Pages dropdown in the top navigation bar.')
numbered(doc, 2, 'Fill in the following fields in the "Send us a Message" form:')

make_table(doc,
    headers=['Field', 'Required', 'Description'],
    rows=[
        ('Name',    'Yes', 'Your full name'),
        ('Email',   'Yes', 'Your email address — a copy of the message will be sent here'),
        ('Phone',   'No',  'Your contact phone number (optional)'),
        ('Subject', 'Yes', 'Brief subject or topic of your message'),
        ('Message', 'Yes', 'Your full message or query'),
    ],
    col_widths=[1.5, 1.0, 3.77],
)

numbered(doc, 3, 'Click the "Get in Touch" button to submit your message.')
numbered(doc, 4, 'A confirmation message will appear on screen when your message is sent successfully.')

note_box(doc,
    'A copy of your message will be sent to the email address you provide. '
    'The KAU Communication Centre will respond within 2 working days.')

h2(doc, '8.2  Contact Information')
make_table(doc,
    headers=['Type', 'Details'],
    rows=[
        ('Address', 'Directorate of Extension, Mannuthy PO, Thrissur — 680651, Kerala'),
        ('Email',   'Visit https://fpolinkage.kau.in/contact for the current email address'),
        ('Phone',   'Visit https://fpolinkage.kau.in/contact for the current phone numbers'),
        ('Website', 'https://fpolinkage.kau.in'),
    ],
    col_widths=[1.5, 4.77],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════════════

h1(doc, '9. Troubleshooting')

make_table(doc,
    headers=['Problem', 'Likely Cause', 'Solution'],
    rows=[
        ('Page does not load',
         'No internet connection or server maintenance',
         'Check your internet connection. If the problem persists, try again after a few minutes.'),
        ('Images not displaying',
         'Slow internet connection or browser cache issue',
         'Refresh the page (Ctrl+R / Cmd+R). Clear browser cache if the issue continues.'),
        ('Language not switching',
         'Page has not reloaded after language change',
         'After clicking the language toggle, wait for the page to reload fully.'),
        ('Contact form not submitting',
         'Required fields left empty or invalid email format',
         'Ensure all required fields (Name, Email, Subject, Message) are filled in correctly.'),
        ('External service link not opening',
         'The external website may be temporarily unavailable',
         'Try the link again after some time. The issue is with the external website, not the KAU-FPO Platform.'),
        ('"Get Started" button leads to an error',
         'Registration server temporarily unavailable',
         'Try again after a few minutes. If the error persists, contact KAU via the Contact Us page.'),
    ],
    col_widths=[1.7, 2.0, 2.57],
)

note_box(doc,
    'For technical support, use the Contact Us page at https://fpolinkage.kau.in/contact '
    'or call the KAU Communication Centre directly. Please include a description of the '
    'issue and the page where it occurred.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# CLOSING
# ══════════════════════════════════════════════════════════════════════════════

spacer(doc, 20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_bottom_border(p, sz='4')

spacer(doc, 8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'Kerala Agricultural University  |  Communication Centre, Mannuthy, Thrissur — 680651  |  '
    'AI-Based Digital Platform for KAU-FPO Linkage Programme  |  fpolinkage.kau.in'
)
r.font.size      = Pt(9)
r.font.name      = 'Cambria'
r.font.color.rgb = GREY


# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
